"""Assemble SCBE non-provisional specification DOCX draft.

This script assembles the current drafting pieces into the USPTO-style order:
Title, Cross-Reference, Background, Summary, Brief Description of Drawings,
Detailed Description, Claims, Abstract.

It does not file anything with the USPTO and does not complete ADS or
oath/declaration forms.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
TITLE = "System and Method for Hyperbolic Geometry-Based Authorization with Topological Control-Flow Integrity"
APPLICATION = "63/961,403"
FILING_DATE = "January 15, 2026"

SOURCE_DETAILED = ROOT / "docs" / "PATENT_DETAILED_DESCRIPTION.md"
SOURCE_CLAIMS = ROOT / "docs" / "legal" / "PATENT_CLAIMS_EXPANDED_v2.md"
SOURCE_ABSTRACT = ROOT / "docs" / "legal" / "PATENT_ABSTRACT_v1.md"
OUT_DIR = ROOT / "docs" / "legal" / "patent-workbench" / "assembled"
OUT_DOCX = OUT_DIR / "SCBE_NONPROVISIONAL_SPEC_DRAFT_v1.docx"
OUT_MD = OUT_DIR / "SCBE_NONPROVISIONAL_SPEC_DRAFT_v1.md"
OUT_JSON = OUT_DIR / "SCBE_NONPROVISIONAL_SPEC_DRAFT_v1.manifest.json"


FIGURE_DESCRIPTIONS = [
    (
        "FIG. 1 is a block diagram of a fourteen-layer authorization pipeline from complex context "
        "ingestion through risk decision gating and telemetry."
    ),
    (
        "FIG. 2 is a graph of a harmonic wall cost function showing nonlinear governance cost "
        "as a function of hyperbolic distance."
    ),
    (
        "FIG. 3 is a Poincare ball cross-section showing concentric security zones, realm centers, "
        "a session centroid, and example authorized and adversarial trajectories."
    ),
    "FIG. 4 is a six-axis semantic weighting diagram showing phi-scaled context channels and orthogonal semantic axes.",
    (
        "FIG. 5 is a deferred-authorization container flow showing multi-predicate gating "
        "and fail-to-noise output on predicate failure."
    ),
    (
        "FIG. 6 is a cheapest-reject-first pre-filter stack showing ordered governance filters "
        "before expensive execution or cryptographic operations."
    ),
    "FIG. 7 is a runtime decision gate diagram showing allow, quarantine, escalate, and deny routing paths.",
    (
        "FIG. 8 is a bijective tamper detection flow showing encode/decode round-trip, "
        "AST canonicalization, identifier canonicality, and governance escalation."
    ),
    (
        "FIG. 9 is a system deployment architecture diagram showing client, API, command-line, "
        "agent-bus, governance, receipt, and audit components."
    ),
]


SUMMARY_PARAGRAPHS = [
    (
        "The disclosed system provides a computer-implemented authorization and "
        "runtime-governance framework for computational actions, including "
        "artificial-intelligence and agentic-tool actions. A request or action is "
        "encoded into a context representation, projected into a bounded nonlinear "
        "geometric domain, and evaluated against a trusted or session reference "
        "state using a distance or drift measure."
    ),
    (
        "In embodiments, the bounded nonlinear geometric domain comprises a "
        "Poincare ball model of hyperbolic space, and governance cost is computed "
        "from hyperbolic distance. The governance cost may be combined with "
        "semantic weighting, temporal drift, spectral coherence, spin coherence, "
        "identifier canonicality, or bijective tamper-detection signals to produce "
        "a composite risk value."
    ),
    (
        "A runtime decision gate emits a governance decision such as allow, review, "
        "quarantine, or deny. The decision controls whether the computational "
        "action is executed, restricted, held for review, or blocked. In some "
        "embodiments, the gate persists session state so that a restarted process "
        "continues from a restored trajectory rather than from a cold start."
    ),
    (
        "Additional embodiments include semantic weighting axes, fail-to-noise "
        "responses, quarantine containment, post-quantum decision receipts, and "
        "tamper or canonicality checks that compare encoded and decoded forms of "
        "source-code or identifier-containing inputs before execution."
    ),
]


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def section(text: str, heading: str, next_headings: list[str]) -> str:
    pattern = re.compile(rf"^## {re.escape(heading)}\s*$", re.MULTILINE)
    match = pattern.search(text)
    if not match:
        raise ValueError(f"Missing section: {heading}")
    start = match.end()
    end = len(text)
    for nxt in next_headings:
        nxt_match = re.search(rf"^## {re.escape(nxt)}\s*$", text[start:], re.MULTILINE)
        if nxt_match:
            end = min(end, start + nxt_match.start())
    return clean_markdown(text[start:end]).strip()


def clean_markdown(text: str) -> str:
    text = text.replace("&nbsp;", " ")
    text = text.replace("—", "-")
    text = text.replace("–", "-")
    text = text.replace("≈", "approximately")
    text = text.replace("φ", "phi")
    text = text.replace("∈", "in")
    text = text.replace("→", "->")
    text = text.replace("≥", ">=")
    text = text.replace("≤", "<=")
    text = text.replace("‖", "||")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"^\s*---\s*$", "", text, flags=re.MULTILINE)
    return text


def extract_background(detailed: str) -> str:
    return section(detailed, "BACKGROUND OF THE INVENTION", ["BRIEF DESCRIPTION OF THE DRAWINGS"])


def extract_detailed_description(detailed: str) -> str:
    return section(detailed, "DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS", ["CLAIMS"])


def extract_abstract(abstract_md: str) -> str:
    body = abstract_md.split("---", 2)
    if len(body) >= 3:
        candidate = body[1]
    else:
        candidate = abstract_md
    lines = []
    for line in candidate.splitlines():
        if line.strip().startswith("#"):
            continue
        if line.strip().startswith("Word count:"):
            continue
        if not line.strip():
            if lines and lines[-1] != "":
                lines.append("")
            continue
        lines.append(line.strip())
    return clean_markdown("\n".join(lines)).strip()


def extract_claims(claims_md: str) -> tuple[str, int]:
    start = claims_md.index("### CLAIM FAMILY A")
    end = claims_md.index("## 4. Independent-Claim")
    block = claims_md[start:end]
    out: list[str] = ["What is claimed is:"]
    claim_count = 0
    in_note = False
    for raw in block.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            if out and out[-1] != "":
                out.append("")
            continue
        if stripped.startswith(">"):
            in_note = True
            continue
        if in_note and not re.match(r"^\*\*\d+\.\*\*", stripped) and not stripped.startswith("###"):
            continue
        in_note = False
        if stripped.startswith("###") or stripped == "---":
            continue
        match = re.match(r"^\*\*(\d+)\.\*\*\s*(.*)$", stripped)
        if match:
            claim_count += 1
            out.append(f"{match.group(1)}. {clean_markdown(match.group(2)).strip()}")
            continue
        cleaned = clean_markdown(stripped).strip()
        if cleaned:
            out.append(cleaned)
    text = "\n".join(out)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip(), claim_count


def add_paragraphs(doc: Document, text: str) -> None:
    for block in re.split(r"\n\s*\n", text.strip()):
        if not block.strip():
            continue
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        first = lines[0]
        if first.startswith("### "):
            doc.add_heading(clean_markdown(first[4:]).strip(), level=3)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
            continue
        if first.startswith("#### "):
            doc.add_heading(clean_markdown(first[5:]).strip(), level=4)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
            continue
        if first.startswith("## "):
            doc.add_heading(clean_markdown(first[3:]).strip(), level=2)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(clean_markdown(" ".join(lines)).strip())


def add_claims(doc: Document, claims_text: str) -> None:
    for block in re.split(r"\n\s*\n", claims_text.strip()):
        if not block.strip():
            continue
        if block.strip() == "What is claimed is:":
            doc.add_paragraph("What is claimed is:")
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(clean_markdown(" ".join(line.strip() for line in block.splitlines())).strip())


def configure_doc(doc: Document) -> None:
    section_obj = doc.sections[0]
    section_obj.top_margin = Inches(1)
    section_obj.bottom_margin = Inches(1)
    section_obj.left_margin = Inches(1)
    section_obj.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        styles[style_name].font.name = "Times New Roman"


def make_markdown(background: str, detailed_description: str, claims: str, abstract: str) -> str:
    return (
        "\n\n".join(
            [
                f"# {TITLE}",
                "## CROSS-REFERENCE TO RELATED APPLICATIONS",
                (
                    f"This application claims priority to U.S. Provisional Application No. {APPLICATION}, "
                    f'filed {FILING_DATE}, entitled "{TITLE}", the entire disclosure of which is '
                    "incorporated herein by reference."
                ),
                "## BACKGROUND OF THE INVENTION",
                background,
                "## SUMMARY OF THE INVENTION",
                "\n\n".join(SUMMARY_PARAGRAPHS),
                "## BRIEF DESCRIPTION OF THE DRAWINGS",
                "\n\n".join(FIGURE_DESCRIPTIONS),
                "## DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
                detailed_description,
                "## CLAIMS",
                claims,
                "## ABSTRACT OF THE DISCLOSURE",
                abstract,
            ]
        )
        + "\n"
    )


def build() -> dict:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    detailed = read(SOURCE_DETAILED)
    background = extract_background(detailed)
    detailed_description = extract_detailed_description(detailed)
    abstract = extract_abstract(read(SOURCE_ABSTRACT))
    claims, claim_count = extract_claims(read(SOURCE_CLAIMS))
    assembled_md = make_markdown(background, detailed_description, claims, abstract)
    OUT_MD.write_text(assembled_md, encoding="utf-8")

    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_heading("CROSS-REFERENCE TO RELATED APPLICATIONS", level=1)
    doc.add_paragraph(
        f"This application claims priority to U.S. Provisional Application No. {APPLICATION}, "
        f'filed {FILING_DATE}, entitled "{TITLE}", the entire disclosure of which is '
        "incorporated herein by reference."
    )

    doc.add_heading("BACKGROUND OF THE INVENTION", level=1)
    add_paragraphs(doc, background)

    doc.add_heading("SUMMARY OF THE INVENTION", level=1)
    for para in SUMMARY_PARAGRAPHS:
        doc.add_paragraph(para)

    doc.add_heading("BRIEF DESCRIPTION OF THE DRAWINGS", level=1)
    for fig in FIGURE_DESCRIPTIONS:
        doc.add_paragraph(fig)

    doc.add_heading("DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS", level=1)
    add_paragraphs(doc, detailed_description)

    doc.add_heading("CLAIMS", level=1)
    add_claims(doc, claims)

    doc.add_heading("ABSTRACT OF THE DISCLOSURE", level=1)
    doc.add_paragraph(abstract)

    doc.save(OUT_DOCX)

    manifest = {
        "schema": "scbe_nonprovisional_spec_assembly_v1",
        "title": TITLE,
        "application_number": APPLICATION,
        "filing_date": FILING_DATE,
        "output_docx": str(OUT_DOCX.relative_to(ROOT)).replace("\\", "/"),
        "output_markdown": str(OUT_MD.relative_to(ROOT)).replace("\\", "/"),
        "source_files": [
            str(SOURCE_DETAILED.relative_to(ROOT)).replace("\\", "/"),
            str(SOURCE_CLAIMS.relative_to(ROOT)).replace("\\", "/"),
            str(SOURCE_ABSTRACT.relative_to(ROOT)).replace("\\", "/"),
        ],
        "sections": [
            "Title",
            "Cross-Reference to Related Applications",
            "Background of the Invention",
            "Summary of the Invention",
            "Brief Description of the Drawings",
            "Detailed Description of the Preferred Embodiments",
            "Claims",
            "Abstract of the Disclosure",
        ],
        "claim_count": claim_count,
        "figure_count": len(FIGURE_DESCRIPTIONS),
        "notes": [
            "Draft assembly only; ADS and oath/declaration are not included.",
            "Claims are assembled from PATENT_CLAIMS_EXPANDED_v2.md with drafting/support notes removed.",
            "Drawings remain separate files in docs/legal/patent-figures/.",
        ],
    }
    OUT_JSON.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return manifest


if __name__ == "__main__":
    print(json.dumps(build(), indent=2))
