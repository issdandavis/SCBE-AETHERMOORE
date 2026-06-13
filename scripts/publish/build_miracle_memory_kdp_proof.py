"""Build KDP-style proof artifacts for The Miracle Was The Memory.

Outputs:
    artifacts/book/kdp/miracle-memory/miracle-memory-kdp-proof.docx
    artifacts/book/kdp/miracle-memory/miracle-memory-kdp-proof.pdf
    artifacts/book/kdp/miracle-memory/build-metadata.json

This is a proof builder, not a final typesetter. It creates a stable paperback
shape for reading, page-count checks, and KDP preview iteration.
"""

from __future__ import annotations

import argparse
import html
import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch as rl_inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

REPO_ROOT = Path(__file__).resolve().parents[2]
PROJECT_DIR = REPO_ROOT / "content" / "projects" / "miracle-memory"
DEFAULT_PROFILE = PROJECT_DIR / "kdp-proof-profile.json"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "artifacts" / "book" / "kdp" / "miracle-memory"


@dataclass(frozen=True)
class Block:
    kind: str
    text: str
    level: int = 0


def load_profile(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def word_count(text: str) -> int:
    body = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
    body = re.sub(r"[*_`>#\[\]()-]+", " ", body)
    return len(re.findall(r"\b[\w']+\b", body))


def clean_source_text(source: str) -> tuple[str, list[str]]:
    """Remove manuscript scaffolding that must never reach the print interior."""
    removed: list[str] = []
    kept: list[str] = []
    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith("Manuscript v1 spine, consolidated "):
            removed.append(stripped)
            continue
        kept.append(line)
    return "\n".join(kept), removed


def parse_markdown(source: str) -> list[Block]:
    blocks: list[Block] = []
    para: list[str] = []

    def flush_para() -> None:
        if para:
            text = " ".join(part.strip() for part in para if part.strip()).strip()
            if text:
                blocks.append(Block("paragraph", text))
            para.clear()

    for raw in source.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_para()
            continue

        if stripped in {"---", "***", "* * *"}:
            flush_para()
            blocks.append(Block("scene_break", ""))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if heading:
            flush_para()
            blocks.append(Block("heading", heading.group(2).strip(), len(heading.group(1))))
            continue

        para.append(stripped)

    flush_para()
    return blocks


def body_blocks(blocks: Iterable[Block], title: str) -> list[Block]:
    """Remove the source H1/title preamble; keep Chapter 0 onward."""
    result: list[Block] = []
    started = False
    for block in blocks:
        if block.kind == "heading" and block.level == 2:
            started = True
        if started:
            result.append(block)
    if not result:
        return [block for block in blocks if not (block.kind == "heading" and block.level == 1 and block.text == title)]
    return result


def contents_blocks(blocks: Iterable[Block]) -> list[Block]:
    """Build a front-matter contents list from part/chapter/back-matter heads."""
    contents: list[Block] = []
    for block in blocks:
        if block.kind != "heading":
            continue
        text = block.text.strip()
        lower = text.lower()
        if block.level == 1 and (lower.startswith("part ") or lower == "back matter"):
            contents.append(block)
        elif block.level == 2 and (lower.startswith("chapter ") or lower.startswith("interlude:")):
            contents.append(block)
        elif block.level == 2 and contents and contents[-1].text.lower() == "back matter":
            contents.append(block)
    return contents


def ensure_docx_font(run, font_name: str, size_pt: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def set_mirror_margins(section) -> None:
    sect_pr = section._sectPr
    mirror = sect_pr.find(qn("w:mirrorMargins"))
    if mirror is None:
        mirror = OxmlElement("w:mirrorMargins")
        sect_pr.append(mirror)


def add_docx_runs(paragraph, text: str, profile: dict) -> None:
    # Small, conservative inline formatter for proof copies.
    parts = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", text)
    for part in parts:
        if not part:
            continue
        bold = False
        italic = False
        content = part
        if part.startswith("***") and part.endswith("***"):
            bold = True
            italic = True
            content = part[3:-3]
        elif part.startswith("**") and part.endswith("**"):
            bold = True
            content = part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            italic = True
            content = part[1:-1]
        run = paragraph.add_run(content)
        ensure_docx_font(
            run,
            profile["typography"]["body_font"],
            profile["typography"]["body_size_pt"],
            bold=bold,
            italic=italic,
        )


def add_docx_centered_page(
    doc: Document,
    text: str,
    profile: dict,
    *,
    size_pt: float = 18,
    bold: bool = False,
    italic: bool = False,
    space_before_pt: float = 180,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(text)
    ensure_docx_font(run, profile["typography"]["body_font"], size_pt, bold=bold, italic=italic)
    doc.add_page_break()


def add_docx_copyright_page(doc: Document, profile: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(215)
    text = (
        f"Copyright © {profile['copyright_year']} {profile['author']}. All rights reserved.\n\n"
        f"{profile['edition_statement']}\n"
        f"{profile['isbn']}\n\n"
        f"{profile['creative_nonfiction_notice']}\n\n"
        "No part of this book may be reproduced, distributed, or transmitted in any form or by any means "
        "without prior written permission, except for brief quotations in reviews, criticism, scholarship, "
        "or other uses permitted by law.\n\n"
        f"Published by {profile['publisher']}"
    )
    run = p.add_run(text)
    ensure_docx_font(run, profile["typography"]["body_font"], 8.5)
    doc.add_page_break()


def add_docx_contents(doc: Document, profile: dict, contents: list[Block]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Contents")
    ensure_docx_font(run, profile["typography"]["body_font"], 17, bold=True)

    for block in contents:
        lower = block.text.lower()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if block.level == 1:
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(block.text)
            ensure_docx_font(run, profile["typography"]["body_font"], 10.5, bold=True)
        else:
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(1)
            label = block.text
            if lower.startswith("chapter "):
                label = re.sub(r"^Chapter\s+(\d+)\.\s+", r"\1. ", label)
            run = p.add_run(label)
            ensure_docx_font(run, profile["typography"]["body_font"], 9.5)


def build_docx(blocks: list[Block], profile: dict, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(profile["trim"]["width_in"])
    section.page_height = Inches(profile["trim"]["height_in"])
    section.top_margin = Inches(profile["margins"]["top_in"])
    section.bottom_margin = Inches(profile["margins"]["bottom_in"])
    section.left_margin = Inches(profile["margins"]["inside_in"])
    section.right_margin = Inches(profile["margins"]["outside_in"])
    set_mirror_margins(section)

    normal = doc.styles["Normal"]
    normal.font.name = profile["typography"]["body_font"]
    normal.font.size = Pt(profile["typography"]["body_size_pt"])
    normal.paragraph_format.line_spacing = Pt(profile["typography"]["leading_pt"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    add_docx_centered_page(doc, profile["title"], profile, size_pt=20, bold=False, space_before_pt=190)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(profile["title"])
    ensure_docx_font(run, profile["typography"]["body_font"], 25, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run(profile["subtitle"])
    ensure_docx_font(run, profile["typography"]["body_font"], 11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"by {profile['author']}")
    ensure_docx_font(run, profile["typography"]["body_font"], 14)
    doc.add_page_break()

    add_docx_copyright_page(doc, profile)

    if profile.get("dedication"):
        add_docx_centered_page(doc, profile["dedication"], profile, size_pt=11, italic=True, space_before_pt=170)

    epigraph = profile.get("epigraph") or {}
    if epigraph.get("enabled") and epigraph.get("text"):
        epigraph_text = epigraph["text"]
        if epigraph.get("attribution"):
            epigraph_text = f"{epigraph_text}\n\n— {epigraph['attribution']}"
        add_docx_centered_page(doc, epigraph_text, profile, size_pt=10.5, italic=True, space_before_pt=165)

    add_docx_contents(doc, profile, contents_blocks(blocks))

    first_after_break = True
    for block in blocks:
        if block.kind == "heading":
            if block.level <= 2:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(58)
                p.paragraph_format.space_after = Pt(28)
                size = profile["typography"]["chapter_title_size_pt"]
                if block.text.lower().startswith("part "):
                    size = profile["typography"]["part_title_size_pt"]
                run = p.add_run(block.text)
                ensure_docx_font(run, profile["typography"]["body_font"], size, bold=True)
                first_after_break = True
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(block.text)
                ensure_docx_font(run, profile["typography"]["body_font"], 12, bold=True)
                first_after_break = True
            continue

        if block.kind == "scene_break":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("*   *   *")
            ensure_docx_font(run, profile["typography"]["body_font"], profile["typography"]["body_size_pt"])
            first_after_break = True
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(profile["typography"]["leading_pt"])
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if not first_after_break:
            p.paragraph_format.first_line_indent = Inches(profile["typography"]["first_line_indent_in"])
        add_docx_runs(p, block.text, profile)
        first_after_break = False

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def register_georgia() -> None:
    font_dir = Path("C:/Windows/Fonts")
    candidates = {
        "Georgia": font_dir / "georgia.ttf",
        "Georgia-Bold": font_dir / "georgiab.ttf",
        "Georgia-Italic": font_dir / "georgiai.ttf",
        "Georgia-BoldItalic": font_dir / "georgiaz.ttf",
    }
    for name, path in candidates.items():
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))


def md_inline_to_reportlab(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", escaped)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\*(.+?)\*", r"<i>\1</i>", escaped)
    return escaped


def build_pdf(blocks: list[Block], profile: dict, out_path: Path) -> int:
    register_georgia()
    page_size = (profile["trim"]["width_in"] * inch, profile["trim"]["height_in"] * inch)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=page_size,
        rightMargin=profile["margins"]["outside_in"] * rl_inch,
        leftMargin=profile["margins"]["inside_in"] * rl_inch,
        topMargin=profile["margins"]["top_in"] * rl_inch,
        bottomMargin=profile["margins"]["bottom_in"] * rl_inch,
        title=profile["title"],
        author=profile["author"],
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Georgia",
        fontSize=profile["typography"]["body_size_pt"],
        leading=profile["typography"]["leading_pt"],
        firstLineIndent=profile["typography"]["first_line_indent_in"] * rl_inch,
        alignment=TA_JUSTIFY,
        spaceBefore=0,
        spaceAfter=0,
    )
    body_first = ParagraphStyle("BodyFirst", parent=body, firstLineIndent=0)
    chapter = ParagraphStyle(
        "Chapter",
        parent=styles["Heading1"],
        fontName="Georgia-Bold",
        fontSize=profile["typography"]["chapter_title_size_pt"],
        leading=profile["typography"]["chapter_title_size_pt"] + 4,
        alignment=TA_CENTER,
        spaceBefore=0.55 * rl_inch,
        spaceAfter=0.25 * rl_inch,
    )
    part = ParagraphStyle(
        "Part",
        parent=chapter,
        fontSize=profile["typography"]["part_title_size_pt"],
        leading=profile["typography"]["part_title_size_pt"] + 4,
        spaceBefore=1.5 * rl_inch,
        spaceAfter=0.2 * rl_inch,
    )
    title_style = ParagraphStyle("Title", parent=chapter, fontSize=24, leading=30, spaceBefore=1.2 * rl_inch)
    half_title_style = ParagraphStyle(
        "HalfTitle", parent=chapter, fontName="Georgia", fontSize=19, leading=24, spaceBefore=2.45 * rl_inch
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontName="Georgia-Italic", fontSize=10.5, leading=14, alignment=TA_CENTER
    )
    center = ParagraphStyle("Center", parent=body_first, alignment=TA_CENTER, firstLineIndent=0)
    small = ParagraphStyle("Small", parent=body_first, fontSize=8.5, leading=11, alignment=TA_LEFT)
    contents_title = ParagraphStyle(
        "ContentsTitle", parent=chapter, fontSize=17, leading=22, spaceBefore=0.3 * rl_inch, spaceAfter=0.2 * rl_inch
    )
    contents_part = ParagraphStyle(
        "ContentsPart", parent=body_first, fontName="Georgia-Bold", fontSize=10, leading=13, spaceBefore=8, spaceAfter=2
    )
    contents_item = ParagraphStyle(
        "ContentsItem",
        parent=body_first,
        fontSize=9.2,
        leading=11.5,
        leftIndent=0.16 * rl_inch,
        spaceBefore=0,
        spaceAfter=1,
    )
    copyright_text = (
        f"Copyright © {html.escape(str(profile['copyright_year']))} {html.escape(profile['author'])}. "
        "All rights reserved.<br/><br/>"
        f"{html.escape(profile['edition_statement'])}<br/>"
        f"{html.escape(profile['isbn'])}<br/><br/>"
        f"{html.escape(profile['creative_nonfiction_notice'])}<br/><br/>"
        "No part of this book may be reproduced, distributed, or transmitted in any form or by any means "
        "without prior written permission, except for brief quotations in reviews, criticism, scholarship, "
        "or other uses permitted by law.<br/><br/>"
        f"Published by {html.escape(profile['publisher'])}"
    )

    story = [
        Paragraph(html.escape(profile["title"]), half_title_style),
        PageBreak(),
        Paragraph(html.escape(profile["title"]), title_style),
        Spacer(1, 0.12 * rl_inch),
        Paragraph(html.escape(profile["subtitle"]), subtitle_style),
        Spacer(1, 0.35 * rl_inch),
        Paragraph(f"by {html.escape(profile['author'])}", center),
        PageBreak(),
        Spacer(1, 3.35 * rl_inch),
        Paragraph(copyright_text, small),
        PageBreak(),
    ]
    if profile.get("dedication"):
        dedication_style = ParagraphStyle(
            "Dedication",
            parent=subtitle_style,
            fontName="Georgia-Italic",
            fontSize=10.75,
            leading=15,
            leftIndent=0.25 * rl_inch,
            rightIndent=0.25 * rl_inch,
            spaceBefore=2.35 * rl_inch,
        )
        story.extend(
            [
                Paragraph(html.escape(profile["dedication"]), dedication_style),
                PageBreak(),
            ]
        )
    epigraph = profile.get("epigraph") or {}
    if epigraph.get("enabled") and epigraph.get("text"):
        epigraph_style = ParagraphStyle(
            "Epigraph",
            parent=subtitle_style,
            fontName="Georgia-Italic",
            fontSize=10.5,
            leading=14,
            leftIndent=0.3 * rl_inch,
            rightIndent=0.3 * rl_inch,
            spaceBefore=2.25 * rl_inch,
        )
        epigraph_text = html.escape(epigraph["text"])
        if epigraph.get("attribution"):
            epigraph_text = f"{epigraph_text}<br/><br/>— {html.escape(epigraph['attribution'])}"
        story.extend([Paragraph(epigraph_text, epigraph_style), PageBreak()])

    story.append(Paragraph("Contents", contents_title))
    for block in contents_blocks(blocks):
        lower = block.text.lower()
        if block.level == 1:
            story.append(Paragraph(html.escape(block.text), contents_part))
        else:
            label = block.text
            if lower.startswith("chapter "):
                label = re.sub(r"^Chapter\s+(\d+)\.\s+", r"\1. ", label)
            story.append(Paragraph(html.escape(label), contents_item))

    first_after_break = True
    for block in blocks:
        if block.kind == "heading":
            if block.level <= 2:
                story.append(PageBreak())
                style = part if block.text.lower().startswith("part ") else chapter
                story.append(Paragraph(html.escape(block.text), style))
                first_after_break = True
            else:
                story.append(Spacer(1, 0.12 * rl_inch))
                story.append(Paragraph(f"<b>{html.escape(block.text)}</b>", body_first))
                story.append(Spacer(1, 0.06 * rl_inch))
                first_after_break = True
            continue

        if block.kind == "scene_break":
            story.append(Spacer(1, 0.12 * rl_inch))
            story.append(Paragraph("*&nbsp;&nbsp;&nbsp;*&nbsp;&nbsp;&nbsp;*", center))
            story.append(Spacer(1, 0.12 * rl_inch))
            first_after_break = True
            continue

        style = body_first if first_after_break else body
        story.append(Paragraph(md_inline_to_reportlab(block.text), style))
        first_after_break = False

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)
    return doc.page


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--profile", type=Path, default=DEFAULT_PROFILE)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    args = parser.parse_args()

    profile = load_profile(args.profile)
    source_path = REPO_ROOT / profile["source"]
    raw_source_text = source_path.read_text(encoding="utf-8")
    source_text, removed_scaffold = clean_source_text(raw_source_text)
    parsed = parse_markdown(source_text)
    blocks = body_blocks(parsed, profile["title"])

    out_dir = args.out_dir
    docx_path = out_dir / "miracle-memory-kdp-proof.docx"
    pdf_path = out_dir / "miracle-memory-kdp-proof.pdf"
    metadata_path = out_dir / "build-metadata.json"

    build_docx(blocks, profile, docx_path)
    pages = build_pdf(blocks, profile, pdf_path)

    headings = [block.text for block in blocks if block.kind == "heading"]
    metadata = {
        "built_at_utc": datetime.now(timezone.utc).isoformat(),
        "profile": str(args.profile.relative_to(REPO_ROOT)),
        "source": profile["source"],
        "docx": str(docx_path.relative_to(REPO_ROOT)),
        "pdf": str(pdf_path.relative_to(REPO_ROOT)),
        "word_count": word_count(source_text),
        "page_count_pdf": pages,
        "heading_count": len(headings),
        "chapter_like_heading_count": sum(1 for h in headings if h.lower().startswith("chapter ")),
        "front_matter": {
            "half_title": True,
            "title_page": True,
            "copyright_page": True,
            "dedication": bool(profile.get("dedication")),
            "epigraph": bool((profile.get("epigraph") or {}).get("enabled")),
            "contents": True,
        },
        "removed_scaffold_lines": removed_scaffold,
        "print_recommendation": {
            "trim": "5.5 x 8.5 in",
            "paper": "cream",
            "ink": "black",
            "cover_finish": "matte",
            "bleed": "no bleed",
            "font": "Georgia",
        },
    }
    metadata_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")
    print(json.dumps(metadata, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
