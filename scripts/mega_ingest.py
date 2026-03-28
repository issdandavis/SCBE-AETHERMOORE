"""
Mega Ingest Pipeline
=====================

Pulls ALL training data from every source across the filesystem into
tongue-tagged SFT pairs ready for Tetris embedding.

Sources:
  1. REPO: All existing JSONL in training-data/ and training/
  2. BOOKS: Six Tongues Protocol (27 chapters + 11 interludes), Reincarnated novel
  3. LORE: Everweave PDFs, Avalon Codex Archives, 500-page theory doc
  4. CHATGPT: Full ChatGPT export (HTML + raw text logs)
  5. NOTION: Workspace export ZIP + Dropbox knowledge JSONs
  6. GEMINI: Google Gemini HTML export
  7. GAME: ChoiceScript scenes, game sessions
  8. OBSIDIAN: Avalon Files vault notes
  9. GROK: Drop zone for Grok export (when available)
  10. DRIVE: Recently moved files

Run:
  python scripts/mega_ingest.py                    # Scan + report only
  python scripts/mega_ingest.py --ingest           # Full ingest to JSONL
  python scripts/mega_ingest.py --ingest --tetris  # Ingest + Tetris embed
  python scripts/mega_ingest.py --ingest --push-hf # Ingest + embed + push
"""

import argparse
import hashlib
import html
import json
import os
import re
import sys
import time
import zipfile
from html.parser import HTMLParser
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
HOME = Path("C:/Users/issda")
sys.path.insert(0, str(ROOT))

# ============================================================
# TONGUE INFERENCE
# ============================================================

TONGUE_KEYS = ["KO", "AV", "RU", "CA", "UM", "DR"]

TONGUE_KEYWORDS = {
    "KO": [
        "command",
        "orchestrat",
        "dispatch",
        "fleet",
        "coordinat",
        "intent",
        "govern",
        "decision",
        "leader",
        "polly",
        "raven",
        "archive",
    ],
    "AV": [
        "transport",
        "navigat",
        "brows",
        "search",
        "web",
        "music",
        "art",
        "spirit",
        "aria",
        "melody",
        "song",
        "tongue",
        "language",
        "lexicon",
    ],
    "RU": [
        "research",
        "hypothes",
        "entropy",
        "chaos",
        "explor",
        "clay",
        "golem",
        "dream",
        "emotion",
        "memory",
        "binding",
        "ancient",
    ],
    "CA": [
        "code",
        "compute",
        "train",
        "deploy",
        "test",
        "implement",
        "math",
        "algorithm",
        "kael",
        "cipher",
        "lattice",
        "quantum",
        "formula",
    ],
    "UM": [
        "secur",
        "govern",
        "audit",
        "threat",
        "scan",
        "seal",
        "shadow",
        "encrypt",
        "vault",
        "ward",
        "protect",
        "shield",
        "zara",
    ],
    "DR": [
        "structur",
        "architect",
        "document",
        "debug",
        "heal",
        "eldrin",
        "tree",
        "root",
        "branch",
        "layer",
        "pipeline",
        "foundation",
    ],
}

# Chapter-level tongue mapping for Six Tongues Protocol
STP_TONGUE_MAP = {
    "ch01": "KO",
    "ch02": "AV",
    "ch03": "RU",
    "ch04": "CA",
    "ch05": "UM",
    "ch06": "DR",
    "ch07": "KO",
    "ch08": "AV",
    "ch09": "RU",
    "ch10": "CA",
    "ch11": "UM",
    "ch12": "DR",
    "ch13": "KO",
    "ch14": "AV",
    "ch15": "RU",
    "ch16": "CA",
    "ch17": "UM",
    "ch18": "DR",
    "ch19": "KO",
    "ch20": "AV",
    "ch21": "RU",
    "ch22": "CA",
    "ch23": "UM",
    "ch24": "DR",
    "ch25": "KO",
    "ch26": "AV",
    "ch27": "RU",
    "ch-ground-level": "DR",
    "ch-rootlight": "AV",
    "ch-the-deep-shelf": "UM",
    "ch-the-relay": "CA",
    "ch-tongue-and-tooth": "KO",
    "interlude-01": "AV",
    "interlude-02": "RU",
    "interlude-03": "CA",
    "interlude-04": "UM",
    "interlude-05": "DR",
    "interlude-06": "KO",
    "interlude-07": "AV",
    "interlude-08": "RU",
    "interlude-09": "CA",
    "interlude-10": "UM",
    "interlude-11": "DR",
}


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._skip_depth = 0
        self._chunks: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        if tag.lower() in {"script", "style"}:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        tag_name = tag.lower()
        if tag_name in {"script", "style"} and self._skip_depth > 0:
            self._skip_depth -= 1
        elif tag_name in {"p", "div", "section", "article", "main", "li", "br"}:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if self._skip_depth == 0 and data:
            self._chunks.append(data)

    def get_text(self) -> str:
        return "".join(self._chunks)


def _html_to_visible_text(raw: str) -> str:
    parser = _VisibleTextParser()
    parser.feed(raw)
    parser.close()
    text = html.unescape(parser.get_text())
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def infer_tongue(text: str) -> str:
    """Infer Sacred Tongue from text content."""
    lower = text.lower()
    best, best_n = "KO", 0
    for tongue, kws in TONGUE_KEYWORDS.items():
        n = sum(1 for kw in kws if kw in lower)
        if n > best_n:
            best, best_n = tongue, n
    return best


def chunk_text(text: str, max_chars: int = 1500, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks, breaking at paragraph boundaries."""
    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) > max_chars and current:
            chunks.append(current.strip())
            # Keep overlap from end of previous chunk
            current = current[-overlap:] + "\n\n" + para if overlap else para
        else:
            current = current + "\n\n" + para if current else para
    if current.strip():
        chunks.append(current.strip())
    return chunks


def text_to_sft_pairs(text: str, source: str, tongue: str = "", chapter: str = "", max_chars: int = 1500) -> list[dict]:
    """Convert a text document into SFT instruction/response pairs."""
    chunks = chunk_text(text, max_chars=max_chars)
    pairs = []
    for i, chunk in enumerate(chunks):
        if len(chunk) < 50:
            continue
        t = tongue or infer_tongue(chunk)
        # Generate instruction from first sentence or title
        first_line = chunk.split("\n")[0][:200].strip()
        if first_line.startswith("#"):
            first_line = first_line.lstrip("#").strip()

        instruction = f"Explain or continue the following from the {source}"
        if chapter:
            instruction += f" ({chapter})"
        instruction += f": {first_line}"

        pairs.append(
            {
                "instruction": instruction,
                "response": chunk,
                "metadata": {
                    "source": source,
                    "tongue": t,
                    "chapter": chapter,
                    "chunk_index": i,
                    "type": "mega_ingest",
                },
                "encoding_tongue": t,
                "timestamp": time.time(),
            }
        )
    return pairs


# ============================================================
# SOURCE LOADERS
# ============================================================


def load_all_repo_jsonl() -> list[dict]:
    """Load ALL JSONL files from training-data/ and training/."""
    pairs = []
    dirs = [
        ROOT / "training-data",
        ROOT / "training",
    ]
    seen_files = set()
    for d in dirs:
        if not d.exists():
            continue
        for jf in sorted(d.rglob("*.jsonl")):
            if jf.name in seen_files:
                continue
            seen_files.add(jf.name)
            try:
                with open(jf, encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if not line:
                            continue
                        try:
                            rec = json.loads(line)
                            if isinstance(rec, dict) and (
                                rec.get("instruction") or rec.get("prompt") or rec.get("response") or rec.get("text")
                            ):
                                rec.setdefault("metadata", {})
                                if isinstance(rec["metadata"], str):
                                    rec["metadata"] = {"raw": rec["metadata"]}
                                rec["metadata"]["source_file"] = jf.name
                                pairs.append(rec)
                        except json.JSONDecodeError:
                            continue
            except Exception:
                continue
    return pairs


def load_six_tongues_protocol() -> list[dict]:
    """Load The Six Tongues Protocol book — all chapters and interludes."""
    book_dir = HOME / "OneDrive" / "Books" / "The Six Tongues Protocol"
    pairs = []
    if not book_dir.exists():
        return pairs

    for md_file in sorted(book_dir.glob("*.md")):
        stem = md_file.stem
        if stem in ("INDEX", "character-designs", "build_kdp", "build_pandoc"):
            continue
        tongue = STP_TONGUE_MAP.get(stem, infer_tongue(md_file.read_text(encoding="utf-8")[:500]))
        text = md_file.read_text(encoding="utf-8")
        chapter_pairs = text_to_sft_pairs(text, source="Six Tongues Protocol", tongue=tongue, chapter=stem)
        pairs.extend(chapter_pairs)
    return pairs


def load_reincarnated_novel() -> list[dict]:
    """Load Reincarnated Into Another World novel."""
    novel_dir = HOME / "OneDrive" / "Books" / "Reincarnated-KDP-Prep-2026-03-11" / "manuscript"
    pairs = []
    for name in ["source-kdp-edit.md", "source.md"]:
        path = novel_dir / name
        if path.exists():
            text = path.read_text(encoding="utf-8")
            pairs.extend(text_to_sft_pairs(text, source="Reincarnated Novel", chapter=name))
            break  # Use the best version
    return pairs


def load_everweave_text() -> list[dict]:
    """Load Everweave API text dump and other raw text exports."""
    pairs = []
    text_sources = [
        (
            HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Exports" / "everweave_api_fullpdf.txt",
            "Everweave API",
        ),
        (
            HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Exports" / "Entire chat log big boy.txt",
            "ChatGPT Avalon Log",
        ),
        (
            HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Exports" / "Entire chat log.txt",
            "ChatGPT Avalon Log",
        ),
    ]

    # Also check alternate locations
    for alt_dir in [
        HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports",
        HOME / "OneDrive" / "Imports" / "issdandavis7795@gmail.com - Dropbox" / "AvalonBook STUFF",
    ]:
        for name in ["httpsstory.everweave.aiapifullpdffa.txt", "Entire chat log big boy.txt", "Entire chat log.txt"]:
            p = alt_dir / name
            if p.exists():
                src = "Everweave API" if "everweave" in name.lower() else "ChatGPT Avalon Log"
                text_sources.append((p, src))

    seen_hashes = set()
    for path, source in text_sources:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="replace")
        h = hashlib.md5(text[:5000].encode()).hexdigest()
        if h in seen_hashes:
            continue
        seen_hashes.add(h)
        pairs.extend(text_to_sft_pairs(text, source=source))
    return pairs


def load_theory_doc() -> list[dict]:
    """Load the 500-page theory document."""
    pairs = []
    for path in [
        HOME / "OneDrive" / "Lore_and_Writing" / "500_page_doc_on_theory.txt",
        HOME / "OneDrive" / "Lore_and_Writing" / "500_page_doc_on_theory.docx",
    ]:
        if path.exists() and path.suffix == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace")
            pairs.extend(text_to_sft_pairs(text, source="500-page Theory Doc"))
            break
    return pairs


def load_lore_drafts() -> list[dict]:
    """Load major lore drafts and manuscripts from OneDrive."""
    pairs = []
    lore_files = [
        # FINAL DRAFTS
        (HOME / "OneDrive" / "FINAL DRAFTS" / "THE SPIRAL 500page Edtiors additons.8.docx", "Spiral 500p Editor"),
        (HOME / "OneDrive" / "FINAL DRAFTS" / "lore book massive.docx", "Lore Book Massive"),
        # Books folder
        (HOME / "OneDrive" / "Books" / "Book_Needs_Editing.docx", "Book Needs Editing"),
        (HOME / "OneDrive" / "Books" / "Review_Sheet.docx", "Review Sheet"),
        (
            HOME / "OneDrive" / "Books" / "FUCKING BOOK" / "Reincarnated Into Another World to Be a God.docx",
            "Reincarnated DOCX",
        ),
        # Lore Drafts
        (HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "MASTER DOC.docx", "Master Doc"),
        (HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "full game logs new 20205.docx", "Game Logs 2025"),
        (
            HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "# The Spiral of Pollyoneth full lengths.txt",
            "Spiral of Pollyoneth",
        ),
        (HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "Narrative framewor docs.txt", "Narrative Framework"),
        (HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "Market overview.txt", "Market Overview"),
        # Documents
        (HOME / "OneDrive" / "Documents" / "# The Spiral of Avalon.txt", "Spiral of Avalon"),
        (HOME / "OneDrive" / "Documents" / "#DarkSetting, Happy Ending.txt", "Dark Setting Chronicle"),
        (HOME / "OneDrive" / "Documents" / "New chat,BIG DUMP.docx", "Big Dump"),
        (HOME / "OneDrive" / "Documents" / "Thought for 13s.docx", "Thought 13s"),
        (HOME / "OneDrive" / "Documents" / "NOTES FR SYSTEM.txt", "System Notes"),
        # Avalon Reference Pack
        (HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Guides" / "foreshadowing.txt", "Foreshadowing Guide"),
        (
            HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Guides" / "narrative_framework_docs.txt",
            "Narrative Framework Ref",
        ),
        # Archive
        (
            HOME / "OneDrive" / "Documents" / "GitHub" / "Avalon" / "archive" / "700000 characters.txt",
            "700K Characters Archive",
        ),
    ]

    seen_hashes = set()
    for path, source in lore_files:
        if not path.exists():
            continue
        if path.suffix == ".txt":
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                continue
        elif path.suffix in (".docx", ".odt", ".rtf"):
            # For binary formats, try python-docx or skip with note
            try:
                import docx

                doc = docx.Document(str(path))
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                print(f"  SKIP (no python-docx): {path.name}")
                continue
            except Exception:
                continue
        else:
            continue

        h = hashlib.md5(text[:5000].encode()).hexdigest()
        if h in seen_hashes:
            continue
        seen_hashes.add(h)
        if len(text) < 100:
            continue
        pairs.extend(text_to_sft_pairs(text, source=source))
    return pairs


def load_choicescript() -> list[dict]:
    """Load ChoiceScript game scenes as interactive fiction training data."""
    pairs = []
    cs_dir = HOME / "OneDrive" / "Documents" / "GitHub" / "Avalon" / "choicescript_game" / "scenes"
    if not cs_dir.exists():
        return pairs
    for scene in sorted(cs_dir.glob("*.txt")):
        text = scene.read_text(encoding="utf-8", errors="replace")
        if len(text) < 100:
            continue
        pairs.extend(text_to_sft_pairs(text, source="ChoiceScript Game", chapter=scene.stem))
    return pairs


def load_obsidian_vault() -> list[dict]:
    """Load Obsidian vault notes from Avalon Files."""
    pairs = []
    vault = HOME / "Documents" / "Avalon Files" / "SCBE Research"
    if not vault.exists():
        return pairs
    for md in sorted(vault.rglob("*.md")):
        text = md.read_text(encoding="utf-8", errors="replace")
        if len(text) < 100:
            continue
        pairs.extend(text_to_sft_pairs(text, source="Obsidian Vault", chapter=md.stem))
    return pairs


def load_notion_knowledge() -> list[dict]:
    """Load Notion knowledge JSONs from Dropbox."""
    pairs = []
    notion_dir = HOME / "Dropbox" / "SCBE" / "knowledge" / "notion"
    if not notion_dir.exists():
        return pairs
    for jf in sorted(notion_dir.rglob("*.json")):
        try:
            data = json.loads(jf.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                title = data.get("title", data.get("name", jf.stem))
                content = data.get("content", data.get("body", data.get("text", "")))
                if isinstance(content, list):
                    content = "\n".join(str(c) for c in content)
                if content and len(str(content)) > 50:
                    tongue = infer_tongue(str(content))
                    pairs.append(
                        {
                            "instruction": f"Describe the {title} concept from the SCBE Notion workspace.",
                            "response": str(content)[:2000],
                            "metadata": {
                                "source": "notion_knowledge",
                                "tongue": tongue,
                                "file": jf.name,
                                "type": "mega_ingest",
                            },
                            "encoding_tongue": tongue,
                            "timestamp": time.time(),
                        }
                    )
        except Exception:
            continue
    return pairs


def load_notion_zip() -> list[dict]:
    """Try to unpack and load the Notion workspace export ZIP."""
    pairs = []
    zip_path = (
        HOME
        / "OneDrive"
        / "Downloads"
        / "91fc3693-09c4-4ff5-b481-a7902c21431b_Export-d3f8086e-07e0-444f-9291-10b7fe375b22"
        / "Export-d3f8086e-07e0-444f-9291-10b7fe375b22-Part-1.zip"
    )
    if not zip_path.exists():
        return pairs

    extract_dir = ROOT / "artifacts" / "notion_export_unpacked"
    if not extract_dir.exists():
        try:
            extract_dir.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(zip_path, "r") as zf:
                zf.extractall(extract_dir)
        except Exception as e:
            print(f"  Notion ZIP extract error: {e}")
            return pairs

    # Read all markdown and CSV from extracted
    for md in sorted(extract_dir.rglob("*.md")):
        try:
            text = md.read_text(encoding="utf-8", errors="replace")
            if len(text) < 100:
                continue
            pairs.extend(text_to_sft_pairs(text, source="Notion Export", chapter=md.stem))
        except Exception:
            continue
    for csv_f in sorted(extract_dir.rglob("*.csv")):
        try:
            text = csv_f.read_text(encoding="utf-8", errors="replace")
            if len(text) < 100:
                continue
            # Convert CSV rows to descriptive text
            lines = text.strip().split("\n")
            if len(lines) > 1:
                header = lines[0]
                for row in lines[1:]:
                    pairs.append(
                        {
                            "instruction": f"What does this Notion record describe? Headers: {header}",
                            "response": row,
                            "metadata": {
                                "source": "notion_csv",
                                "tongue": infer_tongue(row),
                                "file": csv_f.name,
                                "type": "mega_ingest",
                            },
                            "encoding_tongue": infer_tongue(row),
                            "timestamp": time.time(),
                        }
                    )
        except Exception:
            continue
    return pairs


def load_gemini_export() -> list[dict]:
    """Parse Google Gemini HTML export into conversation pairs."""
    pairs = []
    gemini_path = HOME / "Downloads" / "Google Gemini.html"
    if not gemini_path.exists():
        return pairs
    try:
        raw = gemini_path.read_text(encoding="utf-8", errors="replace")
        text = _html_to_visible_text(raw)
        if len(text) > 200:
            pairs.extend(text_to_sft_pairs(text, source="Google Gemini Export"))
    except Exception as e:
        print(f"  Gemini parse error: {e}")
    return pairs


def load_chatgpt_html() -> list[dict]:
    """Parse ChatGPT Data Export HTML into conversation pairs."""
    pairs = []
    chatgpt_paths = [
        HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack" / "Exports" / "ChatGPT Data Export.html",
        HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "ChatGPT Data Export.html",
    ]
    for path in chatgpt_paths:
        if not path.exists():
            continue
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            text = _html_to_visible_text(raw)
            if len(text) > 500:
                pairs.extend(text_to_sft_pairs(text, source="ChatGPT Export"))
            break  # Only use one copy
        except Exception as e:
            print(f"  ChatGPT HTML parse error: {e}")
    return pairs


def load_grok_drop() -> list[dict]:
    """Load Grok export from drop zone (user will place it here)."""
    pairs = []
    drop_zone = ROOT / "training" / "intake" / "grok"
    drop_zone.mkdir(parents=True, exist_ok=True)

    for f in sorted(drop_zone.iterdir()):
        if f.suffix in (".txt", ".md", ".json", ".html", ".jsonl"):
            try:
                text = f.read_text(encoding="utf-8", errors="replace")
                if f.suffix == ".html":
                    text = _html_to_visible_text(text)
                if len(text) > 100:
                    pairs.extend(text_to_sft_pairs(text, source="Grok Export", chapter=f.stem))
            except Exception:
                continue
    return pairs


def load_spiralverse_mega() -> list[dict]:
    """Load the 16.8MB Spiralverse Mega Document from phone backup."""
    pairs = []
    mega_path = (
        HOME
        / "OneDrive"
        / "Organized"
        / "PhoneBackups"
        / "CrossDevice"
        / "Issac's S24+ (1)"
        / "storage"
        / "Download"
        / "Spiralverse_Mega_Document.txt"
    )
    if mega_path.exists():
        text = mega_path.read_text(encoding="utf-8", errors="replace")
        pairs.extend(text_to_sft_pairs(text, source="Spiralverse Mega Document", max_chars=2000))
    return pairs


def load_phone_backup_lore() -> list[dict]:
    """Load lore files from phone backup in OneDrive/Organized."""
    pairs = []
    phone_dir = (
        HOME / "OneDrive" / "Organized" / "PhoneBackups" / "CrossDevice" / "Issac's S24+ (1)" / "storage" / "Download"
    )
    if not phone_dir.exists():
        return pairs

    lore_files = [
        ("enhanced_character_codex.md", "Character Codex"),
        ("spiralverse_chronicles_complete.txt", "Spiralverse Chronicles"),
        ("SHORE_TO_KING_MASTER_CONSOLIDATION-1.md", "Shore to King"),
        ("Entire_chat_log_copy.txt", "Chat Log Phone Copy"),
    ]
    seen = set()
    for name, source in lore_files:
        path = phone_dir / name
        if path.exists():
            text = path.read_text(encoding="utf-8", errors="replace")
            h = hashlib.md5(text[:3000].encode()).hexdigest()
            if h not in seen:
                seen.add(h)
                pairs.extend(text_to_sft_pairs(text, source=source))

    # ChoiceScript game from phone
    cs_dir = phone_dir / "spiralverse_game_updated"
    if cs_dir.exists():
        for f in sorted(cs_dir.glob("*.txt")):
            text = f.read_text(encoding="utf-8", errors="replace")
            if len(text) > 100:
                pairs.extend(text_to_sft_pairs(text, source="Phone ChoiceScript", chapter=f.stem))

    # Phone notes
    for notes in phone_dir.parent.glob("Notes_*.txt"):
        text = notes.read_text(encoding="utf-8", errors="replace")
        if len(text) > 200:
            pairs.extend(text_to_sft_pairs(text, source="Phone Notes", chapter=notes.stem))

    return pairs


def load_dropbox_story_series() -> list[dict]:
    """Load story chapters from Dropbox live sync."""
    pairs = []
    story_dir = (
        HOME
        / "Dropbox"
        / "local-workspace-sync (Selective Sync Conflict)"
        / "SCBE-AETHERMOORE-live-sync"
        / "content"
        / "articles"
        / "story_series"
    )
    if not story_dir.exists():
        return pairs

    for md in sorted(story_dir.glob("*.md")):
        text = md.read_text(encoding="utf-8", errors="replace")
        if len(text) > 100:
            tongue = infer_tongue(text[:500])
            pairs.extend(text_to_sft_pairs(text, source="Dropbox Story Series", tongue=tongue, chapter=md.stem))

    # Expanded editions
    expanded = story_dir / "expanded"
    if expanded.exists():
        for md in sorted(expanded.glob("*.md")):
            text = md.read_text(encoding="utf-8", errors="replace")
            if len(text) > 100:
                pairs.extend(text_to_sft_pairs(text, source="Expanded Story", chapter=md.stem))
    return pairs


def load_dropbox_stp_editorial() -> list[dict]:
    """Load Six Tongues Protocol editorial + reference from Dropbox/Izack Realmforge."""
    pairs = []
    stp_dir = HOME / "OneDrive" / "Dropbox" / "Izack Realmforge" / "Story Files" / "The Six Tongues Protocol"
    if not stp_dir.exists():
        # Try alternate
        stp_dir = (
            HOME
            / "OneDrive"
            / "Imports"
            / "issdandavis7795@gmail.com - Dropbox"
            / "Izack Realmforge"
            / "Story Files"
            / "The Six Tongues Protocol"
        )
    if not stp_dir.exists():
        return pairs

    for md in sorted(stp_dir.rglob("*.md")):
        text = md.read_text(encoding="utf-8", errors="replace")
        if len(text) > 200:
            pairs.extend(text_to_sft_pairs(text, source="STP Editorial", chapter=md.stem))
    return pairs


def load_dropbox_spiral_chapters() -> list[dict]:
    """Load Spiral of Pollyoneth chapter files from Dropbox/Izack Realmforge."""
    pairs = []
    spiral_dir = HOME / "OneDrive" / "Dropbox" / "Izack Realmforge" / "Story Files" / "Spiral of Pollyoneth"
    if not spiral_dir.exists():
        spiral_dir = (
            HOME
            / "OneDrive"
            / "Imports"
            / "issdandavis7795@gmail.com - Dropbox"
            / "Izack Realmforge"
            / "Story Files"
            / "Spiral of Pollyoneth"
        )
    if not spiral_dir.exists():
        return pairs

    for md in sorted(spiral_dir.rglob("*.md")):
        text = md.read_text(encoding="utf-8", errors="replace")
        if len(text) > 200:
            pairs.extend(text_to_sft_pairs(text, source="Spiral of Pollyoneth", chapter=md.stem))
    return pairs


def load_scbe_archive_docs() -> list[dict]:
    """Load large SCBE documentation from Dropbox live sync."""
    pairs = []
    docs_dir = (
        HOME / "Dropbox" / "local-workspace-sync (Selective Sync Conflict)" / "SCBE-AETHERMOORE-live-sync" / "docs"
    )
    if not docs_dir.exists():
        return pairs

    for md in sorted(docs_dir.rglob("*.md")):
        try:
            text = md.read_text(encoding="utf-8", errors="replace")
            if len(text) > 500:
                pairs.extend(text_to_sft_pairs(text, source="SCBE Docs Archive", chapter=md.stem))
        except Exception:
            continue
    return pairs


def load_izack_realmforge() -> list[dict]:
    """Load AI Workspace, SCBE Architecture, Game Design, Math Research from Izack Realmforge."""
    pairs = []
    base = HOME / "OneDrive" / "Dropbox" / "Izack Realmforge"
    if not base.exists():
        base = HOME / "OneDrive" / "Imports" / "issdandavis7795@gmail.com - Dropbox" / "Izack Realmforge"
    if not base.exists():
        return pairs

    for subdir in ["AI Workspace", "SCBE Architecture", "Game Design", "Math Research"]:
        d = base / subdir
        if not d.exists():
            continue
        for f in sorted(d.rglob("*")):
            if f.suffix in (".md", ".txt") and f.stat().st_size > 500:
                text = f.read_text(encoding="utf-8", errors="replace")
                pairs.extend(text_to_sft_pairs(text, source=f"Realmforge/{subdir}", chapter=f.stem))
    return pairs


def load_dropbox_mega_texts() -> list[dict]:
    """Load the big text dumps from OneDrive/Dropbox mirror."""
    pairs = []
    big_files = [
        (HOME / "OneDrive" / "Dropbox" / "AVALON BOOK SHIT" / "EVERYTHING FILE.txt", "Everything File"),
        (HOME / "OneDrive" / "Dropbox" / "AVALON BOOK SHIT" / "ChatGPT 4.5.txt", "ChatGPT 4.5"),
        (HOME / "OneDrive" / "Dropbox" / "AVALON BOOK SHIT" / "A FUll log.txt", "Full Log"),
        (HOME / "OneDrive" / "Dropbox" / "700000 characters.txt", "700K Characters"),
        (HOME / "OneDrive" / "Documents" / "Document_root_3.docx", "Document Root 3"),
        (HOME / "OneDrive" / "Documents" / "New chat,BIG DUMP.docx", "Big Dump"),
        (HOME / "OneDrive" / "Documents" / "Thought for 13s.docx", "Thought 13s"),
        (HOME / "OneDrive" / "Documents" / "Maasive thing code.docx", "Massive Code"),
        (HOME / "OneDrive" / "Documents" / "Structure of the book i think.docx", "Book Structure"),
        (HOME / "OneDrive" / "Documents" / "This beautifully crafted Prologue.docx", "Prologue"),
        (HOME / "OneDrive" / "Documents" / "Long logs.docx", "Long Logs"),
        (HOME / "OneDrive" / "Documents" / "Terramare.docx", "Terramare"),
        (HOME / "OneDrive" / "Documents" / "Entire chat log 2 missing begining.txt", "Chat Log 2"),
        (HOME / "OneDrive" / "Documents" / "#Below is the refined, full-length.txt", "Refined Full-Length"),
        (HOME / "OneDrive" / "Documents" / "The Spiral of Avalon A Complete N.txt", "Spiral Complete"),
        (HOME / "OneDrive" / "Documents" / "The Spiral of Avalon.VERY GOOD.1st draft.txt", "Spiral 1st Draft"),
        (HOME / "OneDrive" / "Documents" / "#DarkSetting, Happy Ending.txt", "Dark Setting"),
        (HOME / "OneDrive" / "Lore_Drafts_and_Chat_Exports" / "full game logs new 20205.odt", "Game Logs ODT"),
        (
            HOME / "OneDrive" / "FINAL DRAFTS" / "Architecture and the bulindg of the Relams of Pollyoneth.rtf",
            "Architecture RTF",
        ),
    ]

    seen = set()
    for path, source in big_files:
        if not path.exists():
            continue
        try:
            if path.suffix == ".txt":
                text = path.read_text(encoding="utf-8", errors="replace")
            elif path.suffix in (".docx",):
                try:
                    import docx

                    doc = docx.Document(str(path))
                    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except ImportError:
                    continue
            elif path.suffix == ".rtf":
                # Basic RTF text extraction
                raw = path.read_text(encoding="utf-8", errors="replace")
                text = re.sub(r"\\[a-z]+\d*\s?", "", raw)
                text = re.sub(r"[{}]", "", text)
            elif path.suffix == ".odt":
                continue  # Skip ODT unless we add odfpy
            else:
                continue

            h = hashlib.md5(text[:5000].encode()).hexdigest()
            if h in seen or len(text) < 200:
                continue
            seen.add(h)
            pairs.extend(text_to_sft_pairs(text, source=source))
        except Exception:
            continue
    return pairs


def load_avalon_reference_pack() -> list[dict]:
    """Load CoreLore and Guides from the Avalon Reference Pack."""
    pairs = []
    ref_dir = HOME / "OneDrive" / "Books" / "Avalon_Reference_Pack"
    if not ref_dir.exists():
        return pairs

    for subdir in ["CoreLore", "Guides"]:
        d = ref_dir / subdir
        if not d.exists():
            continue
        for f in sorted(d.iterdir()):
            if f.suffix in (".txt", ".md") and f.stat().st_size > 500:
                text = f.read_text(encoding="utf-8", errors="replace")
                pairs.extend(text_to_sft_pairs(text, source=f"Avalon Ref/{subdir}", chapter=f.stem))
    return pairs


def load_aethermoore_protocol() -> list[dict]:
    """Load SCBE patent specs and protocol docs from Downloads."""
    pairs = []
    proto_dir = HOME / "OneDrive" / "Downloads" / "Aethremore protocol"
    if not proto_dir.exists():
        return pairs
    for f in sorted(proto_dir.iterdir()):
        if f.suffix in (".md", ".txt") and f.stat().st_size > 500:
            text = f.read_text(encoding="utf-8", errors="replace")
            pairs.extend(text_to_sft_pairs(text, source="Aethermoore Protocol", chapter=f.stem))
    return pairs


def load_avalon_github_archive() -> list[dict]:
    """Load writing drafts and archive from Documents/GitHub/Avalon."""
    pairs = []
    avalon_dir = HOME / "OneDrive" / "Documents" / "GitHub" / "Avalon"
    if not avalon_dir.exists():
        return pairs

    seen = set()
    for subdir in ["writing_drafts", "archive", "lore", "docs/avalon_materials"]:
        d = avalon_dir / subdir
        if not d.exists():
            continue
        for f in sorted(d.rglob("*")):
            if f.suffix in (".txt", ".md", ".markdown") and f.stat().st_size > 500:
                try:
                    text = f.read_text(encoding="utf-8", errors="replace")
                    h = hashlib.md5(text[:3000].encode()).hexdigest()
                    if h not in seen:
                        seen.add(h)
                        pairs.extend(text_to_sft_pairs(text, source=f"Avalon GitHub/{subdir}", chapter=f.stem))
                except OSError:
                    continue
    return pairs


def load_arxiv_metadata() -> list[dict]:
    """Convert arXiv paper metadata into SFT pairs about research topics."""
    pairs = []
    arxiv_dir = ROOT / "training" / "intake" / "arxiv"
    if not arxiv_dir.exists():
        return pairs
    for jf in sorted(arxiv_dir.rglob("*.json")):
        try:
            data = json.loads(jf.read_text(encoding="utf-8"))
            title = data.get("title", "")
            abstract = data.get("abstract", data.get("summary", ""))
            if title and abstract:
                tongue = infer_tongue(title + " " + abstract)
                pairs.append(
                    {
                        "instruction": f"Summarize this research paper: {title}",
                        "response": abstract[:1500],
                        "metadata": {
                            "source": "arxiv",
                            "tongue": tongue,
                            "arxiv_id": data.get("id", ""),
                            "type": "mega_ingest",
                        },
                        "encoding_tongue": tongue,
                        "timestamp": time.time(),
                    }
                )
        except Exception:
            continue
    return pairs


def load_six_tongues_wiki() -> list[dict]:
    """Load Six Tongues enhanced wiki and codex from training/raw/."""
    pairs = []
    raw_dir = ROOT / "training" / "raw"
    if not raw_dir.exists():
        return pairs
    for name in [
        "six_tongues_enhanced_v2.md",
        "six_tongues_full_wiki_20260218.md",
        "spiralverse_canonical_linguistic_codex_v1_seed_20260218.md",
    ]:
        path = raw_dir / name
        if path.exists():
            text = path.read_text(encoding="utf-8", errors="replace")
            pairs.extend(text_to_sft_pairs(text, source="Six Tongues Wiki", chapter=name))
    return pairs


# ============================================================
# DEDUPLICATION
# ============================================================


def deduplicate(records: list[dict]) -> list[dict]:
    """Deduplicate by content hash of instruction+response."""
    seen = set()
    deduped = []
    for rec in records:
        key_text = rec.get("instruction", rec.get("prompt", "")) + rec.get("response", rec.get("text", ""))
        h = hashlib.md5(key_text.encode(errors="replace")).hexdigest()
        if h not in seen:
            seen.add(h)
            deduped.append(rec)
    return deduped


# ============================================================
# TONGUE BALANCE REPORT
# ============================================================


def tongue_report(records: list[dict]) -> dict:
    """Count tongue distribution and flag imbalances."""
    counts = {t: 0 for t in TONGUE_KEYS}
    for rec in records:
        meta = rec.get("metadata", {})
        if isinstance(meta, str):
            meta = {}
        tongue = meta.get("tongue", rec.get("encoding_tongue", ""))
        if tongue not in TONGUE_KEYS:
            tongue = infer_tongue(rec.get("instruction", "") + " " + rec.get("response", ""))
        counts[tongue] = counts.get(tongue, 0) + 1

    total = sum(counts.values())
    target = total / 6 if total > 0 else 0
    imbalance = {}
    for t, c in counts.items():
        diff = c - target
        pct = (c / total * 100) if total else 0
        imbalance[t] = {"count": c, "pct": round(pct, 1), "delta": round(diff)}

    return {"total": total, "tongues": imbalance, "target_per_tongue": round(target)}


# ============================================================
# MAIN
# ============================================================


def main():
    parser = argparse.ArgumentParser(description="Mega Ingest Pipeline — ALL sources")
    parser.add_argument("--ingest", action="store_true", help="Run full ingest (not just scan)")
    parser.add_argument("--tetris", action="store_true", help="Run Tetris embedding after ingest")
    parser.add_argument("--push-hf", action="store_true", help="Push to HuggingFace after embed")
    parser.add_argument("--skip-chatgpt", action="store_true", help="Skip ChatGPT HTML (slow)")
    parser.add_argument("--skip-gemini", action="store_true", help="Skip Gemini HTML")
    args = parser.parse_args()

    print("=" * 70)
    print("MEGA INGEST PIPELINE")
    print("All Sources -> Tongue-Tagged SFT -> Tetris Embedding")
    print("=" * 70)

    sources = {}
    t_total = time.time()

    loaders = [
        ("repo_jsonl", "ALL repo JSONL files", load_all_repo_jsonl),
        ("six_tongues_protocol", "Six Tongues Protocol book", load_six_tongues_protocol),
        ("six_tongues_wiki", "Six Tongues Wiki + Codex", load_six_tongues_wiki),
        ("reincarnated", "Reincarnated novel", load_reincarnated_novel),
        ("everweave_text", "Everweave corpus + ChatGPT logs", load_everweave_text),
        ("theory_doc", "500-page theory document", load_theory_doc),
        ("spiralverse_mega", "Spiralverse Mega Document (16.8MB)", load_spiralverse_mega),
        ("phone_backup", "Phone backup lore + ChoiceScript", load_phone_backup_lore),
        ("lore_drafts", "Lore drafts + manuscripts (DOCX)", load_lore_drafts),
        ("dropbox_mega_texts", "Dropbox mega text dumps", load_dropbox_mega_texts),
        ("dropbox_story_series", "Dropbox story series chapters", load_dropbox_story_series),
        ("stp_editorial", "STP editorial + reference notes", load_dropbox_stp_editorial),
        ("spiral_chapters", "Spiral of Pollyoneth chapters", load_dropbox_spiral_chapters),
        ("izack_realmforge", "Izack Realmforge workspace", load_izack_realmforge),
        ("avalon_ref_pack", "Avalon Reference Pack (lore+guides)", load_avalon_reference_pack),
        ("avalon_github", "Avalon GitHub archive + drafts", load_avalon_github_archive),
        ("aethermoore_protocol", "Aethermoore Protocol patent specs", load_aethermoore_protocol),
        ("scbe_archive_docs", "SCBE Docs Archive (Dropbox sync)", load_scbe_archive_docs),
        ("choicescript", "ChoiceScript game scenes", load_choicescript),
        ("obsidian", "Obsidian vault notes", load_obsidian_vault),
        ("arxiv", "arXiv paper metadata", load_arxiv_metadata),
    ]

    total_loaders = len(loaders) + 2  # +2 for notion and ai_exports

    for idx, (name, desc, loader) in enumerate(loaders, 1):
        print(f"\n[{idx}/{total_loaders}] Loading {desc}...")
        t0 = time.time()
        sources[name] = loader()
        print(f"  {len(sources[name])} records ({time.time()-t0:.1f}s)")

    # Notion (special — conditionally unzips)
    idx = len(loaders) + 1
    print(f"\n[{idx}/{total_loaders}] Loading Notion (knowledge JSON + workspace ZIP)...")
    t0 = time.time()
    notion_k = load_notion_knowledge()
    notion_z = load_notion_zip() if args.ingest else []
    sources["notion"] = notion_k + notion_z
    print(f"  {len(sources['notion'])} records ({time.time()-t0:.1f}s)")

    # External AI exports (special — has skip flags)
    idx = total_loaders
    print(f"\n[{idx}/{total_loaders}] Loading external AI exports (ChatGPT HTML, Gemini, Grok)...")
    t0 = time.time()
    ext = []
    if not args.skip_chatgpt:
        ext.extend(load_chatgpt_html())
    if not args.skip_gemini:
        ext.extend(load_gemini_export())
    ext.extend(load_grok_drop())
    sources["ai_exports"] = ext
    print(f"  {len(sources['ai_exports'])} records ({time.time()-t0:.1f}s)")

    # ================================================================
    # COMBINE + DEDUP
    # ================================================================
    print("\n" + "=" * 70)
    print("COMBINING AND DEDUPLICATING...")
    all_records = []
    for _name, recs in sources.items():
        all_records.extend(recs)
    pre_dedup = len(all_records)
    all_records = deduplicate(all_records)
    print(f"  {pre_dedup} raw -> {len(all_records)} after dedup")

    # ================================================================
    # TONGUE DISTRIBUTION
    # ================================================================
    report = tongue_report(all_records)
    print(f"\nTONGUE DISTRIBUTION ({report['total']} total, target ~{report['target_per_tongue']}/tongue):")
    for t in TONGUE_KEYS:
        info = report["tongues"][t]
        bar = "#" * (info["count"] // max(1, report["target_per_tongue"] // 20))
        status = "OK" if abs(info["delta"]) < report["target_per_tongue"] * 0.3 else "SKEW"
        print(f"  {t}: {info['count']:>6} ({info['pct']:>5.1f}%) {bar} [{status}]")

    # ================================================================
    # SOURCE BREAKDOWN
    # ================================================================
    print(f"\nSOURCE BREAKDOWN:")
    for name, recs in sorted(sources.items(), key=lambda x: -len(x[1])):
        print(f"  {name:.<35} {len(recs):>6} records")

    # ================================================================
    # EXPORT
    # ================================================================
    if args.ingest:
        print(f"\nEXPORTING...")
        output_path = ROOT / "training-data" / "mega_ingest_sft.jsonl"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            for rec in all_records:
                f.write(json.dumps(rec, ensure_ascii=False, default=str) + "\n")
        size_mb = output_path.stat().st_size / (1024 * 1024)
        print(f"  Wrote {len(all_records)} records to {output_path.name} ({size_mb:.1f}MB)")

        # Also write a manifest
        manifest = {
            "timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ"),
            "total_records": len(all_records),
            "pre_dedup": pre_dedup,
            "tongue_report": report,
            "sources": {k: len(v) for k, v in sources.items()},
            "output_file": str(output_path),
            "output_size_mb": round(size_mb, 2),
        }
        manifest_path = ROOT / "artifacts" / "mega_ingest_manifest.json"
        manifest_path.parent.mkdir(parents=True, exist_ok=True)
        manifest_path.write_text(json.dumps(manifest, indent=2, default=str))
        print(f"  Manifest: {manifest_path}")

        # Tetris embedding
        if args.tetris or args.push_hf:
            print(f"\nTETRIS EMBEDDING...")
            os.environ["TOKENIZERS_PARALLELISM"] = "false"
            from src.kernel.tetris_embedder import TetrisEmbedder

            tetris = TetrisEmbedder("all-MiniLM-L6-v2")
            tetris_out = ROOT / "training-data" / "mega_tetris_enriched_sft.jsonl"

            # Import the export function from the tetris pipeline
            sys.path.insert(0, str(ROOT / "scripts"))
            from tetris_training_pipeline import export_tetris_sft, benchmark_embeddings, infer_tongue as tt_infer

            texts = [
                (r.get("instruction", r.get("prompt", "")) + " " + r.get("response", ""))[:500] for r in all_records
            ]
            tongues = [tt_infer(r) for r in all_records]
            tiers = [
                r.get("metadata", {}).get("tier", 1) if isinstance(r.get("metadata"), dict) else 1 for r in all_records
            ]

            # Benchmark a sample
            sample_n = min(500, len(texts))
            metrics = benchmark_embeddings(tetris, texts[:sample_n], tongues[:sample_n], tiers[:sample_n])
            print(f"  Diversity (sample {sample_n}): {metrics['diversity_score']:.0f}/100")
            print(f"  Octree buckets: {metrics['n_buckets']}")

            t0 = time.time()
            written = export_tetris_sft(all_records, tetris, tetris_out)
            print(f"  Wrote {written} Tetris-enriched records in {time.time()-t0:.1f}s")
            tetris_mb = tetris_out.stat().st_size / (1024 * 1024)
            print(f"  Output: {tetris_out.name} ({tetris_mb:.1f}MB)")

            if args.push_hf:
                print(f"\nPUSHING TO HUGGINGFACE...")
                from tetris_training_pipeline import push_to_hf, load_hf_token

                hf_token = load_hf_token()
                repo_id = "issdandavis/scbe-aethermoore-training-data"
                if hf_token:
                    ok1 = push_to_hf(tetris_out, repo_id, hf_token)
                    ok2 = push_to_hf(output_path, repo_id, hf_token)
                    print(f"  Tetris JSONL: {'OK' if ok1 else 'FAILED'}")
                    print(f"  Raw JSONL: {'OK' if ok2 else 'FAILED'}")
                else:
                    print("  SKIP: No HF token found")
    else:
        print(f"\n  DRY RUN — use --ingest to write output")
        print(f"  Grok drop zone: training/intake/grok/ (place export there)")

    elapsed = time.time() - t_total
    print(f"\n{'='*70}")
    print(f"MEGA INGEST COMPLETE — {len(all_records)} records in {elapsed:.1f}s")
    print(f"{'='*70}")


if __name__ == "__main__":
    main()
