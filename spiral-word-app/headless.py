"""
@file headless.py
@module spiral-word-app/headless
@layer Layer 12, Layer 13, Layer 14
@component Headless Word Document Generation

Produces governed .docx (Microsoft Word) artifacts from SpiralWord
documents without requiring the UI. Two surfaces:

- render_docx(doc, ...) -> bytes
    Convert any SyncEngine Document into a .docx file. Paragraphs split
    on blank lines; Markdown-style "# " / "## " prefixes promote to H1/H2.

- generate_record(prompt, provider, options, doc_id, site_id, ...) -> dict
    Full AI-driven pipeline: classify intent (L12), gate via governance
    (L13), drive an AI provider, replace doc content, render .docx, and
    write a deterministic audit record (L14). Returns both the binary
    artifact and the record entry for downstream archival.

Records are SHA-256 keyed on the canonical (doc_id, content, site_id,
timestamp) tuple so the same generation is reproducible from the audit
log alone.
"""

from __future__ import annotations

import hashlib
import io
import logging
import time
from dataclasses import dataclass
from typing import Optional

from sync_engine import Document
from governance import audit_log, classify_intent
from ai_ports import ai_ports
from braid_ledger import get_braid_ledger

logger = logging.getLogger("spiralword.headless")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class HeadlessRecord:
    """Audit record returned for every headless generation."""

    record_id: str
    doc_id: str
    site_id: str
    provider: str
    tongue: str
    confidence: float
    prompt_sha256: str
    docx_sha256: str
    docx_bytes: int
    timestamp: float
    title: Optional[str] = None
    author: Optional[str] = None
    phdm_node_id: Optional[str] = None
    loop_index: Optional[int] = None
    braid_receipt: Optional[str] = None
    loop_root: Optional[str] = None

    def to_dict(self) -> dict:
        data = {
            "record_id": self.record_id,
            "doc_id": self.doc_id,
            "site_id": self.site_id,
            "provider": self.provider,
            "tongue": self.tongue,
            "confidence": self.confidence,
            "prompt_sha256": self.prompt_sha256,
            "docx_sha256": self.docx_sha256,
            "docx_bytes": self.docx_bytes,
            "timestamp": self.timestamp,
            "title": self.title,
            "author": self.author,
        }
        if self.phdm_node_id is not None:
            data["phdm_node_id"] = self.phdm_node_id
        if self.loop_index is not None:
            data["loop_index"] = self.loop_index
        if self.braid_receipt is not None:
            data["braid_receipt"] = self.braid_receipt
        if self.loop_root is not None:
            data["loop_root"] = self.loop_root
        return data


def _split_paragraphs(text: str) -> list[str]:
    """Split a flat string into paragraph blocks on blank-line boundaries."""
    if not text:
        return []
    blocks: list[str] = []
    buf: list[str] = []
    for line in text.splitlines():
        if line.strip() == "":
            if buf:
                blocks.append("\n".join(buf))
                buf = []
        else:
            buf.append(line)
    if buf:
        blocks.append("\n".join(buf))
    return blocks


def render_docx(
    doc: Document,
    *,
    title: Optional[str] = None,
    author: Optional[str] = None,
    include_audit_footer: bool = True,
) -> bytes:
    """
    Render a SyncEngine Document as a .docx file.

    Markdown-lite handling:
        - Lines starting with "# "  -> Heading 1
        - Lines starting with "## " -> Heading 2
        - Blank line                -> paragraph break

    Args:
        doc: The Document to export.
        title: Optional title injected as Heading 1 above the body.
        author: Optional core-properties author string.
        include_audit_footer: If True, append a "Record" line with
            doc_id, version, and op count (matches the audit_log entry).

    Returns:
        The .docx file as raw bytes.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for headless Word export; " "install it via `pip install python-docx`"
        ) from exc

    docx = DocxDocument()

    if author:
        docx.core_properties.author = author
    if title:
        docx.core_properties.title = title
        docx.add_heading(title, level=1)

    for block in _split_paragraphs(doc.text):
        first_line, *rest = block.split("\n", 1)
        if first_line.startswith("## "):
            docx.add_heading(first_line[3:].strip(), level=2)
            if rest:
                docx.add_paragraph(rest[0])
        elif first_line.startswith("# "):
            docx.add_heading(first_line[2:].strip(), level=1)
            if rest:
                docx.add_paragraph(rest[0])
        else:
            docx.add_paragraph(block)

    if include_audit_footer:
        snap = doc.snapshot()
        footer = docx.add_paragraph()
        footer.add_run(
            f"Record: doc_id={snap['doc_id']}  version={snap['version']}  " f"ops={snap['op_count']}"
        ).italic = True

    buf = io.BytesIO()
    docx.save(buf)
    return buf.getvalue()


def _sha256_hex(data: bytes | str) -> str:
    if isinstance(data, str):
        data = data.encode("utf-8")
    return hashlib.sha256(data).hexdigest()


def generate_record(
    *,
    sync,
    doc_id: str,
    prompt: str,
    provider: str = "echo",
    options: Optional[dict] = None,
    site_id: str = "headless",
    title: Optional[str] = None,
    author: Optional[str] = None,
    replace: bool = True,
) -> tuple[bytes, HeadlessRecord, str]:
    """
    Full headless generation pipeline.

    1. Call AI provider via ai_ports (which runs L12 classify + L13 gate).
    2. Replace (or append to) the document with the AI output.
    3. Render the document as .docx.
    4. Write a deterministic audit_log entry.
    5. Return (docx_bytes, record, ai_text).

    Args:
        sync: The SyncEngine instance owning the documents.
        doc_id: Target document identifier (created if missing).
        prompt: The AI prompt to dispatch.
        provider: AI port name (default "echo" for offline/testing).
        options: Provider options (model, temperature, ...).
        site_id: Audit attribution for the writer.
        title: Optional document title (used in .docx core properties).
        author: Optional .docx core-properties author.
        replace: If True, replace doc content. If False, append.

    Returns:
        (docx_bytes, record, ai_text)

    Raises:
        ValueError: When governance blocks the prompt (status string from
            the AI port starts with "[BLOCKED]").
        RuntimeError: When the AI provider fails.
    """
    ai_text = ai_ports.call(prompt, provider=provider, options=options)
    if ai_text.startswith("[BLOCKED]"):
        raise ValueError(ai_text[len("[BLOCKED]") :].strip() or "blocked by governance")
    if ai_text.startswith("[ERROR]"):
        raise RuntimeError(ai_text[len("[ERROR]") :].strip() or "ai provider failed")

    doc = sync.get_or_create(doc_id)
    if replace:
        doc.replace_all(ai_text, site_id=site_id)
    else:
        doc.insert(doc.length, ai_text, site_id=site_id)

    docx_bytes = render_docx(doc, title=title, author=author)

    tongue, confidence = classify_intent(prompt)
    timestamp = time.time()
    docx_hash = _sha256_hex(docx_bytes)
    prompt_hash = _sha256_hex(prompt)
    braid_receipt = get_braid_ledger().commit(prompt_hash, docx_hash)
    if not braid_receipt.tube_ok:
        raise ValueError("braid tube violation")
    record_id = _sha256_hex(
        f"{braid_receipt.loop_root}|{braid_receipt.phdm_node_id}|{prompt_hash}|{docx_hash}"
    )[:32]

    record = HeadlessRecord(
        record_id=record_id,
        doc_id=doc_id,
        site_id=site_id,
        provider=provider,
        tongue=tongue,
        confidence=confidence,
        prompt_sha256=prompt_hash,
        docx_sha256=docx_hash,
        docx_bytes=len(docx_bytes),
        timestamp=timestamp,
        title=title,
        author=author,
        phdm_node_id=braid_receipt.phdm_node_id,
        loop_index=braid_receipt.loop_index,
        braid_receipt=braid_receipt.hmac_tag,
        loop_root=braid_receipt.loop_root,
    )

    audit_log.record(
        doc_id=doc_id,
        site_id=site_id,
        action="headless_generate",
        op_checksum=docx_hash[:16],
        governance_decision=f"AI:{provider} record={record_id}",
        tongue=tongue,
        confidence=confidence,
        phdm_node_id=braid_receipt.phdm_node_id,
        loop_index=braid_receipt.loop_index,
        braid_receipt=braid_receipt.hmac_tag,
    )

    logger.info(
        "headless: doc=%s provider=%s record=%s bytes=%d",
        doc_id,
        provider,
        record_id,
        len(docx_bytes),
    )
    return docx_bytes, record, ai_text
