import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from scripts.publish.kdp_acceptance_gate import run_gate
from scripts.publish.kdp_story_quality_gate import score_packet
from scripts.publish.kdp_visual_format_report import run_report


def test_story_quality_gate_scores_publishable_packet():
    packet = {
        "title": "Demo",
        "author": "Issac Daniel Davis",
        "minimum_publishable_score": 75,
        "minimum_portfolio_score": 85,
        "dimensions": {
            "whole_story_coherence": {"score": 8, "evidence_refs": ["a"]},
            "character_continuity": {"score": 8, "evidence_refs": ["a"]},
            "plot_causality": {"score": 8, "evidence_refs": ["a"]},
            "prose_readability": {"score": 7, "evidence_refs": ["a"]},
            "originality_and_reader_value": {"score": 8, "evidence_refs": ["a"]},
            "portfolio_fit": {"score": 8, "evidence_refs": ["a"]},
            "production_readiness": {"score": 7, "evidence_refs": ["a"]},
        },
    }

    report = score_packet(packet)

    assert report["decision"] == "PASS"
    assert report["quality_tier"] == "publishable_with_review"
    assert report["score"] >= 75


def test_acceptance_gate_requires_disclosure_and_review_reports(tmp_path: Path):
    book_root = tmp_path / "book"
    reader = book_root / "reader-edition"
    stage = tmp_path / "stage"
    reader.mkdir(parents=True)
    stage.mkdir()

    (reader / "ch01.md").write_text("# Chapter 1\n\n" + ("good story words " * 8000), encoding="utf-8")
    (reader / "ch02.md").write_text("# Chapter 2\n\n" + ("more coherent words " * 8000), encoding="utf-8")
    (reader / "zz-back-matter.md").write_text("# Back Matter\n\n" + ("author note " * 5000), encoding="utf-8")
    (book_root / "build_kdp.py").write_text('AUTHOR = "Issac Daniel Davis"', encoding="utf-8")
    (book_root / "HOUSE_STYLE.md").write_text("Issac Daniel Davis\n", encoding="utf-8")
    (book_root / "INDEX.md").write_text("Issac Daniel Davis\n", encoding="utf-8")
    for name in ["KDP_PAPERBACK_FORMAT_SPEC.md", "FINAL_TOPOGRAPHY.md", "MARKET_COMP_ANALYSIS_2025_2026.md"]:
        (book_root / name).write_text("review evidence", encoding="utf-8")
    artifact = book_root / "the-six-tongues-protocol-kdp.docx"
    artifact.write_bytes(b"x" * 20_000)
    packet = book_root / "kdp_submission_packet.json"
    packet.write_text(
        json.dumps(
            {
                "ai_content_disclosure": {
                    "text": "ai_generated",
                    "cover_images": "ai_generated",
                    "interior_images": "none",
                    "translations": "none",
                },
                "human_review": {"approved_for_kdp_upload": True},
            }
        ),
        encoding="utf-8",
    )
    (stage / "story-quality-gate.json").write_text(json.dumps({"decision": "PASS", "score": 80}), encoding="utf-8")
    (stage / "visual-format-report.json").write_text(json.dumps({"decision": "PASS"}), encoding="utf-8")

    report = run_gate(book_root, packet, artifact, stage)

    assert report["decision"] == "PASS"
    assert report["score"] == report["max_score"]


def test_visual_format_report_reads_trim_margins_and_font(tmp_path: Path):
    docx = tmp_path / "demo.docx"
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)
    for i in range(12):
        heading = doc.add_paragraph()
        run = heading.add_run(f"Chapter {i + 1}: Test")
        run.font.name = "Garamond"
        run.font.size = Pt(18)
        para = doc.add_paragraph()
        body = para.add_run("This is body text for visual rhythm inspection.")
        body.font.name = "Garamond"
        body.font.size = Pt(11)
    doc.save(docx)

    report = run_report(docx)

    assert report["decision"] == "PASS"
    assert report["page"]["page_width_in"] == 5.5
    assert report["page"]["page_height_in"] == 8.5
