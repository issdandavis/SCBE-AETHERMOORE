"""
Generate USPTO filing documents as Word (.docx) files.

Creates:
  1. PTO_SB_15A_Micro_Entity_Certification.docx
  2. PTO_SB_16_Cover_Sheet.docx
  3. Cover_Letter_Missing_Parts.docx
  4. Combined_Filing_Package.docx (all-in-one)

Usage:
  python generate_filing_docs.py
"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Applicant Info ──────────────────────────────────────────────────
APPLICANT = {
    "full_name": "Issac Daniel Davis",
    "address": "2361 E 5th Ave",
    "city": "Port Angeles",
    "state": "WA",
    "zip": "98362",
    "country": "US",
    "phone": "360-808-0876",
    "email": "",  # ← ADD YOUR EMAIL HERE
}

APPLICATION = {
    "number": "63/961,403",
    "filing_date": "January 15, 2026",
    "title": "Context-Bound Cryptographic Authorization System",
    "docket": "SCBE-2026-001",
    "claims_count": 16,
    "entity_status": "Micro Entity",
    "fee": "$65.00",
}

TODAY = date.today().strftime("%m/%d/%Y")
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_field_table(doc, rows):
    """Add a two-column field:value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, (field, value) in enumerate(rows):
        cell_f = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_f.text = field
        cell_v.text = str(value)
        for paragraph in cell_f.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
        for paragraph in cell_v.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    doc.add_paragraph("")
    return table


def add_checkbox_line(doc, text, checked=True):
    mark = "[X]" if checked else "[ ]"
    p = doc.add_paragraph(f"{mark}  {text}")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


# ────────────────────────────────────────────────────────────────────
# DOCUMENT 1: Cover Letter
# ────────────────────────────────────────────────────────────────────
def build_cover_letter():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{APPLICANT['full_name']}\n")
    p.add_run(f"{APPLICANT['address']}\n")
    p.add_run(f"{APPLICANT['city']}, {APPLICANT['state']} {APPLICANT['zip']}\n")
    p.add_run(f"Phone: {APPLICANT['phone']}\n")
    if APPLICANT["email"]:
        p.add_run(f"Email: {APPLICANT['email']}\n")
    p.add_run(f"\n{TODAY}")

    doc.add_paragraph("")

    # Addressee
    p = doc.add_paragraph()
    p.add_run("Commissioner for Patents\n")
    p.add_run("United States Patent and Trademark Office\n")
    p.add_run("P.O. Box 1450\n")
    p.add_run("Alexandria, VA 22313-1450")

    doc.add_paragraph("")

    # Subject line
    p = doc.add_paragraph()
    run = p.add_run("Re: Response to Notice to File Missing Parts of Application")
    run.bold = True
    doc.add_paragraph("")

    add_field_table(doc, [
        ("Application Number", APPLICATION["number"]),
        ("Filing Date", APPLICATION["filing_date"]),
        ("Title of Invention", APPLICATION["title"]),
        ("Applicant", APPLICANT["full_name"]),
        ("Docket Number", APPLICATION["docket"]),
    ])

    # Body
    p = doc.add_paragraph("Dear Commissioner:")
    doc.add_paragraph("")

    doc.add_paragraph(
        "In response to the Notice to File Missing Parts of the above-identified "
        "provisional patent application, the undersigned applicant respectfully "
        "submits the following:"
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "1.  Certification of Micro Entity Status (PTO/SB/15A) — attached hereto."
    )
    doc.add_paragraph(
        "2.  Provisional Application Cover Sheet (PTO/SB/16) — attached hereto."
    )
    doc.add_paragraph(
        f"3.  Filing fee payment of {APPLICATION['fee']} (Micro Entity rate) — "
        "submitted via Patent Center electronic payment."
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "The applicant certifies qualification as a Micro Entity under "
        "37 CFR 1.29(a). The required certification form (PTO/SB/15A) is "
        "included with this submission."
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        "The applicant respectfully requests that the Office accept this "
        "response and complete processing of the above-identified provisional "
        "patent application."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Respectfully submitted,")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature block
    p = doc.add_paragraph()
    p.add_run("_" * 40)
    doc.add_paragraph(APPLICANT["full_name"])
    doc.add_paragraph("Pro Se Applicant")
    p = doc.add_paragraph(f"Date: {TODAY}")
    p = doc.add_paragraph(f"Telephone: {APPLICANT['phone']}")

    path = os.path.join(OUTPUT_DIR, "Cover_Letter_Missing_Parts.docx")
    doc.save(path)
    print(f"  Created: {path}")
    return doc


# ────────────────────────────────────────────────────────────────────
# DOCUMENT 2: PTO/SB/15A — Micro Entity Certification
# ────────────────────────────────────────────────────────────────────
def build_pto_sb_15a():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    h = doc.add_heading("CERTIFICATION OF MICRO ENTITY STATUS", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(37 CFR 1.29(a) — Gross Income Basis)")
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PTO/SB/15A")
    run.bold = True
    run.font.size = Pt(10)

    doc.add_paragraph("")

    # Application info
    add_heading(doc, "Application Information", level=2)
    add_field_table(doc, [
        ("Application Number", APPLICATION["number"]),
        ("Filing Date", APPLICATION["filing_date"]),
        ("First Named Inventor", APPLICANT["full_name"]),
        ("Title of Invention", APPLICATION["title"]),
        ("Docket Number (optional)", APPLICATION["docket"]),
    ])

    # Certification statements
    add_heading(doc, "Applicant Hereby Certifies the Following:", level=2)

    add_checkbox_line(
        doc,
        "The applicant qualifies as a micro entity as defined in "
        "37 CFR 1.29(a).",
        checked=True,
    )

    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor nor a joint inventor "
        "has been named as the inventor or a joint inventor on more than "
        "four (4) previously filed patent applications, other than "
        "applications filed in another country, provisional applications "
        "under 35 U.S.C. 111(b), or international applications under "
        "35 U.S.C. 371.",
        checked=True,
    )

    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor nor a joint inventor, "
        "in the calendar year preceding the calendar year in which the "
        "applicable fee is being paid, had a gross income (as defined "
        "in section 61(a) of the Internal Revenue Code of 1986) exceeding "
        "three times the median household income for that preceding "
        "calendar year, as most recently reported by the Bureau of the "
        "Census (current posted threshold: $251,190).",
        checked=True,
    )

    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor nor a joint inventor "
        "has assigned, granted, or conveyed, and is not under an obligation "
        "by contract or law to assign, grant, or convey, a license or "
        "other ownership interest in the application concerned to an entity "
        "that, in the calendar year preceding the calendar year in which "
        "the applicable fee is being paid, had a gross income exceeding "
        "three times the median household income for that preceding "
        "calendar year.",
        checked=True,
    )

    doc.add_paragraph("")

    # Signature
    add_heading(doc, "Signature", level=2)
    add_field_table(doc, [
        ("Signature", f"/{APPLICANT['full_name']}/"),
        ("Name (Printed)", APPLICANT["full_name"]),
        ("Date", TODAY),
        ("Registration Number", "N/A — Pro Se Applicant"),
        ("Telephone", APPLICANT["phone"]),
    ])

    # Applicant info
    add_heading(doc, "Applicant Information", level=2)
    add_field_table(doc, [
        ("Name", APPLICANT["full_name"]),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
        ("Telephone", APPLICANT["phone"]),
        ("Email", APPLICANT["email"] or "(to be provided)"),
    ])

    # Warning
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("WARNING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    p.add_run(
        "Improperly claiming micro entity status may result in the "
        "patent being held unenforceable. Applicant certifies that all "
        "statements herein are true and correct to the best of their "
        "knowledge."
    )

    path = os.path.join(OUTPUT_DIR, "PTO_SB_15A_Micro_Entity_Certification.docx")
    doc.save(path)
    print(f"  Created: {path}")
    return doc


# ────────────────────────────────────────────────────────────────────
# DOCUMENT 3: PTO/SB/16 — Provisional Cover Sheet
# ────────────────────────────────────────────────────────────────────
def build_pto_sb_16():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    h = doc.add_heading("PROVISIONAL APPLICATION FOR PATENT COVER SHEET", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PTO/SB/16")
    run.bold = True
    run.font.size = Pt(10)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(
        "This is a request for filing a PROVISIONAL APPLICATION FOR PATENT "
        "under 37 CFR 1.53(c)."
    )
    run.italic = True

    doc.add_paragraph("")

    # Title of Invention
    add_heading(doc, "Title of Invention", level=2)
    p = doc.add_paragraph(APPLICATION["title"])
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    doc.add_paragraph("")

    # Inventor
    add_heading(doc, "Inventor(s)", level=2)
    add_field_table(doc, [
        ("Given Name", "Issac Daniel"),
        ("Family Name", "Davis"),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
    ])

    # Correspondence Address
    add_heading(doc, "Correspondence Address", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Direct all correspondence to:")
    run.italic = True
    add_field_table(doc, [
        ("Name", APPLICANT["full_name"]),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
        ("Telephone", APPLICANT["phone"]),
        ("Email", APPLICANT["email"] or "(to be provided)"),
    ])

    # Entity Status
    add_heading(doc, "Entity Status", level=2)
    add_checkbox_line(doc, "Micro Entity (37 CFR 1.29)", checked=True)
    add_checkbox_line(doc, "Small Entity (37 CFR 1.27)", checked=False)
    add_checkbox_line(doc, "Regular Undiscounted", checked=False)
    p = doc.add_paragraph(
        "PTO/SB/15A (Certification of Micro Entity Status) is attached."
    )
    p.runs[0].italic = True

    doc.add_paragraph("")

    # Docket
    add_heading(doc, "Docket Number", level=2)
    doc.add_paragraph(APPLICATION["docket"])

    doc.add_paragraph("")

    # Specification
    add_heading(doc, "Specification", level=2)
    add_field_table(doc, [
        ("Number of Claims", str(APPLICATION["claims_count"])),
        ("Abstract", "1 page"),
        ("Application Number (if existing)", APPLICATION["number"]),
        ("Original Filing Date", APPLICATION["filing_date"]),
    ])

    # Signature
    add_heading(doc, "Signature", level=2)
    add_field_table(doc, [
        ("Signature", f"/{APPLICANT['full_name']}/"),
        ("Name (Printed)", APPLICANT["full_name"]),
        ("Date", TODAY),
        ("Registration Number", "N/A — Pro Se Applicant"),
    ])

    path = os.path.join(OUTPUT_DIR, "PTO_SB_16_Cover_Sheet.docx")
    doc.save(path)
    print(f"  Created: {path}")
    return doc


# ────────────────────────────────────────────────────────────────────
# DOCUMENT 4: Combined Filing Package
# ────────────────────────────────────────────────────────────────────
def build_combined():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title page
    for _ in range(6):
        doc.add_paragraph("")

    h = doc.add_heading("USPTO Filing Package", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Response to Notice to File Missing Parts\n\n").bold = True
    p.add_run(f"Application No. {APPLICATION['number']}\n\n")
    p.add_run(f'"{APPLICATION["title"]}"\n\n')
    p.add_run(f"Applicant: {APPLICANT['full_name']}\n\n")
    p.add_run(f"Date: {TODAY}\n\n")
    p.add_run(f"Filing Fee: {APPLICATION['fee']} (Micro Entity)\n")

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, "Table of Contents", level=1)
    doc.add_paragraph("1.  Cover Letter")
    doc.add_paragraph("2.  PTO/SB/15A — Certification of Micro Entity Status")
    doc.add_paragraph("3.  PTO/SB/16 — Provisional Application Cover Sheet")
    doc.add_paragraph("4.  Filing Checklist")

    doc.add_page_break()

    # ── Section 1: Cover Letter ──
    add_heading(doc, "1. Cover Letter", level=1)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{APPLICANT['full_name']}\n")
    p.add_run(f"{APPLICANT['address']}\n")
    p.add_run(f"{APPLICANT['city']}, {APPLICANT['state']} {APPLICANT['zip']}\n")
    p.add_run(f"Phone: {APPLICANT['phone']}\n")
    p.add_run(f"\n{TODAY}")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Commissioner for Patents\n")
    p.add_run("United States Patent and Trademark Office\n")
    p.add_run("P.O. Box 1450\n")
    p.add_run("Alexandria, VA 22313-1450")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(
        "Re: Response to Notice to File Missing Parts of Application"
    )
    run.bold = True

    doc.add_paragraph("")

    add_field_table(doc, [
        ("Application Number", APPLICATION["number"]),
        ("Filing Date", APPLICATION["filing_date"]),
        ("Title of Invention", APPLICATION["title"]),
        ("Applicant", APPLICANT["full_name"]),
        ("Docket Number", APPLICATION["docket"]),
    ])

    doc.add_paragraph("Dear Commissioner:")
    doc.add_paragraph("")
    doc.add_paragraph(
        "In response to the Notice to File Missing Parts of the above-identified "
        "provisional patent application, the undersigned applicant respectfully "
        "submits the following:"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "1.  Certification of Micro Entity Status (PTO/SB/15A) — attached hereto."
    )
    doc.add_paragraph(
        "2.  Provisional Application Cover Sheet (PTO/SB/16) — attached hereto."
    )
    doc.add_paragraph(
        f"3.  Filing fee payment of {APPLICATION['fee']} (Micro Entity rate) — "
        "submitted via Patent Center electronic payment."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "The applicant certifies qualification as a Micro Entity under "
        "37 CFR 1.29(a). The required certification form (PTO/SB/15A) is "
        "included with this submission."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "The applicant respectfully requests that the Office accept this "
        "response and complete processing of the above-identified provisional "
        "patent application."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Respectfully submitted,")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("_" * 40)
    doc.add_paragraph(APPLICANT["full_name"])
    doc.add_paragraph("Pro Se Applicant")
    doc.add_paragraph(f"Date: {TODAY}")

    doc.add_page_break()

    # ── Section 2: PTO/SB/15A ──
    add_heading(doc, "2. Certification of Micro Entity Status (PTO/SB/15A)", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("37 CFR 1.29(a) — Gross Income Basis")
    run.italic = True

    doc.add_paragraph("")

    add_heading(doc, "Application Information", level=2)
    add_field_table(doc, [
        ("Application Number", APPLICATION["number"]),
        ("Filing Date", APPLICATION["filing_date"]),
        ("First Named Inventor", APPLICANT["full_name"]),
        ("Title of Invention", APPLICATION["title"]),
        ("Docket Number", APPLICATION["docket"]),
    ])

    add_heading(doc, "Applicant Hereby Certifies:", level=2)

    add_checkbox_line(
        doc,
        "The applicant qualifies as a micro entity as defined in 37 CFR 1.29(a).",
    )
    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor has been named as inventor "
        "on more than four (4) previously filed patent applications "
        "(excluding foreign, provisional under 35 U.S.C. 111(b), and "
        "international under 35 U.S.C. 371).",
    )
    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor had gross income in the "
        "preceding calendar year exceeding three times the median household "
        "income (current posted threshold: $251,190).",
    )
    add_checkbox_line(
        doc,
        "Neither the applicant nor the inventor has assigned or is obligated "
        "to assign ownership to an entity exceeding the income threshold.",
    )

    doc.add_paragraph("")
    add_heading(doc, "Signature", level=2)
    add_field_table(doc, [
        ("Signature", f"/{APPLICANT['full_name']}/"),
        ("Name (Printed)", APPLICANT["full_name"]),
        ("Date", TODAY),
        ("Registration Number", "N/A — Pro Se Applicant"),
        ("Telephone", APPLICANT["phone"]),
    ])

    add_heading(doc, "Applicant Information", level=2)
    add_field_table(doc, [
        ("Name", APPLICANT["full_name"]),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
        ("Telephone", APPLICANT["phone"]),
        ("Email", APPLICANT["email"] or "(to be provided)"),
    ])

    doc.add_page_break()

    # ── Section 3: PTO/SB/16 ──
    add_heading(doc, "3. Provisional Application Cover Sheet (PTO/SB/16)", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "Request for filing a PROVISIONAL APPLICATION FOR PATENT "
        "under 37 CFR 1.53(c)."
    )
    run.italic = True

    doc.add_paragraph("")

    add_heading(doc, "Title of Invention", level=2)
    p = doc.add_paragraph(APPLICATION["title"])
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)

    doc.add_paragraph("")

    add_heading(doc, "Inventor", level=2)
    add_field_table(doc, [
        ("Given Name", "Issac Daniel"),
        ("Family Name", "Davis"),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
    ])

    add_heading(doc, "Correspondence Address", level=2)
    add_field_table(doc, [
        ("Name", APPLICANT["full_name"]),
        ("Street Address", APPLICANT["address"]),
        ("City", APPLICANT["city"]),
        ("State", APPLICANT["state"]),
        ("ZIP Code", APPLICANT["zip"]),
        ("Country", APPLICANT["country"]),
        ("Telephone", APPLICANT["phone"]),
        ("Email", APPLICANT["email"] or "(to be provided)"),
    ])

    add_heading(doc, "Entity Status", level=2)
    add_checkbox_line(doc, "Micro Entity (37 CFR 1.29)", checked=True)
    add_checkbox_line(doc, "Small Entity (37 CFR 1.27)", checked=False)
    add_checkbox_line(doc, "Regular Undiscounted", checked=False)

    doc.add_paragraph("")
    add_heading(doc, "Application Details", level=2)
    add_field_table(doc, [
        ("Application Number", APPLICATION["number"]),
        ("Original Filing Date", APPLICATION["filing_date"]),
        ("Number of Claims", str(APPLICATION["claims_count"])),
        ("Docket Number", APPLICATION["docket"]),
    ])

    add_heading(doc, "Signature", level=2)
    add_field_table(doc, [
        ("Signature", f"/{APPLICANT['full_name']}/"),
        ("Name (Printed)", APPLICANT["full_name"]),
        ("Date", TODAY),
        ("Registration Number", "N/A — Pro Se Applicant"),
    ])

    doc.add_page_break()

    # ── Section 4: Filing Checklist ──
    add_heading(doc, "4. Filing Checklist", level=1)

    doc.add_paragraph(
        "Use this checklist when filing at patentcenter.uspto.gov:"
    )
    doc.add_paragraph("")

    steps = [
        "Log in to Patent Center (patentcenter.uspto.gov)",
        "Navigate to application 63/961,403",
        "Select 'File a Document' or 'Follow-on Submissions'",
        "Upload Cover Letter (Section 1 of this document)",
        "Upload PTO/SB/15A — Micro Entity Certification (Section 2)",
        "Upload PTO/SB/16 — Cover Sheet (Section 3)",
        "Select entity status: Micro Entity",
        "Verify fee amount shows $65.00 base micro-entity filing fee",
        "Enter payment information (credit/debit card)",
        "Submit payment",
        "Save confirmation receipt / screenshot",
        "Verify submission shows 'Received' the next business day",
    ]

    for i, step in enumerate(steps, 1):
        add_checkbox_line(doc, f"Step {i}: {step}", checked=False)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("USPTO Customer Service: 1-800-786-9199")
    run.bold = True
    doc.add_paragraph("")

    # Reference codes
    add_heading(doc, "Reference Codes", level=2)
    add_field_table(doc, [
        ("Form PTO/SB/15A", "Certification of Micro Entity Status (37 CFR 1.29)"),
        ("Form PTO/SB/16", "Provisional Application Cover Sheet (37 CFR 1.53(c))"),
        ("37 CFR 1.29", "Micro Entity definition and requirements"),
        ("37 CFR 1.29(a)", "Gross Income Basis for Micro Entity"),
        ("37 CFR 1.53(c)", "Provisional application filing requirements"),
        ("35 U.S.C. 111(b)", "Provisional patent application statute"),
        ("35 U.S.C. 119(e)", "Priority claim from provisional"),
        ("MPEP 509.04", "Micro Entity status guidance"),
        ("Fee Code 1005", "Provisional filing fee — Micro Entity"),
    ])

    path = os.path.join(OUTPUT_DIR, "Combined_Filing_Package.docx")
    doc.save(path)
    print(f"  Created: {path}")
    return doc


# ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating USPTO filing documents...\n")
    build_cover_letter()
    build_pto_sb_15a()
    build_pto_sb_16()
    build_combined()
    print(f"\nAll documents saved to: {OUTPUT_DIR}")
    print("\nNEXT STEPS:")
    print("  1. Add your email to the script (line 20) and re-run, OR edit the .docx files directly")
    print("  2. Review each document")
    print("  3. Save/export as PDF for USPTO upload")
    print("  4. Go to patentcenter.uspto.gov when ready")
