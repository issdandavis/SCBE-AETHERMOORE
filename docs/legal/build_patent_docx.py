#!/usr/bin/env python3
"""
Build USPTO-format non-provisional patent specification DOCX.

Output: docs/legal/SCBE_NONPROVISIONAL_SPEC_v1.docx

USPTO formatting:
  Font          Times New Roman 12pt
  Line spacing  At least 1.5 lines
  Margins       1.5" left, 1" top/right/bottom
  Page numbers  Bottom center
  Sections      Title → Cross-Ref → Background → Summary →
                Brief Description → Detailed Description → Claims → Abstract
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
HERE = Path(__file__).parent
BASE = HERE.parent.parent          # SCBE-AETHERMOORE root
SPEC_MD = BASE / "docs" / "PATENT_DETAILED_DESCRIPTION.md"
OUTPUT = HERE / "SCBE_NONPROVISIONAL_SPEC_v1.docx"

# ---------------------------------------------------------------------------
# Static content
# ---------------------------------------------------------------------------
TITLE = (
    "System and Method for Hyperbolic Geometry-Based Authorization "
    "with Topological Control-Flow Integrity"
)

CROSS_REFERENCE = (
    'This application claims priority to U.S. Provisional Application '
    'No. 63/961,403, filed January 15, 2026, entitled '
    '"System and Method for Hyperbolic Geometry-Based Authorization with '
    'Topological Control-Flow Integrity," the entire disclosure of which is '
    'incorporated herein by reference.'
)

SUMMARY = (
    "The present invention provides a computer-implemented system and method "
    "for governing execution of computational actions in artificial-intelligence "
    "and agentic systems. The system embeds input actions into a Poincare ball "
    "model of bounded hyperbolic space, measures hyperbolic distance from a "
    "session-adaptive centroid representing prior authorized behavior, and applies "
    "a nonlinear governance cost function that increases governance severity "
    "as measured deviation increases. A four-tier decision gate routes proposed "
    "actions to allow, review, quarantine, or deny outcomes.\n\n"

    "In one aspect, the invention provides a method comprising: receiving a "
    "request associated with a computational action; generating a context "
    "representation comprising semantic, operational, or temporal features; "
    "transforming the context representation into an embedded point in a bounded "
    "hyperbolic space comprising a Poincare ball model via a tanh-normalized "
    "projection with epsilon clamping; maintaining a session centroid updated as "
    "a function of prior embedded points within the session; computing a "
    "hyperbolic distance between the embedded point and the session centroid; "
    "computing a governance cost using a nonlinear cost function that increases "
    "as the hyperbolic distance increases; combining the governance cost with "
    "additional governance signals to produce a composite risk value; and emitting "
    "a governance decision -- allow, review, quarantine, or deny -- that controls "
    "whether the computational action is executed.\n\n"

    "In another aspect, the invention provides a system comprising at least one "
    "processor and a non-transitory memory storing persistent runtime state and "
    "instructions that, when executed, perform the governance method, wherein the "
    "persistent runtime state is restored after a process restart.\n\n"

    "In a further aspect, the invention provides a non-transitory computer-"
    "readable medium storing instructions that perform bijective tamper detection "
    "by generating a re-encoded form of an input, computing canonical abstract "
    "syntax tree (AST) representations of the input and re-encoded form, computing "
    "a tamper signal based on divergence between the AST representations, and "
    "providing the tamper signal to a governance gate that escalates or blocks "
    "proposed computational actions.\n\n"

    "Preferred embodiments include: session-adaptive centroid accumulation; a "
    "multi-layer pre-filter stack ordered cheapest-reject-first; a fail-to-noise "
    "response returning deterministic pseudorandom-looking audit output upon "
    "denial; durable state persistence and rollback across "
    "process restarts; quarantine non-error containment that restricts tool access "
    "and execution resources without terminating the session; post-quantum decision "
    "receipts signed under ML-DSA-65 and encapsulated under ML-KEM-768; six-axis "
    "semantic weighting using golden-ratio-scaled tongue weights; and physics-based "
    "task coordination for multi-agent governance."
)

FIGURE_DESCRIPTIONS = [
    (
        "FIG. 1",
        "Block diagram of the fourteen-layer processing pipeline showing data "
        "flow from Complex Context State (Layer 1) through Risk Decision Gate "
        "(Layer 13) and Audio Axis (Layer 14), with layer dependencies and "
        "functional group boundaries for Groups I-IV indicated.",
    ),
    (
        "FIG. 2",
        "Alternative harmonic cost functions plotted against measured drift, "
        "including H(d,R) = R^(d^2), a bounded reciprocal score "
        "1/(1+d+2*pd), and a clamped RuntimeGate cost "
        "pi^(phi*min(d*,d_max)).",
    ),
    (
        "FIG. 3",
        "Poincare ball cross-section showing the open unit ball B^n with the "
        "session centroid as the trusted reference region at the ball interior, "
        "an embedded point representing a proposed action, and the hyperbolic "
        "distance arc d_H = arccosh(1 + 2||u-v||^2 / ((1-||u||^2)(1-||v||^2))) "
        "between them.",
    ),
    (
        "FIG. 4",
        "Six Sacred Tongues semantic-weighting diagram showing golden-ratio "
        "weights phi^k for k = 0 through 5 as a bar chart and spoke diagram, "
        "with tongue names Kor'aelin (KO, weight 1.00), Avali (AV, 1.62), "
        "Runethic (RU, 2.62), Cassisivadan (CA, 4.24), Umbroth (UM, 6.85), "
        "and Draumric (DR, 11.09).",
    ),
    (
        "FIG. 5",
        "Sacred Egg five-predicate deferred authorization flowchart showing the "
        "conjunction P_tongue AND P_geo AND P_path AND P_quorum AND P_crypto, "
        "wherein failure of any predicate branches to a fail-to-noise or "
        "pseudorandom-looking audit output.",
    ),
    (
        "FIG. 6",
        "Multi-layer pre-filter stack showing cheapest-reject-first ordering: "
        "(i) script-origin gate, (ii) instruction-safety regex gate, "
        "(iii) semantic pattern filter, (iv) small-language-model router, and "
        "(v) hyperbolic-distance computation, wherein rejection at any stage "
        "prevents invocation of subsequent costlier stages.",
    ),
    (
        "FIG. 7",
        "Runtime decision gate diagram illustrating how the composite risk value "
        "and auxiliary signals enter a threshold decision tree to route a proposed "
        "action to one of four governance decisions: ALLOW (risk below allow "
        "threshold), REVIEW (between allow and quarantine thresholds), QUARANTINE "
        "(non-error containment applied), or DENY (fail-to-noise response "
        "returned).",
    ),
    (
        "FIG. 8",
        "Bijective tamper detection flow diagram showing input source code "
        "processed by: (a) encode operation to token sequence, (b) decode "
        "operation to decoded source, (c) canonical AST parse of both, "
        "(d) SHA-256 fingerprint of each canonical AST dump, (e) comparison "
        "producing a tamper signal classified as none, NFC-normalization, "
        "structural-divergence, or syntax-failure, and (f) routing of the "
        "tamper signal to the governance gate.",
    ),
    (
        "FIG. 9",
        "System deployment architecture showing three deployment surfaces -- "
        "a REST API endpoint, an agent-bus service (scbe-agent-bus package), "
        "and a command-line interface -- all sharing a common governance runtime "
        "that maintains durable session state and interfaces with the 14-layer "
        "pipeline, the pre-filter stack, and the post-quantum receipt module.",
    ),
]

ABSTRACT = (
    "A computer-implemented system and method govern execution of computational "
    "actions in artificial-intelligence and agentic systems. Input actions are "
    "encoded as six-dimensional context vectors weighted by successive powers of "
    "the golden ratio and embedded into a Poincare ball model of hyperbolic space. "
    "Hyperbolic distance between the embedded point and a session centroid "
    "representing prior authorized behavior drives a nonlinear authorization "
    "cost or safety score used to control execution. "
    "Governance decisions -- allow, review, quarantine, or deny -- combine the "
    "authorization cost with bijective tamper detection, instruction-safety "
    "pattern matching, and spectral coherence signals. Denied actions return "
    "deterministic pseudorandom-looking audit noise instead of a structured "
    "failure category. "
    "Authorized actions receive post-quantum cryptographic receipts signed under "
    "ML-DSA-65 and encapsulated under ML-KEM-768."
)

# ---------------------------------------------------------------------------
# 26 Claims
# ---------------------------------------------------------------------------
CLAIMS = [
    # --- Independent Claim 1 ---
    (
        1, True,
        "A computer-implemented method for governing execution of a computational "
        "action, comprising:\n"
        "    receiving, by one or more processors, a request associated with the "
        "computational action;\n"
        "    generating a context representation comprising one or more semantic, "
        "operational, or temporal features;\n"
        "    transforming the context representation into an embedded point in a "
        "bounded hyperbolic space comprising a Poincare ball model, via a "
        "tanh-normalized projection with epsilon clamping that constrains the "
        "embedded point to an open unit ball;\n"
        "    maintaining a session centroid as a trusted reference region, updated "
        "as a function of a plurality of embedded points corresponding to prior "
        "requests within a session;\n"
        "    computing a hyperbolic distance between the embedded point and the "
        "session centroid;\n"
        "    computing a governance cost from the hyperbolic distance using a "
        "nonlinear cost function that increases as the hyperbolic distance "
        "increases;\n"
        "    combining the governance cost with at least one additional governance "
        "signal from: semantic weighting, temporal drift, spectral coherence, spin "
        "coherence, identifier canonicality, or bijective tamper detection, to "
        "produce a composite risk value;\n"
        "    adjusting a severity of the composite risk value as a function of "
        "trajectory drift of the embedded points across the plurality of prior "
        "requests; and\n"
        "    emitting a governance decision, from: allow, review, quarantine, or "
        "deny, that controls whether the computational action is executed;\n"
        "    whereby the governance cost is a nonlinear increasing function of "
        "measured drift from the session centroid and is used to control "
        "execution of the computational action through the emitted governance "
        "decision."
    ),
    (
        2, False,
        "The method of claim 1, wherein the hyperbolic distance is computed as "
        "d_H = arccosh(1 + 2||u - v||^2 / ((1 - ||u||^2)(1 - ||v||^2))), "
        "where u is the embedded point and v is the session centroid."
    ),
    (
        3, False,
        "The method of claim 1, wherein the nonlinear cost function comprises "
        "one of: (i) a function of the form H(d, R) = R^(d^2), where d is a "
        "distance measure and R is a base greater than one; (ii) a bounded "
        "safety-score function of the form H = 1 / (1 + d + 2*pd), where pd is "
        "a phase-deviation term; or (iii) a function of the form pi^(phi*d), "
        "where pi is the mathematical constant pi and phi is the golden ratio."
    ),
    (
        4, False,
        "The method of claim 1, wherein the additional governance signal comprises "
        "a six-axis semantic weighting in which each of the six axes corresponds to "
        "a distinct semantic domain of the context representation and is assigned a "
        "predetermined weight equal to a power of the golden ratio, phi^k for a "
        "respective index k from 0 through 5, with phi approximately equal to 1.618, "
        "such that the weighting applies exponentially increasing emphasis to "
        "higher-order semantic domains and the combined six-axis score scales the "
        "governance signal across semantically diverse dimensions of the context."
    ),
    (
        5, False,
        "The method of claim 1, wherein maintaining the session centroid comprises "
        "an incremental update of the form centroid_new = (n * centroid_old + "
        "coord) / (n + 1), where n is a running count of prior embedded points "
        "and coord is the embedded point of the current request."
    ),
    (
        6, False,
        "The method of claim 1, further comprising maintaining a hash-indexed "
        "adversarial-memory set and a hash-indexed safe-memory set, and, prior "
        "to said transforming, returning a deny decision for a request whose "
        "content hash is a member of the adversarial-memory set, and returning "
        "an allow decision for a request whose content hash is a member of the "
        "safe-memory set, in each case without computing the hyperbolic distance."
    ),
    (
        7, False,
        "The method of claim 1, further comprising, responsive to a deny "
        "decision, generating a deterministic pseudorandom-looking noise output by "
        "computing a seed as a cryptographic hash of a fixed prefix concatenated "
        "with a content hash of the denied request, iteratively re-hashing the "
        "seed until a target length is reached, and returning the noise output "
        "in place of an error response, such that the noise output is identical "
        "for identical denied requests and is reproducible by an auditor from "
        "the content hash."
    ),
    (
        8, False,
        "The method of claim 1, further comprising periodically persisting, to "
        "a durable store, at least the session centroid, a cumulative governance "
        "cost, a query count, a trust history, and the adversarial-memory set of "
        "claim 6; and, after a process restart, restoring the persisted values "
        "so that the session continues from the restored trajectory rather than "
        "from a cold start."
    ),
    # --- Independent Claim 9 ---
    (
        9, True,
        "A system for runtime governance of agentic or artificial-intelligence "
        "actions, comprising:\n"
        "    at least one processor; and\n"
        "    a non-transitory memory storing a persistent runtime state and "
        "instructions that, when executed, cause the system to:\n"
        "        classify a proposed action into a context representation;\n"
        "        map the context representation into a point in a Poincare ball "
        "model of bounded hyperbolic space;\n"
        "        measure drift of the proposed action as a hyperbolic distance "
        "between the point and a session centroid, the session centroid being "
        "part of the persistent runtime state and comprising at least a centroid "
        "vector, a cumulative cost, and a query count;\n"
        "        calculate a harmonic governance cost from the measured drift "
        "using a nonlinear cost function that increases as the drift increases;\n"
        "        apply a decision gate to the harmonic governance cost and one or "
        "more auxiliary signals; and\n"
        "        route the proposed action according to allow, review, quarantine, "
        "or deny;\n"
        "    wherein the session centroid is updated from prior proposed actions "
        "within a session, and the persistent runtime state is restored after a "
        "process restart."
    ),
    (
        10, False,
        "The system of claim 9, wherein the quarantine route applies a non-error "
        "containment state that, without crashing or terminating the session, "
        "restricts available tools to an allowed subset, enforces an "
        "execution-time deadline, restricts execution permissions, and limits "
        "outbound network or filesystem effects of the proposed action."
    ),
    (
        11, False,
        "The system of claim 9, wherein the auxiliary signals are produced by a "
        "multi-layer pre-filter stack ordered cheapest-reject-first, comprising: "
        "(i) a script-origin gate computing a coverage score as a fraction of "
        "UTF-8 bytes falling within the printable ASCII range and rejecting when "
        "the coverage score is below a threshold; (ii) an instruction-safety gate "
        "matching compiled regular-expression patterns for instruction-override, "
        "persona-manipulation, or dangerous-tool-invocation text; (iii) a semantic "
        "pattern filter matching the proposed action against an adversarial-intent "
        "corpus; and (iv) a small-language-model router that rejects an input for "
        "which no applicable semantic band is classified; wherein each gate "
        "operates before the hyperbolic-distance computation, and rejection by "
        "any gate prevents invocation of subsequent gates."
    ),
    (
        12, False,
        "The system of claim 9, wherein, responsive to an allow decision, the "
        "system computes a content-addressed identifier of the proposed action as "
        "a cryptographic hash of a canonical representation, signs the identifier "
        "together with an authorization score and a timestamp using a post-quantum "
        "digital signature algorithm in accordance with FIPS 204 (ML-DSA-65), "
        "encapsulates a session key using a post-quantum key-encapsulation "
        "mechanism in accordance with FIPS 203 (ML-KEM-768), and returns a "
        "structured receipt comprising the decision, the score, signal "
        "identifiers, the timestamp, the signature, and a key-encapsulation "
        "ciphertext; and wherein a downstream executor verifies the signature "
        "before executing the allowed action."
    ),
    (
        13, False,
        "The system of claim 9, wherein the system emits an audit receipt "
        "comprising at least the decision, the harmonic governance cost, signal "
        "identifiers, and decision-relevant metadata including a session query "
        "count and a cumulative cost."
    ),
    (
        14, False,
        "The system of claim 9, wherein the system is deployed as at least one "
        "of: a REST API endpoint, an agent bus service, a command-line interface, "
        "or a programmatic client library."
    ),
    # --- Independent Claim 15 ---
    (
        15, True,
        "A non-transitory computer-readable medium storing instructions that, "
        "when executed, cause one or more processors to:\n"
        "    receive an input comprising source code or an "
        "identifier-containing input;\n"
        "    generate a re-encoded form of the input by applying a bijective "
        "encode operation followed by a decode operation mapping the input to a "
        "token sequence and back to a decoded input;\n"
        "    compute a first canonical abstract syntax tree (AST) representation "
        "of the input and a second canonical AST representation of the decoded "
        "input, each comprising a content-derived fingerprint;\n"
        "    compute a tamper signal based on at least one of: (i) a divergence "
        "between the first and second canonical AST representations; (ii) a "
        "failure of the decoded input to parse into a valid AST; (iii) a Unicode "
        "canonicality failure; or (iv) a confusable-identifier condition; and\n"
        "    provide the tamper signal to a governance gate that escalates or "
        "blocks a proposed computational action when the tamper signal exceeds a "
        "threshold;\n"
        "    wherein the tamper signal is distinct from a tokenizer "
        "reconstruction-quality measure in that it is derived from a comparison "
        "of abstract syntax tree representations and gates execution of a "
        "proposed action."
    ),
    (
        16, False,
        "The medium of claim 15, wherein each content-derived fingerprint "
        "comprises a SHA-256 digest of the corresponding canonical AST, such "
        "that two inputs that are semantically equal but differ in whitespace "
        "or comments produce identical fingerprints."
    ),
    (
        17, False,
        "The medium of claim 15, wherein the confusable-identifier condition is "
        "detected by parsing the input into an AST, extracting identifier names, "
        "and, for each identifier, determining whether any non-ASCII codepoint "
        "maps to an ASCII-confusable codepoint, whether the identifier mixes two "
        "or more distinct writing scripts, or whether the identifier contains an "
        "invisible or bidirectional control codepoint; and computing a "
        "confusable-identifier score as a function of the fraction of suspicious "
        "identifiers."
    ),
    (
        18, False,
        "The medium of claim 15, wherein the governance gate denies the proposed "
        "computational action on a syntax-divergence class in which the decoded "
        "input fails to parse, quarantines the proposed computational action on "
        "a structural-divergence class in which the canonical ASTs diverge while "
        "both parse, and allows with annotation on a normalization-divergence "
        "class attributable to Unicode NFC normalization."
    ),
    (
        19, False,
        "The medium of claim 15, wherein, when a tokenizer artifact is absent, "
        "the instructions substitute a normalization stub that performs the "
        "encode operation by applying Unicode NFC normalization and encoding as "
        "UTF-8, the stub preserving the bijective round-trip property for ASCII "
        "inputs."
    ),
    (
        20, False,
        "The medium of claim 15, wherein the tamper signal, a divergence "
        "classification, and the content-derived fingerprint are recorded in an "
        "audit trail."
    ),
    (
        21, False,
        "The method of claim 1, further comprising generating a cryptographic "
        "authorization container that is unlocked only when N predetermined "
        "predicates are satisfied, where N is at least three, the predicates "
        "including at least: a semantic predicate evaluating whether the context "
        "representation of the proposed action satisfies an authorized semantic "
        "profile; a geometric predicate measuring whether the embedded point lies "
        "within a predetermined hyperbolic distance from the session centroid; "
        "and a cryptographic predicate verifying a post-quantum signature; "
        "wherein failure of any predicate causes the container to return a noise "
        "or pseudorandom-looking output generated by the fail-to-noise function "
        "of claim 7 rather than a structured predicate-failure response."
    ),
    (
        22, False,
        "The method of claim 21, wherein the noise output is generated by the "
        "deterministic re-hashing of claim 7, such that a repeated failure path "
        "for the same denied request produces an audit-reproducible output of "
        "a predetermined length while avoiding disclosure of which predicate "
        "failed."
    ),
    (
        23, False,
        "The method of claim 1, further comprising: prior to emitting the "
        "governance decision, determining whether the computational action matches "
        "a predetermined reroute rule associated with a class of actions; and, "
        "when a match is found, substituting a replacement action for the proposed "
        "computational action and emitting an allow decision for the replacement "
        "action, such that high-risk classes of actions are redirected to "
        "lower-risk alternatives without exposing a denial response to the "
        "requesting entity."
    ),
    (
        24, False,
        "The method of claim 1, further comprising computing a null-space "
        "anomaly score by determining whether per-axis deviations of the context "
        "representation from the session centroid each fall below a predetermined "
        "threshold; incrementing the null-space anomaly score when all per-axis "
        "deviations are below the threshold; and incorporating the null-space "
        "anomaly score into the composite risk value; wherein a null-space "
        "anomaly score above a predetermined level is treated as a governance "
        "signal indicating an action that is deliberately mimicking baseline "
        "behavior to evade the governance cost."
    ),
    (
        25, False,
        "The system of claim 9, wherein the system coordinates task execution "
        "across a plurality of agent slots using a physics-based juggling model "
        "in which tasks are modeled as balls having inertia proportional to a "
        "task priority, agent slots are modeled as hands having readiness states, "
        "handoffs are modeled as throws having predicted catch windows, and a "
        "governance cost of a task increases when a trajectory of the task "
        "deviates from a predicted flight arc, such that higher-risk tasks are "
        "assigned higher arcs and fewer handoffs."
    ),
    (
        26, False,
        "The method of claim 4, wherein each axis of the semantic weighting "
        "employs a bijective token alphabet comprising a number of tokens equal "
        "to the Cartesian product of a first predetermined prefix set and a "
        "second predetermined suffix set, each token uniquely formed by "
        "concatenating a prefix element, a predetermined separator character, "
        "and a suffix element, such that the complete token vocabulary for each "
        "axis bijects onto a contiguous range of integer byte indices, and wherein "
        "a serialized token form includes an axis designator "
        "that makes serialized token vocabularies of distinct axes pairwise "
        "disjoint and makes the axis of origin of any serialized token determinable "
        "from the serialized token without additional context."
    ),
    (
        27, False,
        "The method of claim 4, wherein each of the six axes of the semantic "
        "weighting is associated with a respective harmonic frequency ratio "
        "selected from integer-ratio musical intervals, and a phase offset equal "
        "to 2*pi*k/6 radians for the respective axis index k from 0 through 5, "
        "such that the six axes are uniformly distributed around the unit circle "
        "at sixty-degree intervals, and the governance signal incorporates the "
        "phase offset of each axis as a fixed structural parameter of the "
        "six-axis weighting configuration."
    ),
    (
        28, False,
        "The method of claim 26, wherein each axis's token vocabulary constitutes "
        "a domain-specific deterministic byte encoding in which each byte value "
        "of a context representation maps uniquely to a token in that axis's "
        "vocabulary, the semantic content of the context thereby constraining the "
        "specific tokens produced within each axis's vocabulary; and wherein the "
        "pairwise-disjoint serialized-vocabulary property ensures that the same "
        "byte sequence encoded under distinct axes produces structurally distinct "
        "token sequences, such that axis-specific encodings of the same input are "
        "mutually distinguishable without decoding."
    ),
]


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1.5)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    # Page numbers (bottom center)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fp.runs
    run = fp.add_run()
    run.element.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run.element.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run.element.append(fld2)
    _fmt_run(run, size=10)


def _fmt_run(run, bold=False, size=12) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def _set_spacing(para, before=0, after=6, line_spacing=1.5) -> None:
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(int(12 * line_spacing))


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p, before=18, after=6)
    r = p.add_run(text.upper())
    _fmt_run(r, bold=True, size=12)


def sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_spacing(p, before=10, after=4)
    r = p.add_run(text)
    _fmt_run(r, bold=True, size=12)


def body(doc: Document, text: str) -> None:
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=6)
        p.paragraph_format.first_line_indent = Inches(0.5)
        r = p.add_run(block)
        _fmt_run(r)


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Markdown → plain text
# ---------------------------------------------------------------------------

def strip_md(text: str) -> str:
    text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`{3}[^\n]*\n.*?\n`{3}", "", text, flags=re.DOTALL)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    text = text.replace("&nbsp;", " ")
    return text.strip()


def trim_detailed_description(md_text: str) -> str:
    """Remove legacy drafting appendices that are assembled separately."""
    stop_patterns = [
        r"^##\s+FILING CLAIMS\s*$",
        r"^##\s+CLAIM SUPPORT MAP\s*$",
        r"^##\s+PRIOR ART CITATIONS\s*$",
        r"^###\s+Independent Claims\s*$",
        r"^###\s+Dependent Claims\s*$",
        r"^##\s+Claims\s*$",
        r"^#\s+Claims\s*$",
        r"^CLAIMS\s*$",
        r"^##\s+Abstract\s*$",
        r"^#\s+Abstract\s*$",
        r"^ABSTRACT\s*$",
    ]
    for pattern in stop_patterns:
        match = re.search(pattern, md_text, flags=re.IGNORECASE | re.MULTILINE)
        if match:
            return md_text[:match.start()].rstrip()
    return md_text


def extract_section(md_text: str, start_header: str, stop_headers: list) -> str:
    """
    Return text between `start_header` and the first of `stop_headers`.
    Matches only H2+ headers (## or ###) so H1 title lines don't steal the match.
    """
    lines = md_text.splitlines()
    capturing = False
    out = []
    for line in lines:
        stripped = line.strip()
        low = stripped.lower()
        # Only match H2+ (## ...), not H1 (#) title lines
        is_h2plus = stripped.startswith("## ") or stripped.startswith("### ")
        if start_header.lower() in low and is_h2plus:
            capturing = True
            continue
        if capturing:
            is_any_h2plus = stripped.startswith("## ") or stripped.startswith("### ")
            if any(h.lower() in low and is_any_h2plus for h in stop_headers):
                break
            out.append(line)
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Claim formatter
# ---------------------------------------------------------------------------

def add_claim(doc: Document, num: int, independent: bool, text: str) -> None:
    # Each non-empty line becomes its own paragraph.
    # First line: claim number + opening text at left margin.
    # Remaining lines (claim elements/limitations): indented 0.25".
    # No bold on any claim text — USPTO standard.
    lines = [ln for ln in text.split("\n") if ln.strip()]

    for i, line in enumerate(lines):
        is_first = i == 0
        is_last = i == len(lines) - 1
        p = doc.add_paragraph()
        _set_spacing(p, before=(6 if is_first else 0), after=(6 if is_last else 0))
        p.paragraph_format.left_indent = Inches(0.0 if is_first else 0.25)
        p.paragraph_format.first_line_indent = Pt(0)
        content = f"{num}. {line.strip()}" if is_first else line.strip()
        r = p.add_run(content)
        _fmt_run(r)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build() -> None:
    print(f"Reading specification source: {SPEC_MD}")
    md_text = SPEC_MD.read_text(encoding="utf-8")

    doc = Document()
    setup_document(doc)

    # ---- 1. TITLE --------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=24, after=12)
    r = p.add_run(TITLE)
    _fmt_run(r, bold=True, size=14)

    # ---- 2. CROSS-REFERENCE ----------------------------------------------
    section_heading(doc, "CROSS-REFERENCE TO RELATED APPLICATIONS")
    body(doc, CROSS_REFERENCE)

    # ---- 3. BACKGROUND ---------------------------------------------------
    section_heading(doc, "BACKGROUND OF THE INVENTION")

    bg_raw = extract_section(
        md_text,
        "BACKGROUND OF THE INVENTION",
        ["BRIEF DESCRIPTION", "DETAILED DESCRIPTION", "SUMMARY", "CLAIMS", "ABSTRACT"],
    )
    bg_clean = strip_md(bg_raw)
    # Split into sub-headings if present
    for chunk in re.split(r"\n{2,}", bg_clean):
        chunk = chunk.strip()
        if not chunk:
            continue
        # Lines starting like "Field of the Invention" or "Description of Related Art"
        if len(chunk) < 60 and not chunk.endswith("."):
            sub_heading(doc, chunk)
        else:
            body(doc, chunk)

    # ---- 4. SUMMARY ------------------------------------------------------
    section_heading(doc, "SUMMARY OF THE INVENTION")
    for block in SUMMARY.split("\n\n"):
        body(doc, block.strip())

    # ---- 5. BRIEF DESCRIPTION OF DRAWINGS --------------------------------
    section_heading(doc, "BRIEF DESCRIPTION OF THE DRAWINGS")
    for fig_label, fig_desc in FIGURE_DESCRIPTIONS:
        p = doc.add_paragraph()
        _set_spacing(p, before=3, after=3)
        p.paragraph_format.left_indent = Inches(0.0)
        r1 = p.add_run(f"{fig_label}—")
        _fmt_run(r1, bold=True)
        r2 = p.add_run(fig_desc)
        _fmt_run(r2)

    # ---- 6. DETAILED DESCRIPTION -----------------------------------------
    section_heading(doc, "DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS")

    dd_raw = extract_section(
        md_text,
        "DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS",
        # Stop at nothing — take to end (no further patent section in this file)
        ["## CLAIMS", "## ABSTRACT"],
    )
    dd_raw = trim_detailed_description(dd_raw)
    dd_clean = strip_md(dd_raw)

    # Parse into section/sub-section headings vs body paragraphs.
    # Original file uses ### SECTION 1:, #### 3.1, etc.
    # After strip_md those become bare text lines starting capital or digits.
    in_code_block = False
    current_block: list[str] = []

    def flush_block():
        if current_block:
            joined = "\n".join(current_block).strip()
            if joined:
                body(doc, joined)
        current_block.clear()

    for line in dd_clean.splitlines():
        stripped = line.strip()
        if not stripped:
            flush_block()
            continue
        # Detect what were originally section headings (short lines, no period)
        if (len(stripped) < 80 and not stripped.endswith(".")
                and not stripped.endswith(",")
                and not stripped.endswith(";")
                and re.match(r"^(SECTION|Section|[0-9]+\.|[0-9]+\.[0-9]+)", stripped)):
            flush_block()
            sub_heading(doc, stripped)
        elif (len(stripped) < 80 and not stripped.endswith(".")
              and not stripped.endswith(",")
              and re.match(r"^Layer [0-9]+:", stripped)):
            flush_block()
            sub_heading(doc, stripped)
        else:
            current_block.append(stripped)
    flush_block()

    # ---- 7. CLAIMS -------------------------------------------------------
    page_break(doc)
    section_heading(doc, "CLAIMS")
    doc.add_paragraph()

    for num, independent, text in CLAIMS:
        add_claim(doc, num, independent, text)

    # ---- 8. ABSTRACT (last page) -----------------------------------------
    page_break(doc)
    section_heading(doc, "ABSTRACT")
    body(doc, ABSTRACT)

    # ---- Save ------------------------------------------------------------
    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    print(f"  Claims: {len(CLAIMS)} total, "
          f"{sum(1 for _, ind, _ in CLAIMS if ind)} independent")
    print(f"  Sections: Title, Cross-Ref, Background, Summary, "
          "Brief Description (FIG 1-9), Detailed Description, Claims, Abstract")


if __name__ == "__main__":
    build()
