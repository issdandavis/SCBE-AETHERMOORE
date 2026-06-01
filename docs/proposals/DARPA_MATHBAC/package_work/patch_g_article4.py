"""Patch G Article 4 performer POC fields + Performer definition in Art. 1B."""
import docx

G = (
    "docs/proposals/DARPA_MATHBAC/package_work/07_attachment_g_model_ot/"
    "Attachment_G_SCBE_FILLED.docx"
)
doc = docx.Document(G)

BLUE = "00B0F0"
PHONE = "(360) 808-0876"
EMAIL = "issdandavis7795@gmail.com"
NAME = "Issac Davis"
TITLE = "Principal Investigator"
PERFORMER_NAME = "ISSAC D DAVIS / SCBE-AETHERMOORE"


def is_blue(run):
    try:
        return run.font.color.rgb and str(run.font.color.rgb) == BLUE
    except Exception:
        return False


changes = 0

for idx, para in enumerate(doc.paragraphs):
    runs = para.runs
    txt = para.text or ""

    # ── P126: Performer definition "PERFORMER " + "NAME" (blue, no INSERT wrapper) ──
    # Pattern: two consecutive blue runs whose concatenation is "PERFORMER NAME"
    for i in range(len(runs) - 1):
        if is_blue(runs[i]) and is_blue(runs[i + 1]):
            combined = (runs[i].text or "") + (runs[i + 1].text or "")
            if combined == "PERFORMER NAME":
                runs[i].text = PERFORMER_NAME
                runs[i + 1].text = ""
                changes += 1
                print(f"  FIXED P{idx:03d}: Performer definition → {PERFORMER_NAME!r}")
                break

    # ── Article 4 POC fields (blue single-run placeholders) ──
    # Identify context: only 4 target paragraphs, identified by paragraph index
    # P226/P232: (NAME) → Issac Davis
    if txt == "(NAME)" and all(is_blue(r) for r in runs if r.text.strip()):
        for r in runs:
            if r.text.strip() == "(NAME)":
                r.text = NAME
                changes += 1
                print(f"  FIXED P{idx:03d}: (NAME) → {NAME!r}")
                break

    # P227/P233: (TITLE) → Principal Investigator
    if txt.strip().startswith("(TITLE)") and all(is_blue(r) for r in runs if r.text.strip()):
        for r in runs:
            if "(TITLE)" in (r.text or ""):
                r.text = r.text.replace("(TITLE)", TITLE)
                changes += 1
                print(f"  FIXED P{idx:03d}: (TITLE) → {TITLE!r}")
                break

    # P228/P234: "Phone Number:" (blue) → "Phone Number: (360) 808-0876"
    if "Phone Number:" in txt and is_blue(runs[0]) if runs else False:
        for r in runs:
            if "Phone Number:" in (r.text or "") and is_blue(r):
                r.text = f"Phone Number: {PHONE}"
                changes += 1
                print(f"  FIXED P{idx:03d}: Phone → {PHONE!r}")
                break

    # P229/P235: "Email:" (blue) → "Email: issdandavis7795@gmail.com"
    if "Email:" in txt and any(is_blue(r) for r in runs):
        for r in runs:
            if r.text and r.text.strip() == "Email:" and is_blue(r):
                r.text = f"Email: {EMAIL}"
                changes += 1
                print(f"  FIXED P{idx:03d}: Email → {EMAIL!r}")
                break

doc.save(G)
print(f"\nPatch complete: {changes} changes saved.")

# Verify
doc2 = docx.Document(G)
print("\n=== Verification ===")
for idx, para in enumerate(doc2.paragraphs):
    t = para.text or ""
    for needle in (PHONE, EMAIL, "Issac Davis", "Principal Investigator", "ISSAC D DAVIS"):
        if needle in t:
            print(f"  OK P{idx:03d}: {t[:100]!r}")
            break
