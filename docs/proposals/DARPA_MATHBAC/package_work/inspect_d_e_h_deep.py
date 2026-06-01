"""Deep inspect D, E, H templates — print all table cells and paragraph text."""
import docx

TEMPLATES = {
    "D": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_D_Proposal_Instructions_and_Volume_I_Template__Technical_and_Management.docx"
    ),
    "E": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_E_Proposal_Instructions_and_Volume_II_Template_Price.docx"
    ),
    "H": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_H_Task_Description_Document_TDD_Template.docx"
    ),
}


def dump_table(tbl, tbl_idx):
    print(f"    Table {tbl_idx}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
    for r_idx, row in enumerate(tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"      [{r_idx},{c_idx}]: {txt[:100]!r}")


for name, path in TEMPLATES.items():
    print(f"\n{'='*70}")
    print(f"  Attachment {name}")
    print(f"{'='*70}")
    doc = docx.Document(path)

    print("\n  --- ALL PARAGRAPHS ---")
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        style = para.style.name if para.style else "?"
        if txt:
            print(f"  P{idx:03d} [{style[:20]}]: {txt[:120]!r}")

    print("\n  --- ALL TABLES ---")
    for t_idx, tbl in enumerate(doc.tables):
        dump_table(tbl, t_idx)
