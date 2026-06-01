"""
Build NSF Common Form Current and Pending (Other) Support DOCX for Issac Davis.

NSF Common Form — November 1, 2023 version.

Projects disclosed:
  1. MATHBAC full proposal (this proposal) — Pending
  2. CLARA full proposal — Pending

In-kind contributions: None >= $5,000 requiring time commitment.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/11_disclosure_forms/"
    "Current_and_Pending_Support_Issac_Davis_DRAFT.docx"
)


def add_heading(doc, text, size_pt=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_subheading(doc, text, size_pt=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_field(doc, label, value, size_pt=11):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(size_pt)
    r2 = p.add_run(value)
    r2.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text, size_pt=11, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size_pt)
    r.italic = italic
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_project(doc, project_num, fields):
    """
    fields: dict with keys matching NSF CPS form fields.
    """
    add_subheading(doc, f"Project/Proposal {project_num}")

    add_field(doc, "*Title:", fields["title"])
    add_field(doc, "*Status of Support:", fields["status"])
    add_field(doc, "Award Number (if available):", fields.get("award_number", "N/A"))
    add_field(doc, "*Source of Support:", fields["source"])
    add_field(doc, "*Primary Place of Performance:", fields["place"])
    add_field(doc, "*Start Date (MM/YYYY):", fields["start"])
    add_field(doc, "*End Date (MM/YYYY):", fields["end"])
    add_field(doc, "*Total Anticipated Amount:", fields["amount"])
    add_field(doc, "*Person-Months Per Year Devoted to Project:", fields["person_months"])
    add_field(doc, "*Overall Objectives:", fields["objectives"])
    add_field(doc, "*Statement of Potential Overlap:", fields["overlap"])
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CURRENT AND PENDING (OTHER) SUPPORT")
r.bold = True
r.font.size = Pt(14)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("NSF Common Form — November 1, 2023 | DARPA-PA-26-05 MATHBAC Submission")
r2.font.size = Pt(11)
r2.italic = True

doc.add_paragraph()

# ── Identifying Information ───────────────────────────────────────────────────

add_heading(doc, "Identifying Information")

add_field(doc, "*Name:", "Davis, Issac D.")
add_field(doc, "Persistent Identifier (ORCID/PID):", "N/A")
add_field(doc, "*Position Title:", "Principal Investigator / Independent AI Safety Researcher")

# ── Organization and Location ─────────────────────────────────────────────────

add_heading(doc, "Organization and Location")

add_field(doc, "*Organization Name:", "ISSAC D DAVIS / SCBE-AETHERMOORE")
add_field(doc, "*Location (City, State, Country):", "Port Angeles, WA, USA")

# ── Section a: Proposals and Active Projects ─────────────────────────────────

add_heading(doc, "a. Proposals and Active Projects")

add_body(doc,
    "Disclosure covers all proposals and active projects in accordance with NSPM-33 Current and "
    "Pending (Other) Support definitions. Two proposals are currently pending. No active projects "
    "with existing federal awards."
)

add_project(doc, "1", {
    "title": "SCBE Mathematical Framework for Agentic Communication Protocols",
    "status": "Pending (this proposal)",
    "award_number": "N/A — proposal not yet awarded",
    "source": "Defense Advanced Research Projects Agency (DARPA), Defense Sciences Office (DSO), "
               "MATHBAC Program. Solicitation: DARPA-PA-26-05.",
    "place": "Port Angeles, WA, USA",
    "start": "10/2026 (estimated; 16 months from anticipated award)",
    "end": "01/2028 (estimated)",
    "amount": "$839,000 (full proposal amount; inclusive of G&A and fee)",
    "person_months": "10 person-months per year (sole PI, full-time effort). "
                      "Year 1: 10 months (months 1–12). Year 2: 4 months (months 13–16).",
    "objectives": (
        "Develop a composed-operator governance framework T = L14 o ... o L1 for multi-agent AI "
        "communication using hyperbolic geometry (Poincare ball model). Map agent communication "
        "to Poincare ball coordinates; enforce five physical axioms (Unitarity, Locality, Causality, "
        "Symmetry, Composition) layer-by-layer. Demonstrate exponential adversarial cost scaling "
        "via harmonic wall H(d,pd) = 1/(1+phi*d_H+2*pd). Validate on NMR spectroscopy subdomains "
        "(Task A: 1H NMR to structure; Task B: 2D COSY assignment) with Karplus equation recovery "
        "as primary falsifiability experiment. Target: 99% adversarial classification, p95 latency "
        "<100ms for 240 cases."
    ),
    "overlap": (
        "This entry IS the MATHBAC proposal submission itself. A concurrent pending proposal "
        "(CLARA, DARPA-PA-25-07-02) uses the same SCBE governance framework in a different task "
        "domain (theorem-proving / course-of-action planning) with no scientific overlap to the "
        "NMR spectroscopy scope proposed here. In the event both proposals are awarded, the PI "
        "will disclose person-month commitments to both Program Managers and adjust effort "
        "allocations to eliminate any budgetary duplication. Person-month overlap would require "
        "renegotiation of at least one award's period of performance."
    ),
})

add_project(doc, "2", {
    "title": (
        "SCBE-AETHERMOORE: Verifiable Defeasible LP Composition "
        "for Trustworthy Multi-Agent AI Governance"
    ),
    "status": "Pending (submitted; award decision anticipated 06/2026)",
    "award_number": "DARPA-PA-25-07-02-CLARA-FP-033 (submission ID confirmed 2026-04-13)",
    "source": "Defense Advanced Research Projects Agency (DARPA), Defense Sciences Office (DSO), "
               "CLARA Program. Solicitation: DARPA-PA-25-07-02.",
    "place": "Port Angeles, WA, USA",
    "start": "06/2026 (if awarded; Phase 1 = 15 months, Phase 2 = 9 months)",
    "end": "06/2028 (full 24-month program if both phases awarded)",
    "amount": "$2,000,000 total ($1,350,000 Phase 1 / $650,000 Phase 2, approximate)",
    "person_months": "10 person-months per year (sole PI, full-time effort Phase 1). "
                      "Phase 2 estimated 9 person-months per year.",
    "objectives": (
        "Develop verifiable defeasible logic program (DLP) composition approaches for trustworthy "
        "multi-agent AI governance in complex systems engineering. SCBE's 14-layer pipeline provides "
        "the compositional AR+ML substrate. TA1 target: create composition approaches that achieve "
        "polynomial-time verifiability without performance loss, composed-task reliability exceeding "
        "state-of-the-art baselines, and sample complexity below SOA benchmarks (Phase 2). "
        "Application domain: cyber defense and/or COA planning (DARPA CLARA TA1 scope)."
    ),
    "overlap": (
        "CLARA and MATHBAC both leverage the SCBE governance framework but address distinct task "
        "domains: CLARA targets theorem-proving / COA planning (DARPA TA1 compositional AR+ML); "
        "MATHBAC targets NMR spectroscopy (DARPA TA1 multi-agent science subdomains). There is no "
        "scientific overlap in objectives, datasets, or validation experiments. The shared technical "
        "substrate (SCBE framework) constitutes Background IP disclosed in both proposals. If both "
        "awards are made simultaneously, person-month commitments will be disclosed and adjusted as "
        "required by both Program Managers to eliminate any budgetary duplication."
    ),
})

# ── Section b: In-Kind Contributions ─────────────────────────────────────────

add_heading(doc, "b. In-Kind Contributions")

add_body(doc,
    "None. No in-kind contributions with an estimated value of $5,000 or more that require a "
    "commitment of the PI's time are currently received or pending."
)

# ── Certification ─────────────────────────────────────────────────────────────

add_heading(doc, "Certification")

add_body(doc,
    "I certify that the information provided is current, accurate, and complete. This includes, "
    "but is not limited to, information related to current, pending, and other support (both "
    "foreign and domestic) as defined in 42 U.S.C. §6605.",
    italic=True
)
add_body(doc,
    "I also certify that, at the time of submission, I am not a party in a malign foreign talent "
    "recruitment program.",
    italic=True
)
add_body(doc,
    "Misrepresentations and/or omissions may be subject to prosecution and liability pursuant to, "
    "but not limited to, 18 U.S.C. §§287, 1001, 1031 and 31 U.S.C. §§3729-3733 and 3802.",
    italic=True
)

doc.add_paragraph()
add_field(doc, "Signature:", "[Sign before submission]")
add_field(doc, "Date:", "[Date — must be within 12 months of submission]")

doc.add_paragraph()
add_body(doc,
    "DRAFT — 2026-06-01. Prepared per NSF NSPM-33 Common Form (November 2023). "
    "OMB Control No. 3145-0279. Public reporting burden estimated 2 hours.",
    italic=True
)

doc.save(DEST)
print(f"Saved: {DEST}")
