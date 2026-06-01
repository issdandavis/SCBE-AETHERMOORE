"""
Build NSF Common Form Biographical Sketch (DOCX) for Issac Davis.

NSF Common Form format — 4 sections:
  1. Personal Statement
  2. Education and Training
  3. Products
  4. Synergistic Activities
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/11_disclosure_forms/"
    "Biographical_Sketch_Issac_Davis_DRAFT.docx"
)


def add_heading(doc, text, level=1, size_pt=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_subheading(doc, text, size_pt=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_body(doc, text, size_pt=11, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.italic = italic
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_labeled(doc, label, value, size_pt=11):
    p = doc.add_paragraph()
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(size_pt)
    r2 = p.add_run(value)
    r2.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table_row(tbl, cells):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        for para in cell.paragraphs:
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
    return row


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# ── Title / Header ────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("BIOGRAPHICAL SKETCH")
r.bold = True
r.font.size = Pt(14)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("NSF Common Form — Senior/Key Person")
r2.font.size = Pt(11)
r2.italic = True

doc.add_paragraph()  # spacer

# ── Identifying Information ───────────────────────────────────────────────────

add_heading(doc, "Identifying Information")

add_labeled(doc, "Name:", "Davis, Issac D.")
add_labeled(doc, "Persistent Identifier (PID/ORCID):", "N/A (not registered)")
add_labeled(doc, "Position Title:", "Principal Investigator / Independent AI Safety Researcher")

# ── Organization and Location ─────────────────────────────────────────────────

add_heading(doc, "Organization and Location")

add_labeled(doc, "Organization:", "ISSAC D DAVIS / SCBE-AETHERMOORE")
add_labeled(doc, "Location:", "Port Angeles, WA, USA")

# ── Section A: Personal Statement ─────────────────────────────────────────────

add_heading(doc, "A. Personal Statement")

add_body(doc,
    "I am an independent AI safety researcher and sole developer of the SCBE-AETHERMOORE governance "
    "framework — a 14-layer composed-operator pipeline that enforces AI safety through hyperbolic "
    "geometry. My research centers on the insight that adversarial behavior in multi-agent systems "
    "can be made computationally infeasible by placing it in the exponentially-expensive periphery "
    "of Poincaré ball space. Governance becomes a geometric property of the system, not a "
    "post-hoc classifier."
)
add_body(doc,
    "I am self-taught in the relevant mathematics — Lyapunov stability analysis, control-barrier "
    "functions (CBF), port-Hamiltonian operator structure, and the Poincaré ball metric — which I "
    "independently derived through primary texts (Khalil; Marsden & Ratiu; Cannon) before "
    "encountering the formal field terminology. This approach — deriving first, naming second — "
    "has produced a framework that maps directly to DARPA MATHBAC's TA1 (multi-agent communication "
    "dynamics) and to the CLARA program's compositional AR+ML requirement."
)
add_body(doc,
    "I am uniquely qualified to execute this proposal because I have already built and benchmarked "
    "the system. Pre-submission evidence includes: 13/13 terminal-bench task parity (zero governance "
    "overhead), 172/173 adversarial classifications correct on the Anthropic Petri benchmark (99.42%), "
    "and a p95 latency of 0.1095ms for 240 RuntimeGate evaluations. The mathematical framework, "
    "benchmarks, and open-source codebase all existed before this proposal was written — this is "
    "evidence of genuine readiness, not proposal-generated claims."
)

# ── Section B: Education and Training ─────────────────────────────────────────

add_heading(doc, "B. Education and Training")

add_body(doc, "* = required field. Entries follow Institution | Degree/Training | Field | Year format.")
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"

hdr = tbl.rows[0].cells
headers = ["Institution / Organization", "Degree / Training Type", "Field", "Year(s)"]
for i, h in enumerate(headers):
    for para in hdr[i].paragraphs:
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
    rn = hdr[i].paragraphs[0].add_run(h)
    rn.bold = True
    rn.font.size = Pt(10)

add_table_row(tbl, [
    "Self-directed study (primary texts: Khalil; Marsden & Ratiu; Cannon; Brattka et al.)",
    "Independent research training",
    "Nonlinear control theory, hyperbolic geometry, port-Hamiltonian systems",
    "2020–present"
])
add_table_row(tbl, [
    "Self-directed study (primary texts: Sutton & Barto; Goodfellow et al.; Devlin et al.)",
    "Independent research training",
    "Machine learning, neural architecture, transformer systems",
    "2019–present"
])
add_table_row(tbl, [
    "Self-directed study (primary texts: Baier & Katoen; Clarke et al.)",
    "Independent research training",
    "Formal methods, model checking, logic-based AI",
    "2021–present"
])
add_table_row(tbl, [
    "Port Angeles, WA (professional practice)",
    "Registered Entity (SAM.gov Active, UEI J4NXHM6N5F59, CAGE 1EXD5)",
    "AI safety research and governance systems development",
    "2023–present"
])

doc.add_paragraph()

# ── Section C: Products ───────────────────────────────────────────────────────

add_heading(doc, "C. Products")

add_subheading(doc, "Most Relevant Products (up to 5)")

add_body(doc,
    "1.  Davis, I. (2025). The Six Tongues Protocol. KDP (Amazon). ASIN: B0GSSFQD9G. "
    "Timestamped prior art for the Langues Weighting System (phi-metric Sacred Tongue tokenization), "
    "which forms the L3–L4 weighted transform in the SCBE 14-layer pipeline."
)
add_body(doc,
    "2.  Davis, I. (2024, filed). USPTO Provisional Patent Application 63/961,403. "
    "SCBE Composed-Operator Governance Framework — Poincaré ball safety geometry, "
    "14-layer axiom-enforced pipeline, harmonic wall scoring H(d,pd) = 1/(1+φ·d_H+2·pd)."
)
add_body(doc,
    "3.  Davis, I. (2024–present). SCBE-AETHERMOORE. Open-source repository (GitHub: issdandavis/SCBE-AETHERMOORE). "
    "62+ TypeScript modules, 40+ Python modules implementing the full 14-layer pipeline, HYDRA orchestration, "
    "Juggling Scheduler, Red/Blue Arena, and the SacredTongues tokenizer. Apache 2.0 license."
)
add_body(doc,
    "4.  Davis, I. (2024). issdandavis/phdm-21d-embedding. HuggingFace Hub. "
    "Polyhedral Hamiltonian Defense Manifold embedding model; 21-dimensional canonical state lift "
    "for AI governance decision consistency."
)
add_body(doc,
    "5.  Davis, I. (2025). issdandavis/spiralverse-ai-federated-v1. HuggingFace Hub. "
    "Federated multi-agent governance model trained on SCBE-governed trajectories."
)

add_subheading(doc, "Other Significant Products (up to 5)")

add_body(doc,
    "6.  Davis, I. (2026). \"Witnessed\" (novel). KDP. ASIN: B0H257QJC2. "
    "Published work demonstrating applied narrative AI ethics concepts; "
    "fictional exploration of agent accountability and moral persistence themes."
)
add_body(doc,
    "7.  Davis, I. (2026). \"The Miracle Was the Memory\" (novel). KDP. "
    "Published 2026-05-18. Companion work; explores memory and identity in AI-adjacent contexts."
)
add_body(doc,
    "8.  Davis, I. (2026). SCBE Governance Evidence Brief (pre-submission benchmark package). "
    "Internal technical report; commit f3fb4aa3c; includes 13/13 terminal-bench results, "
    "172/173 Petri adversarial classification results, and 240-case RuntimeGate latency data."
)

# ── Section D: Synergistic Activities ─────────────────────────────────────────

add_heading(doc, "D. Synergistic Activities")

add_body(doc,
    "1.  Open-source governance framework dissemination (2024–present). The SCBE-AETHERMOORE "
    "repository is publicly available under Apache 2.0, enabling external researchers and "
    "practitioners to examine, reproduce, and build on the 14-layer governance pipeline."
)
add_body(doc,
    "2.  HuggingFace model releases (2024–2025). Released two governance models (phdm-21d-embedding; "
    "spiralverse-ai-federated-v1) to the public HuggingFace Hub under user issdandavis, "
    "providing the AI safety community with trained governance-aware embedding artifacts."
)
add_body(doc,
    "3.  Prior art documentation and publication (2025). Published The Six Tongues Protocol "
    "(KDP, ASIN B0GSSFQD9G) as a timestamped technical record of the Sacred Tongues tokenization "
    "system, enabling independent verification of SCBE's pre-DARPA mathematical foundations."
)
add_body(doc,
    "4.  DARPA DSO non-traditional proposer participation (2026). Submitted full proposals to "
    "DARPA-PA-25-07-02 (CLARA) and DARPA-PA-26-05 (MATHBAC) as an independent sole proprietor, "
    "demonstrating that foundational AI governance research is feasible outside traditional "
    "academic and large-corporate settings."
)
add_body(doc,
    "5.  Collaboration with APEX Accelerator, Port Angeles, WA (2026–present). Receiving "
    "free government contracting guidance to build pathways for minority-owned small business "
    "access to federal research programs."
)

# ── Footer note ───────────────────────────────────────────────────────────────

doc.add_paragraph()
add_body(doc,
    "DRAFT — 2026-06-01. Prepared in accordance with NSF NSPM-33 Common Form (November 2023). "
    "Signature required for final submission. Certification: I certify that the information "
    "provided is current, accurate, and complete.",
    italic=True
)

doc.save(DEST)
print(f"Saved: {DEST}")
