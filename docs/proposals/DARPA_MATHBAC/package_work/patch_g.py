"""Patch remaining G issues after fill_g_and_i.py run."""
import docx

G_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/07_attachment_g_model_ot/"
    "Attachment_G_SCBE_FILLED.docx"
)
doc = docx.Document(G_DEST)

PERFORMER_NAME = "ISSAC D DAVIS / SCBE-AETHERMOORE"

changes = 0

for idx, para in enumerate(doc.paragraphs):
    runs = para.runs

    # ── P035: non-blue split "(INSERT PERFORMER NAME)" across runs 12-14 ──
    # Detect by run text pattern and fix in-place
    for i in range(len(runs) - 2):
        if (
            (runs[i].text or "").startswith("(INSERT ")
            and (runs[i + 1].text or "") == "PERFORMER"
            and (runs[i + 2].text or "").startswith(" NAME)")
        ):
            runs[i].text = PERFORMER_NAME
            runs[i + 1].text = ""
            # strip " NAME) " or " NAME)" from beginning of run[i+2]
            tail = runs[i + 2].text
            if tail.startswith(" NAME) "):
                runs[i + 2].text = tail[7:]
            elif tail.startswith(" NAME)"):
                runs[i + 2].text = tail[6:]
            changes += 1
            print(f"  FIXED P{idx:03d} non-blue performer name runs {i}-{i+2}")
            break

    # ── P708: split-bracket checkbox — mark "is not [X]" ──
    # Structure: run4="is not " run5="[ " run6="  ]"
    # After finding "is not" run, the NEXT two runs form the bracket
    for i, run in enumerate(runs):
        if (run.text or "") == "is not " and i + 2 < len(runs):
            next1 = runs[i + 1].text or ""
            next2 = runs[i + 2].text or ""
            if next1.startswith("[ ") and next2.strip().endswith("]"):
                runs[i + 1].text = "[X"
                runs[i + 2].text = "]"
                changes += 1
                print(f"  FIXED P{idx:03d} split bracket checkbox runs {i+1}-{i+2}")
                break

    # ── P402: "(xx)" remnant after "3 (three)" ──
    # Runs: ..., "3 (three)", "xx", ")", " years..."
    for i in range(len(runs) - 2):
        if (runs[i].text or "") == "3 (three)" and (runs[i + 1].text or "") == "xx":
            if (runs[i + 2].text or "") == ")":
                runs[i + 1].text = ""
                runs[i + 2].text = ""
                changes += 1
                print(f"  FIXED P{idx:03d} xx) remnant at runs {i+1}-{i+2}")
                break

doc.save(G_DEST)
print(f"\nPatch complete: {changes} changes saved to {G_DEST}")

# Final verification
doc2 = docx.Document(G_DEST)
print("\n=== Final verification ===")
remaining_issues = 0
for i, p in enumerate(doc2.paragraphs):
    t = p.text or ""
    if "(INSERT PERFORMER" in t or "(ENTER CAGE" in t or "(ENTER UEI" in t or "xx)" in t:
        print(f"  STILL UNFILLED P{i:03d}: {t[:100]!r}")
        remaining_issues += 1
if remaining_issues == 0:
    print("  All performer placeholders resolved.")

for i, p in enumerate(doc2.paragraphs):
    t = p.text or ""
    if "[   ]" in t or ("[  ]" in t and "is not" in t):
        if "is not [X]" not in t and "is not [X" not in t:
            print(f"  UNCHECKED checkbox P{i:03d}: {t[:100]!r}")
