"""Fill Attachment H (Task Description Document) from official template."""
import shutil
import copy
from docx import Document
from docx.shared import Pt
from lxml import etree
from docx.oxml.ns import qn

H_SRC = (
    "docs/proposals/DARPA_MATHBAC/official_templates/"
    "Attachment_H_Task_Description_Document_TDD_Template.docx"
)
H_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/08_attachment_h_tdd/"
    "Attachment_H_SCBE_FILLED.docx"
)

shutil.copy2(H_SRC, H_DEST)
doc = Document(H_DEST)

# ─────────────────────────────────────────────────────────────────────────────
# Task content from ATTACHMENT_H_TDD_DRAFT_2026-05-31.md
# ─────────────────────────────────────────────────────────────────────────────

TASKS = [
    {
        "title": "Task 1: Program Kickoff, Baseline Definition, and Data Architecture",
        "period": "Month 0–1",
        "payment": "$84,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (limited to Sub Background IP review if needed)",
        "objective": (
            "Establish the Phase I mathematical framework, science subdomains, foundation models, "
            "baseline protocol, personnel assignments, initial evaluation rubric, and IV&V artifact structure."
        ),
        "description": (
            "SCBE will present the selected mathematical, information-theoretic, and systems-theory basis "
            "for the TA1 framework. The framework treats agents as composed operators over communication "
            "streams and latent-state readouts, with protocol progress measured through hyperbolic geometry, "
            "protocol graph structure, and proposer-added metrics. SCBE will define the initial task families, "
            "foundation models (ChemBERTa-77M, Qwen2.5-Coder-0.5B-Instruct), latent-space access plan, "
            "handcrafted baseline protocol, governance pipeline, data schema, and initial compute plan. "
            "The kickoff will include the IV&V bundle schema for communication streams, latent states, "
            "governance event logs, and metric summaries."
        ),
        "milestones": [
            "Kickoff briefing and foundation model / latent-space access plan delivered",
            "Baseline protocol description and initial TA1 rubric mapping delivered",
            "IV&V data bundle schema v0 delivered; personnel and level-of-effort confirmed",
        ],
        "deliverables": (
            "Kickoff briefing; foundation model and latent-space access plan; baseline protocol description; "
            "initial TA1 rubric mapping; personnel level-of-effort confirmation; draft IV&V data bundle schema; "
            "initial compute-by-month plan."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 2: Initial Mathematical Framework and Early Baseline Experiments",
        "period": "Month 1–3",
        "payment": "$109,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (optional context only if verified artifacts available)",
        "objective": (
            "Produce the first technical report on early successes/failures, baseline model behavior, "
            "and initial mathematical-framework implementation."
        ),
        "description": (
            "SCBE will implement the initial operator-chain instrumentation for selected fixed pretrained agents "
            "and run early baseline tasks. Work includes measuring communication acts, extracting latent probes "
            "where available, recording token streams, producing initial hyperbolic states, and scoring the "
            "resulting interactions through SCBE metrics. SCBE will document successes and failures, including "
            "where latent-state signals are weak, where output-only operation is required, and where protocol "
            "graph instrumentation requires calibration. Lyapunov constants eta and b will be measured for the "
            "NMR subdomain. Oracle thresholds (I_min, gamma_min, MEE_min) calibrated. ACV field and MEE score "
            "added to event schema."
        ),
        "milestones": [
            "M3 technical report delivered with initial mathematical framework description",
            "Baseline experiment results including initial hyperbolic state and SCBE metric scores",
            "IV&V bundle v1 transmitted; ACV and MEE fields live in event schema",
        ],
        "deliverables": (
            "M3 technical report; initial mathematical framework description; baseline experiment results; "
            "personnel level-of-effort confirmation; updated risk register; IV&V bundle v1."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 3: Framework Demonstration, Metrics Progress, and IV&V Data Pipeline",
        "period": "Month 3–6",
        "payment": "$151,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (bounded review of DAVA-related support material if included)",
        "objective": (
            "Demonstrate the framework at the first PI meeting, show progress against Phase I metrics, "
            "and deliver a functioning IV&V data capture and sharing pipeline."
        ),
        "description": (
            "SCBE will expand the mathematical framework, implement the metric capture path for MEE, ACV, "
            "CDPTI, and PIS where available, and demonstrate the ability to capture and process agentic "
            "communication streams and latent-space data. SCBE will package run artifacts in a reproducible "
            "IV&V bundle. The demonstration will include concrete examples in the NMR task families and a "
            "comparison to the handcrafted baseline protocol. CDPTI live computation will be demonstrated. "
            "At least two NMR subtask variants will demonstrate adaptability."
        ),
        "milestones": [
            "M6 PI meeting briefing delivered; CDPTI live computation demonstrated",
            "IV&V artifact bundle v2 demonstrated; governance event schema and latent-state archive format delivered",
            "Framework description showing progress against all four proposer-added metrics",
        ],
        "deliverables": (
            "M6 PI meeting briefing; framework description; metric progress report; "
            "IV&V artifact bundle demonstration; governance event schema; latent-state archive format; "
            "quarterly technical report."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 4: Software Suite, Reduced-Order Models, and Challenge-Problem Progress",
        "period": "Month 6–9",
        "payment": "$151,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (optional Sub Background IP / telemetry context only if finalized)",
        "objective": (
            "Deliver an initial software-suite report, progress on IV&V challenge problems, "
            "and reduced-order models for selected agents."
        ),
        "description": (
            "SCBE will formalize the software-suite structure around the operator framework, protocol graph, "
            "metrics, and computational design tool components. The reduced-order model work will estimate "
            "selected agent outputs at specified fidelity using observed communication and latent-state records. "
            "SCBE will adapt to IV&V challenge problem inputs when available and document any mismatch between "
            "proposal task families and IV&V task definitions. Cheminformatics (Hammett) secondary generalization "
            "lane results delivered."
        ),
        "milestones": [
            "Initial software-suite report delivered; ROMs for ChemBERTa-77M and Qwen2.5-Coder-0.5B",
            "IV&V challenge-problem progress report; Task V (cheminformatics) results delivered",
            "Updated compute usage and forecast; quarterly technical report",
        ],
        "deliverables": (
            "Initial software-suite report; reduced-order model report; IV&V challenge-problem progress report; "
            "updated compute usage and forecast; quarterly technical report."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 5: Side-by-Side Baseline Comparison and Initial Protocol Set",
        "period": "Month 9–13",
        "payment": "$151,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (only for bounded corroborating context if approved)",
        "objective": (
            "Provide a second PI meeting update, side-by-side baseline comparison, and initial set "
            "of superior protocols with mathematical explanation."
        ),
        "description": (
            "SCBE will compare the SCBE-governed protocol against baseline protocols (Mixtral-8x7B ungoverned) "
            "using the living IV&V rubric. The comparison will report success rate, speedup, adaptability, "
            "generalization behavior, governance overhead, and all four proposer-added metrics. SCBE will "
            "identify protocol patterns that improve task success or auditability and explain why those patterns "
            "work or fail using the mathematical framework. PIS catalog v1 will be delivered."
        ),
        "milestones": [
            "M13 PI meeting briefing delivered; baseline comparison table with mathematical explanation",
            "Initial superior-protocol set (PIS catalog v1) delivered",
            "IV&V challenge-problem progress report; quarterly technical report",
        ],
        "deliverables": (
            "M13 PI meeting briefing; side-by-side baseline comparison; initial superior-protocol set; "
            "mathematical explanation of protocol performance; IV&V challenge-problem progress report; "
            "quarterly technical report."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 6: Computational Design Tool Demonstration",
        "period": "Month 13–14",
        "payment": "$84,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (none unless specifically approved)",
        "objective": (
            "Demonstrate the computational design tool and provide test results with IV&V."
        ),
        "description": (
            "SCBE will demonstrate a computational tool that takes task conditions, agent access mode, "
            "and metric constraints as inputs and produces candidate communication-protocol designs. "
            "The demonstration will include test runs, generated protocol candidates, metric reports, "
            "and replayable artifact bundles. Test results will be shared with IV&V. "
            "Calibration time for new subdomains targeted at ≤ 1 week."
        ),
        "milestones": [
            "Computational design tool demonstration delivered with test results",
            "Tool user notes / draft manual delivered",
            "Replayable run artifacts delivered to IV&V",
        ],
        "deliverables": (
            "Computational design tool demonstration; test results with IV&V; "
            "tool user notes / draft manual; replayable run artifacts."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
    {
        "title": "Task 7: Final Report, Protocol Catalog, and Phase II Plan",
        "period": "Month 14–16",
        "payment": "$109,000",
        "lead": "SCBE / Issac Davis (Prime)",
        "support": "Hoags Inc. (final Sub Background IP review if included)",
        "objective": (
            "Deliver the final Phase I report, catalog of protocol design principles, domain-specific "
            "optimal protocols, implementation summary, and Phase II plan."
        ),
        "description": (
            "SCBE will summarize mathematical, informational, and systems-theory developments; document "
            "software implementations; present the protocol catalog; provide final benchmark and IV&V results; "
            "and describe a Phase II path for integrating TA1 and TA2 capabilities. The final report will "
            "include a grounded Phase II example (Karplus rediscovery through self-evolving agent protocol), "
            "anticipated reward structure (MEE + ACV + CDPTI), known limitations, and transition plan."
        ),
        "milestones": [
            "Final technical report delivered with complete IV&V data archive",
            "Catalog of protocol design principles and domain-specific optimal protocol set delivered",
            "Phase II plan and ROM update delivered; phase completion report within 30 days of phase end",
        ],
        "deliverables": (
            "Final technical report; quantitative mathematical framework summary; computational design tool "
            "package summary; catalog of protocol design principles; domain-specific optimal protocol set; "
            "Phase II plan and ROM update; phase completion report within 30 days of phase end."
        ),
        "human_subjects": "No",
        "location": "2361 E 5th Ave, Port Angeles, WA 98362 (remote / local)",
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell(cell, text, font_pt=10):
    """Clear and set text in a table cell."""
    for para in cell.paragraphs:
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
    if not cell.paragraphs:
        cell.add_paragraph()
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(font_pt)


def fill_task_table(tbl, task, task_num):
    """Fill a 9-row x 4-col task table from task dict."""
    # Row 0: Task title (merged across all 4 cols — they share the same value due to merge)
    title_text = f"Task {task_num}: {task['title'].split(':', 1)[1].strip() if ':' in task['title'] else task['title']}"
    title_text += f"  |  {task['period']}  |  Milestone Payment: {task['payment']}"
    for c in range(4):
        try:
            set_cell(tbl.cell(0, c), title_text if c == 0 else "", font_pt=10)
        except Exception:
            pass

    # Row 1: Objective | Task Description | Location | Location
    try:
        set_cell(tbl.cell(1, 0), "Objective", font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(1, 1), task["objective"], font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(1, 2), task["location"], font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(1, 3), task["location"], font_pt=10)
    except Exception:
        pass

    # Row 2: (empty in template — use for work description)
    try:
        set_cell(tbl.cell(2, 0), "Work Description", font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(2, 1), task["description"], font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(2, 2), f"Lead: {task['lead']}\nSupport: {task['support']}", font_pt=10)
    except Exception:
        pass

    # Row 3: Primary Organization
    try:
        set_cell(tbl.cell(3, 0), "Primary Organization Responsible", font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(3, 1), task["lead"], font_pt=10)
    except Exception:
        pass

    # Row 4: Human Subjects
    try:
        set_cell(tbl.cell(4, 0), "Human Subjects or Animal Research?", font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(4, 1), task["human_subjects"], font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(4, 2), task["human_subjects"], font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(4, 3), task["human_subjects"], font_pt=10)
    except Exception:
        pass

    # Row 5: Milestones | Deliverables
    try:
        set_cell(tbl.cell(5, 0), "Associated Milestones", font_pt=10)
    except Exception:
        pass
    try:
        set_cell(tbl.cell(5, 1), task["deliverables"], font_pt=10)
    except Exception:
        pass

    # Rows 6, 7, 8: milestone items
    for i, ms in enumerate(task["milestones"]):
        row = 6 + i
        if row < len(tbl.rows):
            try:
                set_cell(tbl.cell(row, 0), f"{i + 1}.", font_pt=10)
            except Exception:
                pass
            try:
                set_cell(tbl.cell(row, 1), ms, font_pt=10)
            except Exception:
                pass


# ─────────────────────────────────────────────────────────────────────────────
# Fill objective paragraph
# ─────────────────────────────────────────────────────────────────────────────

for para in doc.paragraphs:
    if "[Insert a general description of the objective" in para.text:
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
        run = para.add_run(
            "The objective of this proposal is: Develop and empirically validate a composed-operator "
            "governance framework T = L₁₄ ∘ ⋯ ∘ L₁ for multi-agent AI communication dynamics in "
            "scientific discovery tasks. The framework maps agent communication acts to points in "
            "a Poincaré ball, measures safety by hyperbolic distance from the safe-operation origin, "
            "and enforces five physical axioms across 14 layers. Phase I demonstrates the framework "
            "on NMR spectroscopy task families (Karplus equation rediscovery) using ChemBERTa-77M "
            "and Qwen2.5-Coder-0.5B-Instruct as small science models."
        )
        run.font.size = Pt(11)
        break

# ─────────────────────────────────────────────────────────────────────────────
# Strategy: the template has 4 tables:
#   Table 0: Task 1 (9x4)
#   Table 1: Subtask 1.A (9x4)
#   Table 2: Subtask 1.B (9x4)
#   Table 3: Task # (9x4) — generic template
#
# We need 7 task tables. Strategy:
#   1. Fill Table 0 as Task 1
#   2. Delete Table 1 (Subtask 1.A) and Table 2 (Subtask 1.B) — no subtasks
#   3. Fill Table 3 as Task 2
#   4. Clone Table 3 (now filled as Task 2) for Tasks 3-7
# ─────────────────────────────────────────────────────────────────────────────

# Fill Task 1 into Table 0
fill_task_table(doc.tables[0], TASKS[0], 1)

# Delete Subtask tables (Table 1 and Table 2)
# Tables in python-docx are inline with body; we remove their parent <w:tbl> elements
for _ in range(2):
    # After deleting, Table 1 becomes the old Table 2, so always delete index 1
    tbl = doc.tables[1]
    tbl._tbl.getparent().remove(tbl._tbl)

# Now doc.tables[1] is the old Table 3 (Task # template) — fill as Task 2
fill_task_table(doc.tables[1], TASKS[1], 2)

# Clone and fill for Tasks 3-7
task_template_tbl = doc.tables[1]
prev_tbl_elem = task_template_tbl._tbl

for i, task in enumerate(TASKS[2:], start=3):  # Tasks 3-7
    # Deep clone the task template
    new_tbl_elem = copy.deepcopy(task_template_tbl._tbl)
    # Insert after previous table element
    prev_tbl_elem.addnext(new_tbl_elem)
    prev_tbl_elem = new_tbl_elem
    # Get the Document-level Table wrapper for the newly inserted table
    # We need to find it in doc.tables — it's the last one
    new_tbl = doc.tables[-1]
    fill_task_table(new_tbl, task, i)

# ─────────────────────────────────────────────────────────────────────────────
# Recurring Deliverables section — add after last table as paragraphs
# ─────────────────────────────────────────────────────────────────────────────

last_tbl = doc.tables[-1]
body = doc.element.body
last_tbl_idx = list(body).index(last_tbl._tbl)

recurring_lines = [
    "",
    "RECURRING DELIVERABLES",
    "",
    "Quarterly Technical Reports: SCBE will provide quarterly technical reports within 10 days of each "
    "quarter end. Reports will summarize technical progress, metric progress, risks, compute usage, "
    "IV&V interactions, and upcoming work.",
    "",
    "Government / IV&V Meetings: SCBE will support regular teleconferences with the Government and "
    "IV&V team. The cost build includes three required three-day meetings: two DC-area and one "
    "San Francisco Bay-area meeting.",
    "",
    "Evaluation Briefings: SCBE will brief DARPA on each major evaluation and provide written "
    "summaries within two weeks when required.",
    "",
    "COMPUTE AND RESOURCES",
    "",
    "SCBE will use a prime-owned Linux AI workstation (RTX 4090-class), local development resources, "
    "and intermittent commercial GPU endpoint capacity. Compute will be tracked by month and reported "
    "in milestone/quarterly updates.",
    "",
    "SUBCONTRACT BOUNDARIES",
    "",
    "Hoags Inc. supporting work is limited to: DAVA background-IP descriptions and data-rights assertions; "
    "DAVA trace-generation / sealed-protocol context if verified; phi_beacon telemetry specification if "
    "artifact support is finalized; net_probe / phi-push context if artifact support is finalized; "
    "review of final Sub Background IP rows.",
    "",
    "Excluded unless explicitly reopened: Hoags as proposal lead or joint investigator; "
    "public GitHub repository content as evidence; voice grammar / language dictionary stack; "
    "unresolved surface claims; any private-code dependency that prevents SCBE from executing Phase I.",
]

for line in recurring_lines:
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "Normal")
    r = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), "20")  # 10pt
    t = etree.SubElement(r, qn("w:t"))
    t.text = line if line else " "
    if line and (line[0] == " " or line[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    body.append(p)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────

doc.save(H_DEST)
print(f"Saved: {H_DEST}")
print(f"Tables in final doc: {len(doc.tables)}")
