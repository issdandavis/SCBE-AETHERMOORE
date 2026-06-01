"""Inspect D, E, H templates for placeholder structure."""
import docx
import re

TEMPLATES = {
    "D": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_D_Proposal_Instructions_and_Volume_I_Template__Technical_and_Management.docx"
    ),
    "E": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_E_Proposal_Instructions_and_Volume_II_Template_Price.docx"
    ),
    "H": (
        "docs/proposals/DARPA_MATHBAC/official_templates/"
        "Attachment_H_Task_Description_Document_TDD_Template.docx"
    ),
}

BLUE = "00B0F0"
INSERT_RE = re.compile(r"\[INSERT|<INSERT|\(INSERT|\[ENTER|\[FILL|\[PROPOSER", re.I)


def is_blue(run):
    try:
        return run.font.color.rgb and str(run.font.color.rgb).upper() == BLUE.upper()
    except Exception:
        return False


for name, path in TEMPLATES.items():
    print(f"\n{'='*60}")
    print(f"  Attachment {name}: {path}")
    print(f"{'='*60}")
    doc = docx.Document(path)
    print(f"  Total paragraphs: {len(doc.paragraphs)}")

    blue_paras = []
    insert_paras = []
    headings = []

    for idx, para in enumerate(doc.paragraphs):
        txt = para.text or ""
        style = para.style.name if para.style else ""

        # Collect headings
        if "Heading" in style:
            headings.append((idx, style, txt[:80]))

        # Blue runs
        if any(is_blue(r) for r in para.runs if r.text.strip()):
            blue_paras.append((idx, txt[:120]))

        # INSERT patterns
        if INSERT_RE.search(txt):
            insert_paras.append((idx, txt[:120]))

    print(f"\n  --- Headings ({len(headings)}) ---")
    for idx, style, txt in headings[:30]:
        print(f"    P{idx:04d} [{style}]: {txt!r}")
    if len(headings) > 30:
        print(f"    ... ({len(headings) - 30} more)")

    print(f"\n  --- Blue runs ({len(blue_paras)}) ---")
    for idx, txt in blue_paras[:30]:
        print(f"    P{idx:04d}: {txt!r}")
    if len(blue_paras) > 30:
        print(f"    ... ({len(blue_paras) - 30} more)")

    print(f"\n  --- INSERT patterns ({len(insert_paras)}) ---")
    for idx, txt in insert_paras[:20]:
        print(f"    P{idx:04d}: {txt!r}")

    # Tables
    print(f"\n  --- Tables: {len(doc.tables)} ---")
    for t_idx, tbl in enumerate(doc.tables[:5]):
        rows = len(tbl.rows)
        cols = len(tbl.columns) if tbl.columns else "?"
        first_cell = tbl.cell(0, 0).text[:60] if rows > 0 else ""
        print(f"    Table {t_idx}: {rows} rows x {cols} cols | first cell: {first_cell!r}")
