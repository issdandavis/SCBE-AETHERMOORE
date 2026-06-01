"""Fill Attachment G and Attachment I performer fields."""
import shutil
import docx
from docx.shared import Pt

BLUE = "00B0F0"


def is_blue(run):
    try:
        return run.font.color.rgb and str(run.font.color.rgb) == BLUE
    except Exception:
        return False


def replace_blue_runs(para, fragment, new_text):
    """
    Find any blue run whose text contains `fragment`, replace that run's
    text with new_text, then clear subsequent blue runs that are part of
    the same placeholder span (up to 3 consecutive runs).
    Returns True on first match.
    """
    runs = para.runs
    n = len(runs)
    for i, run in enumerate(runs):
        if not is_blue(run):
            continue
        # Check single-run match
        if fragment in (run.text or ""):
            run.text = new_text
            return True
        # Check multi-run match (up to 3 runs)
        for j in range(i + 1, min(i + 4, n)):
            combined = "".join(r.text or "" for r in runs[i : j + 1])
            if fragment in combined:
                runs[i].text = new_text
                for k in range(i + 1, j + 1):
                    runs[k].text = ""
                return True
    return False


def replace_any_run(para, old, new):
    for run in para.runs:
        if old in (run.text or ""):
            run.text = (run.text or "").replace(old, new)
            return True
    return False


# ──────────────────────────────────────────────────────────────────────────────
# ATTACHMENT G
# ──────────────────────────────────────────────────────────────────────────────
G_SRC = (
    "docs/proposals/DARPA_MATHBAC/official_templates/"
    "Attachment_G_Model_Other_Transaction_for_Research_Agreement_MATHBAC_.docx"
)
G_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/07_attachment_g_model_ot/"
    "Attachment_G_SCBE_FILLED.docx"
)
shutil.copy2(G_SRC, G_DEST)
doc_g = docx.Document(G_DEST)

PERFORMER_NAME = "ISSAC D DAVIS / SCBE-AETHERMOORE"
PERFORMER_ADDR = "ISSAC D DAVIS / SCBE-AETHERMOORE, 2361 E 5th Ave, Port Angeles, WA 98362"
R_AND_D = (
    "a SCBE mathematical framework for governing multi-agent AI communication "
    "protocols using a composed Poincare-ball operator with five physical axioms "
    "(Unitarity, Locality, Causality, Symmetry, Composition)"
)
GOAL = (
    "develop a mathematically rigorous framework for governing multi-agent AI "
    "communication protocols using a composed Poincare-ball operator with five "
    "physical axioms, demonstrated on NMR spectroscopy task families, producing "
    "certified convergence bounds, a protocol coherence drift index (CDPTI), and "
    "a computational design tool for applying the framework to new scientific subdomains"
)
EXEC = "Principal Investigator (sole proprietor; no higher organizational authority applies)"

blue_subs = [
    ("(INSERT PERFORMER AND ADDRESS)", PERFORMER_ADDR),
    ("(INSERT PERFORMER NAME)", PERFORMER_NAME),
    ("(INSERT PERFORMER)", PERFORMER_NAME),
    ("(INSERT RESEARCH AND DEVELOPMENT EFFORT)", R_AND_D),
    ("(INSERT GOAL(S) OF AGREEMENT)", GOAL),
    ("(ENTER CAGE CODE)", "1EXD5"),
    ("(ENTER UEI)", "J4NXHM6N5F59"),
    ("(ENTER TIN)", "[TIN — ENTER DIRECTLY IN FINAL SIGNED DOCUMENT; DO NOT ADD TO REPO]"),
    (
        "(INSERT A LEVEL OF EXECUTIVE FAR ENOUGH REMOVED FROM THE PROGRAM TO MAINTAIN A GREATER LEVEL OF IMPARTIALITY)",
        EXEC,
    ),
    ("(The senior POC in paragraph 4 must be at least 1 level higher than the senior POC being referenced within paragraph 3.)", ""),
    ("[ONLY INCLUDE IF APPLICABLE]", ""),
]

changed_g = 0
for para in doc_g.paragraphs:
    t = para.text
    for fragment, replacement in blue_subs:
        if fragment in t:
            if replace_blue_runs(para, fragment, replacement):
                changed_g += 1
                t = para.text  # refresh
                break

    # Number of months
    if "(INSERT NUMBER OF MONTHS)" in (para.text or ""):
        if replace_blue_runs(para, "(INSERT NUMBER OF MONTHS)", "16 (sixteen)"):
            replace_any_run(para, "(xx)", "")
            changed_g += 1

    # Number of years — two instances
    if "(INSERT NUMBER OF YEARS)" in (para.text or ""):
        if replace_blue_runs(para, "(INSERT NUMBER OF YEARS)", "3 (three)"):
            replace_any_run(para, "(xx)", "")
            changed_g += 1

    # Address block
    if "(ENTER ADDRESS)" in (para.text or ""):
        replace_any_run(para, "(ENTER ADDRESS)", "2361 E 5th Ave")
        replace_any_run(para, "(XXXX, XXXX, XXXX)", "Port Angeles, Clallam County, WA")
        replace_any_run(para, "(XXXXX)", "98362")
        changed_g += 1

# Checkbox paragraphs — mark "is not" for both corporation felony/tax lines
for para in doc_g.paragraphs:
    txt = para.text or ""
    if "It is" in txt and "is not" in txt and "corporation" in txt:
        # Find the run containing second "[   ]" bracket (after "is not")
        found_is_not = False
        for run in para.runs:
            rt = run.text or ""
            if "is not" in rt:
                found_is_not = True
                continue
            if found_is_not and ("[   ]" in rt or "[  ]" in rt):
                run.text = rt.replace("[   ]", "[X]").replace("[  ]", "[X]")
                changed_g += 1
                found_is_not = False
                break

doc_g.save(G_DEST)
print(f"G: {changed_g} fills applied -> {G_DEST}")

# Spot-check G
doc_gv = docx.Document(G_DEST)
print("\n=== G spot checks ===")
for para in doc_gv.paragraphs:
    t = para.text or ""
    for needle in ("ISSAC", "1EXD5", "J4NXHM6N5F59", "3 (three)", "16 (sixteen)", "[X]", "TIN —"):
        if needle in t:
            print(f"  OK [{needle}]: {t[:90]!r}")
            break
    for bad in ("(INSERT PERFORMER", "(ENTER CAGE", "(ENTER UEI", "(INSERT NUMBER"):
        if bad in t:
            print(f"  UNFILLED: {t[:90]!r}")

# ──────────────────────────────────────────────────────────────────────────────
# ATTACHMENT I
# ──────────────────────────────────────────────────────────────────────────────
I_SRC = (
    "docs/proposals/DARPA_MATHBAC/official_templates/"
    "Attachment_I_Other_Transaction_Certifications.docx"
)
I_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/08_attachment_i_certifications/"
    "Attachment_I_SCBE_FILLED.docx"
)
shutil.copy2(I_SRC, I_DEST)
doc_i = docx.Document(I_DEST)

changed_i = 0
entity_filled = False
for para in doc_i.paragraphs:
    t = para.text or ""

    # Entity name blank line — first paragraph that is entirely underscores
    if not entity_filled and t.strip("_").strip() == "" and len(t.strip()) > 5:
        for run in para.runs:
            run.text = "ISSAC D DAVIS / SCBE-AETHERMOORE"
            entity_filled = True
            changed_i += 1
            break

    # Address
    if "[Street Address]" in t:
        replace_any_run(para, "[Street Address]", "2361 E 5th Ave")
        replace_any_run(para, "[City, County, State]", "Port Angeles, Clallam County, WA")
        replace_any_run(para, "[Zip code]", "98362")
        changed_i += 1

    # Typed name / org
    if "[Typed Name and Title of Official Responsible for This Transaction]" in t:
        replace_any_run(
            para,
            "[Typed Name and Title of Official Responsible for This Transaction]",
            "Issac Davis, Principal Investigator",
        )
        replace_any_run(
            para, "[Name of Organization/Institution]", "SCBE-AETHERMOORE"
        )
        changed_i += 1

    # Tax/felony checkboxes — mark "is not"
    if "It is" in t and "is not" in t and ("corporation" in t or "felony" in t or "unpaid" in t):
        for j, run in enumerate(para.runs):
            rt = run.text or ""
            if "is not" in rt:
                # Insert [X] right after "is not" in this run
                run.text = rt.replace("is not", "is not [X]", 1)
                changed_i += 1
                break

doc_i.save(I_DEST)
print(f"\nI: {changed_i} fills applied -> {I_DEST}")

print("\n=== I spot checks ===")
doc_iv = docx.Document(I_DEST)
for para in doc_iv.paragraphs:
    t = para.text or ""
    for needle in ("ISSAC", "Port Angeles", "[X]", "Issac Davis"):
        if needle in t:
            print(f"  OK [{needle}]: {t[:90]!r}")
            break
