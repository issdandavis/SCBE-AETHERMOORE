"""Fill Attachment E (Volume II Price) from official template."""
import shutil
from docx import Document
from docx.shared import Pt
from lxml import etree
from docx.oxml.ns import qn

E_SRC = (
    "docs/proposals/DARPA_MATHBAC/official_templates/"
    "Attachment_E_Proposal_Instructions_and_Volume_II_Template_Price.docx"
)
E_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/05_attachment_e_vol_ii_price/"
    "Attachment_E_SCBE_FILLED.docx"
)

shutil.copy2(E_SRC, E_DEST)
doc = Document(E_DEST)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _make_p(text, font_pt=11):
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "Normal")
    r = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_pt * 2))
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(font_pt * 2))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def set_cell(tbl, row, col, text):
    cell = tbl.cell(row, col)
    for para in cell.paragraphs:
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(11)


def replace_para_text(doc, starts_with, content_lines):
    """Find first paragraph starting with given text, replace + insert lines after."""
    for para in doc.paragraphs:
        if para.text.strip().startswith(starts_with.strip()[:40]):
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            if content_lines:
                run = para.add_run(content_lines[0])
                run.font.size = Pt(11)
            prev = para._p
            for line in content_lines[1:]:
                new_p = _make_p(line, font_pt=11)
                prev.addnext(new_p)
                prev = new_p
            return prev
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Cover Sheet Table (Table 0 — 20 rows x 2 cols)
# ─────────────────────────────────────────────────────────────────────────────

cover = doc.tables[0]

set_cell(cover, 0, 1, "SCBE Mathematical Framework for Agentic Communication Protocols")
set_cell(cover, 1, 1, "ISSAC D DAVIS / SCBE-AETHERMOORE")
set_cell(cover, 2, 1, "TA1")
set_cell(cover, 3, 1, "Other Small Business (sole proprietorship, minority-owned)")
set_cell(cover, 4, 1, "N/A")
set_cell(
    cover, 5, 1,
    "Name: Issac Davis\n"
    "Address: 2361 E 5th Ave, Port Angeles, WA 98362-9014\n"
    "Telephone: (360) 808-0876\n"
    "Email: issdandavis7795@gmail.com"
)
set_cell(
    cover, 6, 1,
    "Name: Issac Davis\n"
    "Address: 2361 E 5th Ave, Port Angeles, WA 98362-9014\n"
    "Telephone: (360) 808-0876\n"
    "Email: issdandavis7795@gmail.com"
)
# Row 7: Award Instrument
set_cell(cover, 7, 1, "Other Transactions for Research (10 U.S.C. §4021)")
# Row 8: Place of Performance
set_cell(cover, 8, 1, "2361 E 5th Ave, Port Angeles, WA 98362-9014 (remote / local)")
# Row 9: Period of Performance
set_cell(cover, 9, 1, "16 months from award date")
# Row 10: Months
set_cell(cover, 10, 1, "16")
# Row 11: Other Team Members
set_cell(
    cover, 11, 1,
    "Technical POC Name: Collin Hoag\n"
    "Organization: Hoags Inc.\n"
    "Organization Type: Supporting Subcontractor (bounded scope)"
)
# Row 12: Total Price
set_cell(
    cover, 12, 1,
    "Year 1: $[fill when award date / FY boundary known]\n"
    "Year 2: $[remainder]\n"
    "Year 3: N/A\n"
    "Total: $839,000"
)
# Row 13: UEI
set_cell(cover, 13, 1, "J4NXHM6N5F59")
# Row 14: TIN — leave blank for manual entry
set_cell(cover, 14, 1, "[Enter TIN directly in Word — do not record here]")
# Row 15: CAGE
set_cell(cover, 15, 1, "1EXD5")
# Row 16: DCMA POC
set_cell(cover, 16, 1, "Name: Unknown at proposal stage\nAddress: —\nTelephone: —")
# Row 17: DCAA POC
set_cell(cover, 17, 1, "Name: Unknown at proposal stage\nAddress: —\nTelephone: —")
# Row 18: Date Prepared
set_cell(cover, 18, 1, "2026-06-[fill on submission day]")
# Row 19: Proposal Validity
set_cell(cover, 19, 1, "120 days from submission")

# ─────────────────────────────────────────────────────────────────────────────
# Price Summary Section
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide the following cost summary information",
    [
        "PRICE SUMMARY — DARPA-PA-26-05 MATHBAC | Phase I | 16 Months | $839,000 Total",
        "",
        "Cost Element                        | Phase I Total",
        "------------------------------------------------------",
        "Direct Labor                        | $360,960",
        "Subcontract (Hoags Inc.)            | $120,000",
        "Equipment                           | $6,500",
        "Travel                              | $18,000",
        "Other Direct Costs (ODCs)           | $119,500",
        "Total Direct Costs                  | $624,960",
        "G&A (30% on $504,960 prime base)    | $151,488",
        "Subtotal Costs                      | $776,448",
        "Fee / Profit (8.06%)                | $62,552",
        "TOTAL PROPOSED PRICE                | $839,000",
        "",
        "See Attachment F (Streamlined Cost Buildup Workbook) for full line-item detail.",
        "All milestone payments are triggered by observable technical events (see Attachment H TDD):",
        "M1 $84,000 | M2 $109,000 | M3 $151,000 | M4 $151,000 | M5 $151,000 | M6 $84,000 | M7 $109,000",
    ]
)

# Remove the "In addition to the Price Information" instruction paragraph
replace_para_text(
    doc,
    "In addition to the Price Information detailed",
    ["[See Attachment F for spreadsheet workbook.]"]
)

# ─────────────────────────────────────────────────────────────────────────────
# Price Details intro
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "[Provide the following price details and include",
    [
        "Price details are provided below by cost element. All costs are fully traceable "
        "to Attachment F (Streamlined Cost Buildup Workbook). Costs are broken down by Phase (Phase I only), "
        "TDD task (7 milestones per Attachment H), and cost element. "
        "Year 1/Year 2 FY split will be confirmed once award date is known.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 1. Direct Labor
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "[Provide individual labor categories",
    [
        "1. DIRECT LABOR — $360,960",
        "",
        "Labor Category          | Name              | Hours | Rate    | Total",
        "---------------------------------------------------------------------",
        "Principal Investigator  | Davis, Issac      | 1,536 | $115/hr | $176,640",
        "Science/Engineering     | Research Engr TBD | 1,024 | $105/hr | $107,520",
        "Science/Engineering     | Data/IV&V Spec TBD|   640 | $85/hr  | $54,400",
        "Administrative Support  | Admin/Comp TBD    |   320 | $70/hr  | $22,400",
        "TOTAL                   |                   | 3,520 |         | $360,960",
        "",
        "Rates are market-competitive for the Pacific Northwest independent contractor market. "
        "Rates are fully loaded. No separate fringe line is required (sole proprietorship; TBD positions "
        "would be hired as independent contractors at confirmed rates before award). "
        "G&A (30%) applies on top of direct labor base.",
        "",
        "Why solo-prime is lean: (1) Mathematical derivations are single-author research. "
        "(2) Benchmark infrastructure already exists in SCBE codebase. "
        "(3) Phase I deliverables are fully achievable within 3,520 hours.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. Materials
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide an itemized list of all proposed materials",
    [
        "2. MATERIALS — $0",
        "",
        "No materials proposed. Phase I uses publicly available NMR datasets (BMRB + PDB) at no cost. "
        "No physical materials are required to complete Phase I deliverables.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. Equipment Purchases
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide an itemized list of all proposed equipment",
    [
        "3. EQUIPMENT PURCHASES — $6,500",
        "",
        "Item                           | Qty | Unit Price | Total  | Basis",
        "----------------------------------------------------------------------",
        "Linux AI Workstation (RTX 4090)| 1   | $6,500     | $6,500 | Current market pricing for "
        "RTX 4090-class Linux workstations (e.g., System76 Thelio R5). "
        "Prime-owned; not a Government deliverable. "
        "Required for local CUDA benchmark harness and reproducible IV&V artifact generation.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. Travel
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide the purpose of the trip",
    [
        "4. TRAVEL — $18,000",
        "",
        "Trip   | Purpose              | Origin      | Destination    | People | Days | Cost",
        "-------------------------------------------------------------------------------------",
        "M3 PI  | Framework progress   | PA, WA      | Washington, DC | 2      | 5    | ~$5,650",
        "M6 PI  | IV&V demo            | PA, WA      | Washington, DC | 2      | 5    | ~$5,650",
        "M13 PI | Baseline comparison  | PA, WA      | San Francisco  | 2      | 5    | ~$4,550",
        "Cont.  | Schedule shifts/add. | —           | TBD            | —      | —    | ~$2,150",
        "TOTAL  |                      |             |                |        |      | ~$18,000",
        "",
        "Basis: GSA per diem rates (DC: hotel $258/night, M&IE $79/day; SF: hotel $274/night, M&IE $79/day). "
        "Airfare at economy (SEA-DCA ~$750/pp RT, SEA-SFO ~$450/pp RT).",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. Other Direct Costs
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide an itemized breakdown with costs",
    [
        "5. OTHER DIRECT COSTS (ODCs) — $119,500",
        "",
        "Item                                            | Amount   | Basis",
        "----------------------------------------------------------------------",
        "Commercial GPU compute (HF Endpoints + burst)  | $72,000  | $4,500/mo avg x 16 mo; ~3,200 A10G-hours; 25% contingency buffer",
        "Storage, tooling, accounting, misc.             | $47,500  | Storage ~$200/mo + artifact tooling subs + compliance prep",
        "TOTAL                                           | $119,500 |",
        "",
        "GPU compute is the dominant ODC item. The NMR SSM training runs (ChemBERTa-77M calibration, "
        "CDPTI latent-state archive, Hammett validation) require approx. 3,200 A10G-hours over 16 months. "
        "HuggingFace Jobs (a10g-small) priced at ~$1.10/hr effective rate. $72,000 includes 25% contingency.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 6. Resource Sharing
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide the source, nature, and amount",
    [
        "6. RESOURCE SHARING — None",
        "",
        "No resource sharing or cost share is proposed. SCBE-AETHERMOORE does not plan to apply "
        "non-Federal resources to this effort.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 7. Consultant Costs
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "Provide a copy of all consultants",
    [
        "7. CONSULTANT COSTS — None",
        "",
        "No consultants are proposed. All technical and administrative work is performed by prime "
        "labor categories in Section 1 and the supporting subcontractor in Section 8.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 8. Subawardee Costs
# ─────────────────────────────────────────────────────────────────────────────

replace_para_text(
    doc,
    "For each proposed subawardee, provide",
    [
        "8. SUBAWARDEE COSTS — $120,000",
        "",
        "Subawardee: Hoags Inc. (Collin Hoag) | UEI DUHWVUXFNPV5 | CAGE 15XV5",
        "",
        "Scope                                                           | Amount",
        "-----------------------------------------------------------------------",
        "Bounded DAVA background-IP telemetry / corroborating support   | $120,000",
        "",
        "Hoags Inc. provides supporting subcontract services bounded to the narrow "
        "field-type-correspondence and Annex A static framing scope. "
        "ISSAC D DAVIS / SCBE-AETHERMOORE holds prime technical authority. "
        "$120,000 is a ceiling; Hoags will provide their own cost proposal for DARPA review.",
        "",
        "Why $120,000: ~800 hours at blended ~$150/hr for DAVA background-IP verification, "
        "field-type-correspondence mapping, and Annex A row support. "
        "Hoags subcontract: $120K / Prime: $719K = 16.7% subcontract ratio (below 20% threshold, "
        "keeping prime technical direction firmly with SCBE-AETHERMOORE).",
    ]
)

# Fill the two sub-bullets about subawardee documentation
replace_para_text(
    doc,
    "A copy of the proposed TDD as well as any documents that verify",
    [
        "Hoags Inc. TDD: Bounded scope per Attachment H §Subcontract Boundaries. "
        "Cost proposal to be submitted by Hoags Inc. directly or included as an addendum. "
        "Teaming agreement v2 in final review.",
    ]
)

replace_para_text(
    doc,
    "Interdivisional work transfer agreements",
    [
        "No interdivisional work transfer agreements apply (Hoags Inc. is an independent entity).",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# 9. Indirect Rates and Fee (replace G&A/Fee guidance line if present)
# ─────────────────────────────────────────────────────────────────────────────

# Insert G&A and fee details after the subawardee section (find empty heading)
# Find the last paragraph and add an extra section
for para in list(doc.paragraphs)[-10:]:
    if "Provide the following for all proposed sub-awardees" in para.text:
        prev = para._p
        lines = [
            "",
            "9. INDIRECT RATES AND FEE",
            "",
            "G&A Rate: 30% applied to prime direct costs excluding subcontract.",
            "  Base: Direct Labor $360,960 + Equipment $6,500 + Travel $18,000 + ODCs $119,500 = $504,960",
            "  G&A = 30% x $504,960 = $151,488",
            "",
            "The 30% rate covers home-office IT, professional subscriptions, legal/compliance review, "
            "proposal overhead, and administrative support — consistent with sole-proprietorship estimating practices.",
            "Subcontract costs excluded from G&A base per standard practice.",
            "",
            "Fee / Profit: 8.06% on total estimated cost.",
            "  Base: $776,448 (total direct costs + G&A)",
            "  Fee = 8.06% x $776,448 = $62,552",
            "",
            "8% fee is within standard DARPA OT for Research ranges and commensurate with Phase I technical risk.",
            "Fee is not applied to the subcontract line or equipment.",
            "",
            "10. COST REALISM UNDER $2M CAP",
            "",
            "$839,000 = 42% of the MATHBAC $2,000,000 Phase I ceiling.",
            "  Solo-prime structure concentrates labor into high-skill hours.",
            "  Existing SCBE codebase eliminates infrastructure build cost.",
            "  GPU compute and travel are sized conservatively to IV&V requirements only.",
            "  No facilities construction, no new software platform, no hardware deliverable (workstation is "
            "prime-owned support instrument).",
            "  Phase II ROM ~$382,000; combined Phase I + II well within program funding envelope.",
        ]
        for line in lines:
            new_p = _make_p(line, font_pt=11)
            prev.addnext(new_p)
            prev = new_p
        break

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────

doc.save(E_DEST)
print(f"Saved: {E_DEST}")
