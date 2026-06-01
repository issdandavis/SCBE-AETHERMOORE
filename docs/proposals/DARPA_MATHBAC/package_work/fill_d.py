"""Fill Attachment D (Volume I Technical & Management) from official template."""
import shutil
import copy
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from lxml import etree
from docx.oxml.ns import qn

D_SRC = (
    "docs/proposals/DARPA_MATHBAC/official_templates/"
    "Attachment_D_Proposal_Instructions_and_Volume_I_Template__Technical_and_Management.docx"
)
D_DEST = (
    "docs/proposals/DARPA_MATHBAC/package_work/04_attachment_d_vol_i_technical_management/"
    "Attachment_D_SCBE_FILLED.docx"
)

shutil.copy2(D_SRC, D_DEST)
doc = Document(D_DEST)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _make_p(text, style_val="Normal", bold=False, italic=False, font_pt=11):
    p = etree.Element(qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), style_val)
    r = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_pt * 2))
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(font_pt * 2))
    if bold:
        etree.SubElement(rPr, qn("w:b"))
    if italic:
        etree.SubElement(rPr, qn("w:i"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def insert_after(ref_p_elem, new_p_elem):
    """Insert new_p_elem immediately after ref_p_elem in the document body."""
    ref_p_elem.addnext(new_p_elem)


def replace_instruction_and_insert(doc, instruction_start, content_lines):
    """
    Find the first paragraph whose text starts with instruction_start,
    replace its text with content_lines[0], and insert remaining lines after it.
    Returns the last inserted paragraph element.
    """
    for para in doc.paragraphs:
        if para.text.strip().startswith(instruction_start.strip()[:40]):
            # Replace runs with first content line
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            run = para.add_run(content_lines[0])
            run.font.size = Pt(11)
            # Insert remaining paragraphs after this one
            prev = para._p
            for line in content_lines[1:]:
                new_p = _make_p(line, font_pt=11)
                prev.addnext(new_p)
                prev = new_p
            return prev
    return None


def set_cell(tbl, row, col, text):
    """Set text in table cell, preserving paragraph structure."""
    cell = tbl.cell(row, col)
    # Clear existing paragraphs
    for para in cell.paragraphs:
        for r in list(para.runs):
            r._r.getparent().remove(r._r)
    # Set text in first paragraph
    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(11)
    return cell


# ─────────────────────────────────────────────────────────────────────────────
# Cover Sheet (Table 1 — 20 rows x 2 cols)
# ─────────────────────────────────────────────────────────────────────────────

cover = doc.tables[1]  # Table index 1 is the cover sheet

# Row 0: Proposal Title
set_cell(cover, 0, 1, "SCBE Mathematical Framework for Agentic Communication Protocols")

# Row 1: Proposer Organization
set_cell(cover, 1, 1, "ISSAC D DAVIS / SCBE-AETHERMOORE")

# Row 2: Topic Area (TA)
set_cell(cover, 2, 1, "TA1")

# Row 3: Type of Organization
set_cell(cover, 3, 1, "Other Small Business (sole proprietorship, minority-owned)")

# Row 4: Proposer Reference Number
set_cell(cover, 4, 1, "N/A")

# Row 5: Technical POC
set_cell(
    cover, 5, 1,
    "Name: Issac Davis\n"
    "Address: 2361 E 5th Ave, Port Angeles, WA 98362-9014\n"
    "Telephone: (360) 808-0876\n"
    "Email: issdandavis7795@gmail.com"
)

# Row 6: Administrative POC (same)
set_cell(
    cover, 6, 1,
    "Name: Issac Davis\n"
    "Address: 2361 E 5th Ave, Port Angeles, WA 98362-9014\n"
    "Telephone: (360) 808-0876\n"
    "Email: issdandavis7795@gmail.com"
)

# Row 7: Place(s) of Performance
set_cell(cover, 7, 1, "2361 E 5th Ave, Port Angeles, WA 98362-9014 (remote / local)")

# Row 8: Period(s) of Performance
set_cell(cover, 8, 1, "16 months from award date")

# Row 9: Months
set_cell(cover, 9, 1, "16")

# Row 10: Other Team Members
set_cell(
    cover, 10, 1,
    "Technical POC Name: Collin Hoag\n"
    "Organization: Hoags Inc.\n"
    "Organization Type: Supporting Subcontractor (bounded scope)"
)

# Row 11: Total Price
set_cell(
    cover, 11, 1,
    "Year 1: $[fill when award date / FY boundary known]\n"
    "Year 2: $[remainder]\n"
    "Year 3: N/A\n"
    "Total: $839,000"
)

# Row 12: SAM.gov UEI
set_cell(cover, 12, 1, "J4NXHM6N5F59")

# Row 13: TIN — LEFT BLANK for manual entry in Word
set_cell(cover, 13, 1, "[Enter TIN directly in Word — do not record here]")

# Row 14: CAGE
set_cell(cover, 14, 1, "1EXD5")

# Row 15: DCMA / ONR POC
set_cell(cover, 15, 1, "Name: Unknown at proposal stage\nAddress: —\nTelephone: —")

# Row 16: DCAA POC
set_cell(cover, 16, 1, "Name: Unknown at proposal stage\nAddress: —\nTelephone: —")

# Row 17: Date Prepared
set_cell(cover, 17, 1, "2026-06-[fill on submission day]")

# Row 18: AI-Generated Material
set_cell(
    cover, 18, 1,
    "Yes. Claude (Anthropic) was used to assist in drafting and formatting sections of this "
    "proposal. All technical content, mathematical derivations, benchmark evidence, and "
    "strategic decisions are authored by Issac Davis / SCBE-AETHERMOORE."
)

# Row 19: Proposal Validity Period
set_cell(cover, 19, 1, "120 days from submission")

# ─────────────────────────────────────────────────────────────────────────────
# Personnel / Key Team Table (Table 2 — 5 rows x 3 cols)
# ─────────────────────────────────────────────────────────────────────────────

team = doc.tables[2]

# Row 1: Prime — Individual Name / Organization / Foreign national? FFRDC?
set_cell(team, 1, 0, "Issac Davis")
set_cell(team, 1, 1, "ISSAC D DAVIS / SCBE-AETHERMOORE")
set_cell(
    team, 1, 2,
    "Non-U.S. Organization: No\nNon-U.S. Individual: No\nFFRDC: No"
)

# Row 2: Prime (secondary — leave blank or repeat)
set_cell(team, 2, 0, "")
set_cell(team, 2, 1, "")
set_cell(team, 2, 2, "Non-U.S. Organization: No\nNon-U.S. Individual: No\nFFRDC: No")

# Row 4: Subawardee
set_cell(team, 4, 0, "Collin Hoag")
set_cell(team, 4, 1, "Hoags Inc.")
set_cell(
    team, 4, 2,
    "Non-U.S. Organization: No\nNon-U.S. Individual: No\nFFRDC: No"
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 1: Official Transmittal Letter
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Attach the official transmittal letter",
    [
        "Issac Davis / SCBE-AETHERMOORE",
        "2361 E 5th Ave, Port Angeles, WA 98362-9014",
        "issdandavis7795@gmail.com | (360) 808-0876",
        "",
        "DARPA / MATHBAC Program",
        "DARPA-PA-26-05 | Full Proposal Reference: DARPA-PA-26-05-MATHBAC-PA-010",
        "",
        "I am pleased to submit this full proposal in response to DARPA-PA-26-05 (MATHBAC). "
        "SCBE-AETHERMOORE proposes a 16-month Phase I effort at $839,000 to develop a formal "
        "mathematical framework for governing multi-agent AI communication dynamics using "
        "hyperbolic geometry. All work will be performed at the prime organization under my "
        "direction as Principal Investigator.",
        "",
        "Issac Davis, Principal Investigator",
        "ISSAC D DAVIS / SCBE-AETHERMOORE | UEI J4NXHM6N5F59 | CAGE 1EXD5",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 2: Proposal Summary
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Provide a top-level synopsis",
    [
        "SCBE-AETHERMOORE proposes to develop a composed-operator governance framework "
        "T = L₁₄ ∘ ⋯ ∘ L₁ for multi-agent AI communication in scientific discovery tasks. "
        "The framework maps agent communication acts to points in a Poincaré ball (hyperbolic space) "
        "and measures safety by hyperbolic distance from the safe-operation origin. Safety cost "
        "scales exponentially with adversarial drift, making attacks computationally infeasible.",
        "",
        "Current state-of-the-art multi-agent AI safety relies on post-hoc filtering, red-teaming, "
        "and RLHF fine-tuning. None of these approaches provide a formal mathematical characterization "
        "of when a multi-agent collective is communicating correctly versus drifting toward misalignment. "
        "Spera (arXiv:2603.15973, 2026) proves that 42.6% of 900 real multi-agent trajectories have "
        "conjunctive safety dependencies that flat classifiers cannot detect.",
        "",
        "Success will enable DARPA to field mathematically-auditable multi-agent AI systems with "
        "verifiable communication integrity across scientific subdomains (NMR spectroscopy, "
        "cheminformatics, and Phase II: materials discovery and theorem-proving).",
        "",
        "Cost: $839,000 | Duration: 16 months | Phase II ROM: ~$382,000.",
        "",
        "The key innovation is governance as geometry. Five physical axioms (Unitarity, Locality, "
        "Causality, Symmetry, Composition) are enforced layer-by-layer; the composition is the "
        "governance certificate. Pre-submission benchmark evidence: 13/13 terminal-bench parity, "
        "172/173 Petri adversarial classification (99.42%), 5/5 real-patch task harness wins.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 3: Goals and Impact
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Describe what you are trying to achieve",
    [
        "Goal: Develop and empirically validate a formal mathematical framework for governing "
        "multi-agent AI communication dynamics. The framework treats governance as a geometric "
        "property of the collective's communication trajectory in Poincaré ball space, not as a "
        "post-hoc classifier applied to individual messages.",
        "",
        "Phase I targets: (1) Lyapunov convergence constants η, b per subdomain (M3). "
        "(2) CDPTI ≥ 95% recall, ≤ 5% FPR vs. Mixtral-8x7B baseline (M13). "
        "(3) Four proposer-added metrics (MEE, ACV, CDPTI, PIS) computable from IV&V bundle at every milestone. "
        "(4) NMR validation: Task A (1H NMR→structure) + Task B (2D COSY), Karplus rediscovery "
        "demonstrated through governed agent communication. "
        "(5) Subdomain agnosticism demonstrated on cheminformatics secondary lane (M9).",
        "",
        "Impact if successful: Scientifically reliable multi-agent AI systems with verifiable "
        "communication integrity. The composed-operator framework provides a common mathematical "
        "substrate for TA1 and feeds directly into TA2’s principle-extraction mandate. "
        "The NMR Karplus rediscovery experiment is a falsifiable, graded demonstration that governed "
        "multi-agent collectives can recover scientific principles through communication alone — "
        "without the principle appearing in any training context.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 4: Technical Plan
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Outline and address technical challenges",
    [
        "TECHNICAL APPROACH: COMPOSED OPERATOR GOVERNANCE",
        "",
        "The 14-layer pipeline T = L₁₄ ∘ ⋯ ∘ L₁ implements five physical axioms across layers: "
        "Unitarity (L2, L4, L7), Locality (L3, L8), Causality (L6, L11, L13), Symmetry (L5, L9, L10, L12), "
        "Composition (L1, L14). Each agent Aᵢ is modeled as Φᵢ : Tᵢ* → Bⁿᵢ, "
        "mapping token strings to points in the Poincaré ball via realification, LWS phi-weighting, "
        "and the radial exponential map.",
        "",
        "MC-1 (Agent Operator Model): One-step latent operator Sᵢ = Phaseᵢ ∘ Breatheᵢ ∘ Moveᵢ. "
        "Agent interactions scheduled by Juggling Scheduler (physics-based, 7 rules). "
        "Three interaction modes: output-only (b1), latent-access (b2), mixed (b3).",
        "",
        "MC-2 (Protocol Graph): Trust-weighted directed Laplacian L_H(G_t) over communication graph G_t. "
        "Edge weights w_ij(t) = sign(H_ij - H_min) · ‖log(S_ij)‖. CDPTI = Re(λ₂(L_H(G_t)))/Re(λ₂(L_H(G_t₀))). "
        "Cascade (monotone DAG) and recurrence (SCC ≥3 agents) detectors trigger ESCALATE.",
        "",
        "MC-3 (Harmonic Wall + Lyapunov): H(d,pd) = 1/(1 + φ·d_H + 2·pd) ∈ (0,1]. "
        "Lyapunov candidate V(t) = -log H(d_H(ψ(t), ψ_safe), pd(t)). "
        "Convergence: V(t+1) ≤ (1-η)V(t) + b. Constants η, b are M3 empirical deliverables.",
        "",
        "MC-4 (Protocol Optimization): Bicriteria MDP: minimize E_π[Σ_t c_t] s.t. Pr_π[success] ≥ τ "
        "and ACV(π) ≥ ACV_min. Lagrangian relaxation with subgradient updates. "
        "Step-cost vector c_k = (token, time, failure, energy, compliance) ∈ ℝ₅₊.",
        "",
        "MC-5 (Oracle): Triplet condition — MI ≥ I_min AND spectral coherence ≥ γ_min AND MEE ≥ MEE_min. "
        "All three thresholds are M3 calibration deliverables. Oracle is falsifiable: ablate any one "
        "check, true-positive rate drops measurably.",
        "",
        "Science Subdomain (TA1-(3)): Primary SSM ChemBERTa-77M (seyonec/ChemBERTa-zinc-base-v1, "
        "768-dim, probe layer 10/12). Orchestrator Qwen2.5-Coder-0.5B-Instruct. "
        "Task Family A: 1H NMR → structure (Karplus equation J = A·cos²θ + B·cosθ + C). "
        "Task Family B: 2D COSY cross-peak assignment. Validation lane: Hammett σ–ρ from reaction corpus.",
        "",
        "Risk mitigations: R1 (ChemBERTa probe noise) — fallback to output-logit operator. "
        "R2 (Mixtral outperforms raw accuracy) — SCBE value is governance-overhead-adjusted parity, not raw supremacy. "
        "R3 (IV&V challenge problem scope mismatch) — FAQ Q1 filed 2026-05-31; secondary lane + Phase II plan reduce exposure. "
        "Full risk register in §2.3 of this volume.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 5: Capabilities
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Describe the organizational experience",
    [
        "SCBE-AETHERMOORE is a sole proprietorship (UEI J4NXHM6N5F59, CAGE 1EXD5, SAM.gov Active, "
        "minority-owned small business) built by PI Issac Davis. The organization has independently "
        "developed the SCBE 14-layer pipeline (62+ TypeScript modules, 40+ Python modules), the HYDRA "
        "multi-agent orchestration system, and the Juggling Scheduler physics-based task-flight coordinator.",
        "",
        "Benchmark evidence (all pre-submission):",
        "  • 13/13 neutral-task oracle parity on terminal-bench-core-0.1.1 (commit f3fb4aa3c, zero governance overhead)",
        "  • 172/173 adversarial seeds correctly classified on Anthropic Petri benchmark (99.42%; 0.58% false-allow after regex v7, 2026-05-08)",
        "  • L13 RuntimeGate fast-path: 240 cases, p95 = 0.1095ms (target: <100ms, PASS)",
        "  • Real-patch task harness: 5/5 SCBE wins vs. 0/5 baseline on deterministic repair fixtures",
        "  • Semantic sphere benchmark: 26/26 surfaces, 505 iterations (2026-05-30T01:41:06Z)",
        "  • Bijective Sacred-Tongue encoder: 25/25 (100%) gate on deterministic test suite (commit 6538f4db)",
        "",
        "Prior art (publicly timestamped): The Six Tongues Protocol (KDP ASIN B0GSSFQD9G, 2025); "
        "USPTO Provisional Patent Application 63/961,403. "
        "Mathematical foundations independently derived: Lyapunov stability, CBF, port-Hamiltonian "
        "operator structure, Poincaré ball metric (Khalil; Marsden/Ratiu; Cannon).",
        "",
        "Prior government contracts: None. No OCI.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 6: Management Plan (heading found at P067)
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Provide a summary of the proposed team",
    [
        "Prime: ISSAC D DAVIS / SCBE-AETHERMOORE. PI Issac Davis holds sole technical and management authority.",
        "",
        "Supporting Subcontractor: Hoags Inc. (Collin Hoag; UEI DUHWVUXFNPV5 / CAGE 15XV5). "
        "Hoags Inc. provides DAVA bare-metal background-IP corroborating support, bounded to "
        "§5.1.6/§11.3 of the Phase I SOW and Annex A Part 2. Hoags is not PI or co-PI. "
        "Phase I is fully executable by SCBE without Hoags support.",
        "",
        "Milestone payment schedule (7 milestones, triggered by observable technical events):",
        "  M1 $84,000 (Mo. 1)  M2 $109,000 (Mo. 3)  M3 $151,000 (Mo. 6)",
        "  M4 $151,000 (Mo. 9)  M5 $151,000 (Mo. 13)  M6 $84,000 (Mo. 14)  M7 $109,000 (Mo. 16)",
        "",
        "Risk register summary: R1 ChemBERTa probe noise (med/med); R2 Mixtral outperforms raw accuracy (low/med); "
        "R3 IV&V scope mismatch (low/high); R4 MC-2 instrumentation delay (med/med); "
        "R5 $2M cap (med/high — addressed by $839K lean budget); R6 BAAT upload deadline (med/critical); "
        "R7 prior publication IP (low/med); R8 teaming agreement execution (med/med).",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 9: Personnel, Qualifications, and Commitments
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[List key personnel",
    [
        "PI: Issac Davis | ISSAC D DAVIS / SCBE-AETHERMOORE | 100% effort Phase I",
        "",
        "Qualifications:",
        "  • Sole developer of SCBE-AETHERMOORE AI safety/governance framework (62+ TypeScript + 40+ Python modules)",
        "  • Independent derivation: Lyapunov stability analysis, control-barrier functions, port-Hamiltonian operator structure, Poincaré ball metric",
        "  • HuggingFace: issdandavis (multiple trained/deployed models including phdm-21d-embedding, spiralverse-ai-federated-v1)",
        "  • Bijective Sacred-Tongue encoder: 25/25 gate; Petri adversarial: 172/173 (99.42%)",
        "  • Prior art: KDP ASIN B0GSSFQD9G; USPTO Provisional 63/961,403",
        "  • SAM.gov Active; UEI J4NXHM6N5F59; CAGE 1EXD5; sole prop; minority-owned",
        "  • No prior Government contracts; no OCI",
        "",
        "Subcontractor: Collin Hoag / Hoags Inc. (bounded supporting scope, DAVA background IP). "
        "Not a key person for Phase I technical deliverables.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 10: OCI Affirmations
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[In accordance with the requirements",
    [
        "10.a. Are any proposed team members currently receiving support from DARPA under a separate award? ☐ No (X)  ☐ Yes",
        "",
        "Answer: No. Neither Issac Davis nor Hoags Inc. is receiving DARPA support under a separate award.",
        "",
        "10.b. Did any proposed team member participate in developing the requirements for this solicitation? ☐ No (X)  ☐ Yes",
        "",
        "Answer: No.",
        "",
        "10.c. Are there other potential OCI concerns? ☐ No (X)  ☐ Yes",
        "",
        "Answer: No. SCBE-AETHERMOORE has no prior Government contracts and is a sole proprietorship "
        "with no overlapping government work that could create a conflict of interest.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 11: Novelty
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[Has the proposed work been submitted",
    [
        "Has the proposed work been submitted to any other Government solicitation? ☐ No (X)  ☐ Yes",
        "",
        "Has the proposed work already received funding or a positive funding decision? ☐ No (X)  ☐ Yes  ☐ Decision pending",
        "",
        "The composed-operator governance framework is novel and has not been submitted to any other "
        "Government solicitation. Prior public disclosures (KDP ASIN B0GSSFQD9G; USPTO Provisional "
        "63/961,403) establish timestamped prior art but do not represent Government funding.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 12: IP
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "Please provide the following information",
    [
        "Technical Data and Computer Software:",
        "",
        "IP restrictions: Yes. The following are asserted as Background IP (Government Purpose Rights proposed):",
        "  1. The Six Tongues Protocol — KDP ASIN B0GSSFQD9G (2025). Prior published work establishing the "
        "Langues Weighting System and Sacred Tongue phi-metric basis.",
        "  2. USPTO Provisional Patent Application 63/961,403 — SCBE composed-operator governance framework. "
        "Filing date establishes priority for the 14-layer pipeline architecture.",
        "  3. SCBE-AETHERMOORE open-source codebase (github.com/issdandavis/SCBE-AETHERMOORE). "
        "Government Purpose Rights are proposed for all Phase I deliverables developed under this agreement.",
        "",
        "Commercial software: No commercial software deliverables are proposed.",
        "",
        "Patents: USPTO Provisional Application 63/961,403 is owned by Issac Davis / SCBE-AETHERMOORE. "
        "No other patents encumber the proposed effort. FAQ Q3 filed 2026-05-31 to confirm DARPA "
        "treatment of prior published work.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Section 13: Bibliography
# ─────────────────────────────────────────────────────────────────────────────

replace_instruction_and_insert(
    doc,
    "[A brief (no more than 5 pages) bibliography",
    [
        "1. Khalil, H.K. (2002). Nonlinear Systems (3rd ed.). Prentice Hall. [Lyapunov Theorem 4.16; CBF]",
        "2. Marsden, J.E. and Ratiu, T.S. (1999). Introduction to Mechanics and Symmetry. Springer. [Port-Hamiltonian structure]",
        "3. Cannon, J.W. et al. (1997). Hyperbolic Geometry. MSRI. [Poincaré ball model]",
        "4. Poppi, G. et al. (2025). Hyperbolic Safety-Aware Vision-Language Models. CVPR 2025. arXiv:2503.12127.",
        "5. Cinà, G. et al. (2026). Harnessing Hyperbolic Geometry for Harmful Prompt Detection. arXiv:2604.06285.",
        "6. Spera, S. (2026). Safety is Non-Compositional: A Formal Framework for Capability-Based AI Systems. arXiv:2603.15973.",
        "7. Bhardwaj, A. (2026). Agent Behavioral Contracts. arXiv:2602.22302.",
        "8. Anthropic. Petri adversarial benchmark (173 seeds, 2026). Internal reference; results documented in SCBE-AETHERMOORE repo.",
        "9. terminal-bench-core-0.1.1. Commit f3fb4aa3c. SCBE-AETHERMOORE governance evidence brief.",
        "10. BMRB (Biological Magnetic Resonance Data Bank). https://bmrb.io/",
        "11. SDBS (Spectral Database for Organic Structure Compounds). https://sdbs.db.aist.go.jp/",
        "12. seyonec/ChemBERTa-zinc-base-v1. HuggingFace.",
        "13. Qwen/Qwen2.5-Coder-0.5B-Instruct. HuggingFace.",
        "14. mistralai/Mixtral-8x7B-Instruct-v0.1. HuggingFace.",
    ]
)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────

doc.save(D_DEST)
print(f"Saved: {D_DEST}")
