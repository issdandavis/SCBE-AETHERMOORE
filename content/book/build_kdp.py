"""
Build KDP-formatted Word document for "The Six Tongues Protocol" by Issac Davis.
6x9 trim paperback, Kindle Direct Publishing standards.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Configuration ──────────────────────────────────────────────

BOOK_TITLE = "The Six Tongues Protocol"
AUTHOR = "Issac Davis"
SUBTITLE = "Book One of The Six Tongues Protocol"

BASE_DIR = r"C:\Users\issda\SCBE-AETHERMOORE\content\book\reader-edition"
OUTPUT = r"C:\Users\issda\SCBE-AETHERMOORE\content\book\the-six-tongues-protocol-kdp.docx"

# Chapter files in reading order (chapters + interludes woven in)
CHAPTER_FILES = [
    "ch01.md",
    "interlude-01-pollys-vigil.md",
    "ch02.md", "ch03.md", "ch04.md", "ch05.md",
    "interlude-06-jorren-records.md",
    "ch06.md", "ch07.md",
    "interlude-04-brams-report.md",
    "ch08.md",
    "interlude-09-tovak-hides.md",
    "ch09.md",
    "interlude-02-the-garden-before.md",
    "ch10.md",
    "interlude-03-sennas-morning.md",
    "ch11.md", "ch12.md", "ch13.md",
    "interlude-07-nadia-runs.md",
    "ch14.md", "ch15.md", "ch16.md",
    "ch17.md", "ch18.md", "ch19.md", "ch20.md",
    "interlude-10-arias-garden.md",
    "ch21.md", "ch22.md", "ch23.md", "ch24.md",
    "interlude-08-the-pipe.md",
    "ch25.md",
    "interlude-05-alexanders-hold.md",
    "ch26.md", "ch27.md",
    "ch-rootlight.md",
]

# Page setup constants (6x9 trim)
PAGE_WIDTH = Inches(6)
PAGE_HEIGHT = Inches(9)
MARGIN_TOP = Inches(0.75)
MARGIN_BOTTOM = Inches(0.75)
MARGIN_INSIDE = Inches(0.875)
MARGIN_OUTSIDE = Inches(0.625)

# Font settings
BODY_FONT = "Georgia"
BODY_SIZE = Pt(11)
CHAPTER_TITLE_SIZE = Pt(18)
LINE_SPACING = Pt(14)
FIRST_LINE_INDENT = Inches(0.3)

# ── Helpers ────────────────────────────────────────────────────

def setup_page(section):
    """Configure page size and margins for 6x9 KDP trim."""
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_INSIDE
    section.right_margin = MARGIN_OUTSIDE
    # Mirror margins via sectPr XML
    sectPr = section._sectPr
    # Set mirror margins
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.set(qn('w:mirrorMargins'), '1')
    # Also set at document level
    mirrorEl = sectPr.find(qn('w:mirrorMargins'))
    if mirrorEl is None:
        mirrorEl = parse_xml(f'<w:mirrorMargins {nsdecls("w")} />')
        sectPr.append(mirrorEl)


def set_font(run, font_name=BODY_FONT, font_size=BODY_SIZE, bold=False, italic=False):
    """Apply font settings to a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic


def add_body_paragraph(doc, text, first_para=False, justify=True, indent=True):
    """Add a body text paragraph with proper formatting."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING
    if justify:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Widow/orphan control
    pPr = p._element.get_or_add_pPr()
    widow = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="1"/>')
    pPr.append(widow)

    if indent and not first_para:
        pf.first_line_indent = FIRST_LINE_INDENT
    else:
        pf.first_line_indent = Pt(0)

    # Parse inline formatting (bold, italic)
    add_formatted_runs(p, text)
    return p


def add_formatted_runs(paragraph, text):
    """Parse markdown inline formatting and add runs with proper font."""
    # Replace straight quotes with curly quotes
    text = smart_quotes(text)
    # Ensure em dashes
    text = text.replace('--', '\u2014')

    # Parse bold and italic markers
    # Pattern: ***bold italic***, **bold**, *italic*
    # We'll process bold+italic first, then bold, then italic
    parts = []
    i = 0
    while i < len(text):
        if i + 2 < len(text) and text[i:i+3] == '***':
            # Bold italic
            end = text.find('***', i + 3)
            if end != -1:
                parts.append(('bi', text[i+3:end]))
                i = end + 3
                continue
        if i + 1 < len(text) and text[i:i+2] == '**':
            # Bold
            end = text.find('**', i + 2)
            if end != -1:
                parts.append(('b', text[i+2:end]))
                i = end + 2
                continue
        if text[i] == '*':
            # Italic
            end = text.find('*', i + 1)
            if end != -1:
                parts.append(('i', text[i+1:end]))
                i = end + 1
                continue
        # Normal text - accumulate until next marker
        j = i
        while j < len(text):
            if text[j] == '*':
                break
            j += 1
        parts.append(('n', text[i:j]))
        i = j

    for style, content in parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        set_font(run,
                 bold=(style in ('b', 'bi')),
                 italic=(style in ('i', 'bi')))


def smart_quotes(text):
    """Convert straight quotes to curly/smart quotes."""
    # Double quotes
    result = []
    in_double = False
    in_single = False
    for i, ch in enumerate(text):
        if ch == '"':
            if in_double:
                result.append('\u201d')  # right double quote
                in_double = False
            else:
                result.append('\u201c')  # left double quote
                in_double = True
        elif ch == "'":
            # Check if it's an apostrophe (preceded by a letter)
            if i > 0 and text[i-1].isalpha():
                result.append('\u2019')  # right single quote / apostrophe
            elif in_single:
                result.append('\u2019')  # right single quote
                in_single = False
            else:
                result.append('\u2018')  # left single quote
                in_single = True
        else:
            result.append(ch)
    return ''.join(result)


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    # Insert a page break via XML
    br = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br)
    return p


def add_section_break(doc):
    """Add a scene separator (centered asterisks)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(18)
    pf.space_after = Pt(18)
    pf.line_spacing = LINE_SPACING
    run = p.add_run('*\u2003*\u2003*')  # asterisks with em spaces
    set_font(run)
    return p


def parse_chapter(filepath):
    """Parse a markdown chapter file, returning (title, list_of_elements).
    Elements are ('para', text), ('break',), or ('first_para', text).
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    title = None
    elements = []
    current_para = []

    def flush_para():
        nonlocal current_para
        if current_para:
            text = ' '.join(current_para).strip()
            if text:
                elements.append(('para', text))
            current_para = []

    for line in lines:
        stripped = line.strip()

        # Chapter title (# heading)
        if stripped.startswith('# ') and title is None:
            title = stripped[2:].strip()
            continue

        # Scene break
        if stripped == '* * *' or stripped == '***' or stripped == '---':
            flush_para()
            elements.append(('break',))
            continue

        # Empty line = paragraph break
        if not stripped:
            flush_para()
            continue

        # Regular text line
        current_para.append(stripped)

    flush_para()

    # Mark first paragraph
    if elements and elements[0][0] == 'para':
        elements[0] = ('first_para', elements[0][1])

    return title, elements


# ── Build Document ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # Remove default styles margin
    style = doc.styles['Normal']
    font = style.font
    font.name = BODY_FONT
    font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING

    # ── Page setup for first section ──
    section = doc.sections[0]
    setup_page(section)

    # ──────────────────────────────────
    # FRONT MATTER
    # ──────────────────────────────────

    # 1. HALF TITLE PAGE
    # Add spacing above
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run(BOOK_TITLE)
    set_font(run, font_size=Pt(24), bold=True)

    add_page_break(doc)

    # 2. TITLE PAGE
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(BOOK_TITLE)
    set_font(run, font_size=Pt(28), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(f'by {AUTHOR}')
    set_font(run, font_size=Pt(16))

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    set_font(run, font_size=Pt(12), italic=True)

    add_page_break(doc)

    # 3. COPYRIGHT PAGE
    for _ in range(18):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    copyright_text = (
        f"Copyright \u00a9 2026 {AUTHOR}. All rights reserved.\n\n"
        "No part of this book may be reproduced, distributed, or transmitted "
        "in any form or by any means, including photocopying, recording, or "
        "other electronic or mechanical methods, without the prior written "
        "permission of the author, except in the case of brief quotations "
        "embodied in critical reviews and certain other noncommercial uses "
        "permitted by copyright law.\n\n"
        "This is a work of fiction. Names, characters, places, and incidents "
        "are products of the author\u2019s imagination or are used fictitiously. "
        "Any resemblance to actual events, locales, or persons, living or dead, "
        "is entirely coincidental.\n\n"
        "First Edition, 2026\n\n"
        "Published by Issac Davis"
    )
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(copyright_text)
    set_font(run, font_size=Pt(9))

    add_page_break(doc)

    # 4. DEDICATION PAGE
    for _ in range(10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run('For the ones who stayed.')
    set_font(run, font_size=Pt(12), italic=True)

    add_page_break(doc)

    # 5. BLANK PAGE (to start Chapter 1 on right-hand page)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    add_page_break(doc)

    # ──────────────────────────────────
    # CHAPTERS
    # ──────────────────────────────────

    total_words = 0
    chapter_count = 0
    chapter_number = 0

    for fname in CHAPTER_FILES:
        fpath = os.path.join(BASE_DIR, fname)
        if not os.path.exists(fpath):
            print(f"  [SKIP] {fname} not found")
            continue

        title, elements = parse_chapter(fpath)
        if not title:
            print(f"  [SKIP] {fname} has no title")
            continue

        chapter_number += 1
        chapter_count += 1

        # Count words in this chapter
        chapter_words = 0
        for etype, *edata in elements:
            if etype in ('para', 'first_para'):
                # Strip markdown markers for word count
                raw = re.sub(r'\*+', '', edata[0])
                chapter_words += len(raw.split())
        total_words += chapter_words

        print(f"  Chapter {chapter_number}: {title} ({chapter_words:,} words)")

        # Page break before chapter (except handled by the blank page for ch1)
        if chapter_number > 1:
            add_page_break(doc)

        # Chapter title with spacing above
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(72)
        pf.space_after = Pt(36)
        pf.line_spacing = LINE_SPACING

        # Use the title from the markdown file
        run = p.add_run(title)
        set_font(run, font_size=CHAPTER_TITLE_SIZE, bold=True)

        # Widow/orphan control on title
        pPr = p._element.get_or_add_pPr()
        widow = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="1"/>')
        pPr.append(widow)

        # Chapter body
        is_first = True
        for element in elements:
            etype = element[0]
            if etype == 'first_para':
                add_body_paragraph(doc, element[1], first_para=True)
                is_first = False
            elif etype == 'para':
                add_body_paragraph(doc, element[1], first_para=is_first)
                is_first = False
            elif etype == 'break':
                add_section_break(doc)
                is_first = True  # First para after break has no indent

    # ──────────────────────────────────
    # BACK MATTER
    # ──────────────────────────────────

    add_page_break(doc)

    # ABOUT THE AUTHOR
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run('About the Author')
    set_font(run, font_size=CHAPTER_TITLE_SIZE, bold=True)

    about_text = (
        "Issac Davis is a writer and technologist who believes that the best stories "
        "are the ones that make you think differently about the systems you live inside. "
        "He writes fiction at the intersection of engineering, magic, and the stubborn "
        "human insistence that distributed consensus is worth the overhead.\n\n"
        "The Six Tongues Protocol is his first novel."
    )
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING
    run = p.add_run(about_text)
    set_font(run)

    add_page_break(doc)

    # ALSO BY
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run('Also by Issac Davis')
    set_font(run, font_size=CHAPTER_TITLE_SIZE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Coming Soon')
    set_font(run, font_size=Pt(12), italic=True)

    # ── Save ──
    doc.save(OUTPUT)

    print(f"\n{'='*60}")
    print(f"  KDP Document Built Successfully")
    print(f"{'='*60}")
    print(f"  Output: {OUTPUT}")
    print(f"  Chapters: {chapter_count}")
    print(f"  Total Words: {total_words:,}")
    print(f"  Format: 6\" x 9\" trim, KDP paperback")
    print(f"{'='*60}")

    return chapter_count, total_words


if __name__ == '__main__':
    build_document()
